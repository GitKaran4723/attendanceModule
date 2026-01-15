"""
Flask PWA Application - Main Entry Point
This is a Progressive Web App (PWA) built with Flask that can be installed on mobile devices.
The app uses SQLite database and is designed with a mobile-first approach.

Project Structure:
- app.py: Main application file with routes
- models.py: Database models and schema
- config.py: Configuration settings
- routes/: Modular route blueprints (for students to work on)
"""

from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, session, send_file
from datetime import datetime, date, time, timedelta
import os
import math
import json
import uuid
from io import BytesIO
from fee_helpers import (
    assign_fee_to_student,
    add_additional_fee,
    remove_additional_fee,
    get_student_fee_breakdown
)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import configuration
from config import Config

# Import database instance from models
from models import db

# Initialize Flask application
app = Flask(__name__)

# Load configuration
app.config.from_object(Config)

# Configure session for PWA persistence (for localhost/development)
app.config['SESSION_COOKIE_NAME'] = 'bca_bub_session'
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # Use Lax for local development
app.config['SESSION_COOKIE_SECURE'] = False  # False for HTTP (local), True for HTTPS (production)
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = 2592000  # 30 days in seconds

# Create instance folder if it doesn't exist
instance_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'instance')
os.makedirs(instance_path, exist_ok=True)

# Initialize database with app
db.init_app(app)

# Initialize Flask-Migrate
from flask_migrate import Migrate
migrate = Migrate(app, db)

# Import models after db initialization
from models import (
    gen_uuid, User, Role, Faculty, Student, Program, Section,
    Semester, Subject, SubjectAllocation, ClassSchedule,
    AttendanceSession, AttendanceRecord, Test, TestResult,
    WorkDiary, ImportLog, Unit, Chapter, Concept,
    CampusCheckIn, CollegeConfig, StudentSubjectEnrollment,
    FacultyAttendance, FeeStructure, FeeReceipt, Holiday, SystemSettings,
    ProgramCoordinator
)

# Import authentication utilities
from auth import (
    login_required, admin_required, faculty_required, student_required, role_required,
    login_user, logout_user, get_current_user, hash_password, verify_password,
    can_edit_work_diary, can_approve_work_diary, can_view_all_diaries,
    coordinator_required, is_coordinator
)

# Import fee management blueprint
from fee_routes import fee_bp

# Register blueprints
app.register_blueprint(fee_bp)

# Import and register enrollment blueprint
from routes.enrollment_routes import enrollment_bp
app.register_blueprint(enrollment_bp)

# Import and register coordinator routes
from coordinator_routes import register_coordinator_routes
register_coordinator_routes(app)

# Import and register coordinator admin routes
from coordinator_admin_routes import register_coordinator_admin_routes
# Import and register department routes
from department_routes import register_department_routes
register_department_routes(app)

register_coordinator_admin_routes(app)




# ============================================
# Custom Jinja2 Filters
# ============================================

@app.template_filter('section_display')
def section_display_filter(section, subject=None):
    """
    Display section name with fallback for subjects without sections.
    Returns section name if exists, otherwise returns 'Program-Subject' format.
    """
    if section and hasattr(section, 'section_name'):
        return section.section_name
    
    # Fallback: Use subject's program or default to 'BCA'
    if subject:
        if hasattr(subject, 'program') and subject.program:
            return f"{subject.program.program_code or subject.program.program_name}-Subject"
        elif hasattr(subject, 'subject_code'):
            return f"BCA-{subject.subject_code}"
    
    return "BCA-Subject"


def get_or_create_virtual_section(subject):
    """
    Get or create a virtual section for a subject that doesn't carry sections.
    Virtual sections are auto-created for subjects with carries_section=False.
    
    Args:
        subject: Subject object
        
    Returns:
        Section object (virtual section)
    """
    if subject.carries_section:
        return None  # Subject uses real sections, no virtual section needed
    
    # Check if virtual section already exists
    virtual_section = Section.query.filter_by(
        linked_subject_id=subject.subject_id,
        is_elective=True,  # Repurposing is_elective as is_virtual
        is_deleted=False
    ).first()
    
    if virtual_section:
        return virtual_section
    
    # Create virtual section
    program_code = subject.program.program_code if subject.program else "BCA"
    section_name = f"{program_code}-{subject.subject_code}"
    
    virtual_section = Section(
        section_name=section_name,
        program_id=subject.program_id if subject.program_id else None,
        current_semester=subject.semester_id,
        is_elective=True,  # Marking as virtual
        linked_subject_id=subject.subject_id
    )
    
    db.session.add(virtual_section)
    db.session.commit()
    
    return virtual_section




# ============================================
# Routes
# ============================================

# ============================================
# Authentication Routes
# ============================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """
    User login page and handler
    Supports student/parent login type selection for USN-based logins
    """
    # debug code
    print("Login attempt at", datetime.now())
    print("Corrent session:", dict(session))

    # If already logged in, redirect to appropriate dashboard
    current_user = get_current_user()

    if current_user:
        # Check if user is admin
        if current_user.role and current_user.role.role_name.lower() == 'admin':
            return redirect(url_for('admin_dashboard'))
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password')
        login_type = request.form.get('login_type', 'student')  # 'student' or 'parent'
        
        user = User.query.filter_by(username=username, is_active=True, is_deleted=False).first()

        print("User fetched:", user)
        
        if user and verify_password(user.password_hash, password):
            login_user(user)
            
            # Store login_type in session for parent view mode
            # This allows same credentials to access as student or as parent
            if login_type == 'parent' and user.role and user.role.role_name.lower() == 'student':
                session['login_as_parent'] = True
            else:
                session['login_as_parent'] = False

            
            # Redirect to next page or appropriate dashboard based on role
            next_page = request.args.get('next')
            # Validate next_page to prevent open redirect
            if next_page and next_page.startswith('/'):
                return redirect(next_page)
            
            # Admin users go to admin dashboard
            if user.role and user.role.role_name.lower() == 'admin':
                return redirect(url_for('admin_dashboard'))
            
            # Student users go to student dashboard (whether logged in as student or parent)
            if user.role and user.role.role_name.lower() == 'student':
                if session.get('login_as_parent'):
                    return redirect(url_for('parent_dashboard'))
                return redirect(url_for('student_dashboard'))
            
            # Legacy: Parent users with separate accounts also go to student dashboard
            if user.role and user.role.role_name.lower() == 'parent':
                return redirect(url_for('student_dashboard'))
                
            # Faculty users go to faculty dashboard
            if user.role and user.role.role_name.lower() == 'faculty':
                return redirect(url_for('faculty_dashboard_view'))
            
            return redirect(url_for('index'))
        else:
            return render_template('login.html', error='Invalid username or password')
    
    return render_template('login.html')



@app.route('/logout')
def logout():
    """
    User logout handler
    """
    logout_user()
    return redirect(url_for('login'))




@app.route('/')
def index():
    """
    Root route - redirects users based on their role
    """
    current_user = get_current_user()
    
    # If not logged in, redirect to login
    if not current_user:
        return redirect(url_for('login'))
        
    # Redirect based on role
    role_name = current_user.role.role_name.lower() if current_user.role else ''
    
    if role_name == 'admin':
        return redirect(url_for('admin_dashboard'))
    elif role_name == 'faculty':
        return redirect(url_for('faculty_dashboard_view'))
    elif role_name == 'student':
        return redirect(url_for('student_dashboard'))
    else:
        # Fallback
        return redirect(url_for('login'))


@app.route('/manifest.json')
def manifest():
    """
    Serve PWA manifest file for app installation
    """
    return app.send_static_file('manifest.json')


@app.route('/service-worker.js')
def service_worker():
    """
    Serve service worker JavaScript file for PWA functionality
    """
    return app.send_static_file('service-worker.js')


# ============================================
# Public Pages (No Authentication Required)
# ============================================

@app.route('/welcome')
def public_home():
    """
    Public landing page showcasing the department
    Accessible without login
    """
    faculty_count = Faculty.query.filter_by(is_deleted=False).count()
    student_count = Student.query.filter_by(is_deleted=False).count()
    
    return render_template('public_home.html',
                         faculty_count=faculty_count,
                         student_count=student_count)


@app.route('/admissions')
def admissions():
    """
    Admission process page
    Shows eligibility, steps, dates, and contact info
    """
    return render_template('admissions.html')


@app.route('/about')
def about():
    """
    About us page
    Shows department info, team members, and history
    """
    return render_template('about.html')


@app.route('/admin/faculty-list')
@admin_required
def admin_faculty_list():
    """
    Display all faculty members (Admin View)
    TODO for students: Add pagination, search, filters
    """
    faculties = Faculty.query.filter_by(is_deleted=False).all()
    return render_template('faculty.html', faculties=faculties)


@app.route('/admin/students-list')
@admin_required
def admin_students_list():
    """
    Display all students (Admin View)
    TODO for students: Add pagination, search, filters by section/program
    """
    students = Student.query.filter_by(is_deleted=False).limit(50).all()
    return render_template('students.html', students=students)


@app.route('/admin/attendance-overview')
@admin_required
def admin_attendance_overview():
    """
    Display attendance management page (Admin View)
    TODO for students: Implement full attendance taking workflow
    """
    # Get recent attendance sessions
    sessions = AttendanceSession.query.order_by(AttendanceSession.taken_at.desc()).limit(10).all()
    return render_template('attendance.html', sessions=sessions)


# ============================================
# API Endpoints - For Students to Expand
# ============================================

@app.route('/api/faculty', methods=['GET', 'POST'])
def api_faculty():
    """
    API endpoint for faculty management
    GET: List all faculty
    POST: Create new faculty
    TODO for students: Add filtering, pagination, validation
    """
    if request.method == 'POST':
        data = request.get_json()
        try:
            # This is simplified - students should add proper validation
            new_faculty = Faculty(
                name=data.get('name'),
                email=data.get('email'),
                phone=data.get('phone')
            )
            db.session.add(new_faculty)
            db.session.commit()
            return jsonify({'success': True, 'faculty': new_faculty.to_dict()}), 201
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 400
    
    # GET request
    faculties = Faculty.query.filter_by(is_deleted=False).all()
    return jsonify([f.to_dict() for f in faculties])


@app.route('/api/students', methods=['GET', 'POST'])
def api_students():
    """
    API endpoint for student management
    GET: List all students
    POST: Create new student
    TODO for students: Add validation, handle user creation, assign to sections
    """
    if request.method == 'POST':
        data = request.get_json()
        try:
            # Simplified version - students should expand this
            new_student = Student(
                usn=data.get('usn'),
                name=data.get('name'),
                email=data.get('email'),
                phone=data.get('phone')
            )
            db.session.add(new_student)
            db.session.commit()
            return jsonify({'success': True, 'student': new_student.to_dict()}), 201
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 400
    
    # GET request
    students = Student.query.filter_by(is_deleted=False).limit(100).all()
    return jsonify([s.to_dict() for s in students])


@app.route('/api/student/profile', methods=['PUT'])
@login_required
def update_student_profile():
    """
    Update student profile information.
    When date_of_birth is updated, automatically update password to match DOB in DDMMYYYY format.
    """
    from werkzeug.security import generate_password_hash
    from datetime import datetime
    
    current_user_obj = get_current_user()
    
    # Only students can update their own profile
    if not hasattr(current_user_obj, 'student') or not current_user_obj.student:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    student = current_user_obj.student
    data = request.get_json()
    
    try:
        # Update basic fields
        if 'phone' in data:
            student.phone = data['phone']
        
        if 'address' in data:
            student.address = data['address']
        
        # Handle date of birth update
        if 'date_of_birth' in data and data['date_of_birth']:
            dob_str = data['date_of_birth']
            
            # Parse the date
            try:
                new_dob = datetime.strptime(dob_str, '%Y-%m-%d').date()
                student.date_of_birth = new_dob
                
                # Update password to match new DOB in DDMMYYYY format (without dashes)
                password_str = new_dob.strftime('%d%m%Y')
                current_user_obj.password_hash = generate_password_hash(password_str)
                
            except ValueError:
                return jsonify({'success': False, 'error': 'Invalid date format'}), 400
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Profile updated successfully',
            'password_updated': 'date_of_birth' in data and data['date_of_birth']
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/attendance/session', methods=['POST'])
@faculty_required
def api_create_attendance_session():
    """
    Create a new attendance session and automatically generate work diary entry
    TODO for students: Implement full attendance workflow
    - Validate schedule exists
    - Auto-populate student list
    - Handle bulk status updates
    """
    data = request.get_json()
    try:
        current_user = get_current_user()
        
        # Create attendance session
        session = AttendanceSession(
            schedule_id=data.get('schedule_id'),
            taken_by_user_id=current_user.user_id,
            status='draft'
        )
        db.session.add(session)
        db.session.flush()  # Get session ID without committing
        
        # Auto-generate work diary entry if schedule exists
        if session.schedule:
            auto_create_work_diary_from_attendance(session)
        
        db.session.commit()
        return jsonify({'success': True, 'session': session.to_dict()}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/subjects/<subject_id>/sections', methods=['GET'])
@login_required
def get_subject_sections(subject_id):
    """
    Get sections for a subject, separated into core and virtual sections.
    This endpoint is used when taking attendance to dynamically load sections
    based on the selected subject.
    
    Returns:
        - core_sections: Regular departmental sections (is_elective=False)
        - virtual_sections: Virtual sections for electives (is_elective=True, linked to subject)
        - subject_carries_section: Whether the subject uses sections
        - subject_category: Type of subject (compulsory, language, specialization, elective)
    """
    try:
        subject = Subject.query.get(subject_id)
        if not subject:
            return jsonify({'error': 'Subject not found'}), 404
        
        core_sections = []
        virtual_sections = []
        
        # Get core sections (regular departmental sections)
        if subject.program_id and subject.semester_id:
            core_sections_query = Section.query.filter_by(
                program_id=subject.program_id,
                current_semester=subject.semester_id,
                is_elective=False,
                is_deleted=False
            ).all()
            core_sections = [s.to_dict() for s in core_sections_query]
        
        # Get virtual sections (for electives/specializations without sections)
        virtual_sections_query = Section.query.filter_by(
            linked_subject_id=subject_id,
            is_elective=True,
            is_deleted=False
        ).all()
        virtual_sections = [s.to_dict() for s in virtual_sections_query]
        
        # If subject doesn't carry sections but no virtual section exists, create one
        if not subject.carries_section and not virtual_sections:
            virtual_section = get_or_create_virtual_section(subject)
            if virtual_section:
                virtual_sections = [virtual_section.to_dict()]
        
        return jsonify({
            'core_sections': core_sections,
            'virtual_sections': virtual_sections,
            'subject_carries_section': subject.carries_section,
            'subject_category': subject.subject_category,
            'success': True
        })
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/settings/<setting_key>', methods=['GET'])
@login_required
def get_setting(setting_key):
    """
    Get a system setting value by key.
    Used to check if features like past attendance are enabled.
    """
    from models import SystemSettings
    
    setting = SystemSettings.query.filter_by(setting_key=setting_key).first()
    
    if setting:
        return jsonify({
            'key': setting.setting_key,
            'value': setting.setting_value,
            'description': setting.description,
            'success': True
        })
    else:
        return jsonify({
            'key': setting_key,
            'value': None,
            'success': True
        })


@app.route('/admin/settings')
@admin_required
def admin_settings():
    """
    Admin settings page to configure system-wide features.
    """
    return render_template('admin_settings.html')


@app.route('/api/admin/settings', methods=['POST'])
@admin_required
def update_setting():
    """
    Update a system setting.
    """
    from models import SystemSettings
    from datetime import datetime
    from flask_login import current_user
    
    data = request.get_json()
    setting_key = data.get('setting_key')
    setting_value = data.get('setting_value')
    
    if not setting_key:
        return jsonify({'success': False, 'error': 'Setting key is required'}), 400
    
    try:
        # Check if setting exists
        setting = SystemSettings.query.filter_by(setting_key=setting_key).first()
        
        if setting:
            # Update existing setting
            setting.setting_value = setting_value
            setting.updated_at = datetime.now()
            if hasattr(current_user, 'user_id'):
                setting.updated_by = current_user.user_id
        else:
            # Create new setting
            setting = SystemSettings(
                setting_key=setting_key,
                setting_value=setting_value,
                description=f'Setting for {setting_key}',
                updated_at=datetime.now()
            )
            if hasattr(current_user, 'user_id'):
                setting.updated_by = current_user.user_id
            db.session.add(setting)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Setting updated successfully',
            'setting': {
                'setting_key': setting.setting_key,
                'setting_value': setting.setting_value
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================
# Work Diary Routes
# ============================================

@app.route('/faculty/work-diary')
@faculty_required
def faculty_work_diary():
    """
    Display work diary entries based on actual attendance taken.
    Shows classes conducted date-wise and faculty-wise.
    """
    current_user = get_current_user()
    
    # Always fetch faculty record for navigation purposes
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    if not faculty_record:
        diaries = []
        return render_template('work_diary.html', diaries=diaries, faculty=None, is_class_teacher=False)
    
    # Base query for AttendanceSession
    query = AttendanceSession.query.join(ClassSchedule).filter(
        AttendanceSession.is_deleted == False
    )
    
    if not can_view_all_diaries():
        # Faculty view: only their own sessions
        query = query.filter(AttendanceSession.taken_by_user_id == current_user.user_id)
    
    # Get recent sessions
    sessions = query.order_by(AttendanceSession.taken_at.desc()).limit(100).all()
    
    # Process sessions into displayable diary entries
    diaries = []
    for idx, session in enumerate(sessions):
        # Fetch related data
        schedule = session.schedule
        subject = schedule.subject
        section = schedule.section
        faculty = session.taken_by.faculty
        
        # Calculate hours
        actual_hours = 0
        if schedule.start_time and schedule.end_time:
            # Simple duration calculation
            dummy_date = date.today()
            start_dt = datetime.combine(dummy_date, schedule.start_time)
            end_dt = datetime.combine(dummy_date, schedule.end_time)
            actual_hours = (end_dt - start_dt).total_seconds() / 3600
        
        # Logic for claiming hours: Lab period reduced by 3/4 (which means 0.75 of actual)
        is_lab = subject.subject_type and subject.subject_type.lower() in ['lab', 'practical']
        claiming_hours = actual_hours * 0.75 if is_lab else actual_hours
        
        # Particulars: Topic taught + Semester info
        particulars = session.topic_taught or "Regular Class"
        sem_info = f"({section.current_semester if section else '?'} Sem)"
        
        # Count students present and total
        # For subjects requiring enrollment (specialization/elective), only count enrolled students
        if subject and subject.subject_category in ['specialization', 'elective']:
            # Get enrolled student IDs for this subject
            enrolled_student_ids = db.session.query(StudentSubjectEnrollment.student_id).filter_by(
                subject_id=subject.subject_id,
                is_deleted=False
            ).all()
            enrolled_ids = [e[0] for e in enrolled_student_ids]
            
            # Count only enrolled students
            present_count = AttendanceRecord.query.filter(
                AttendanceRecord.attendance_session_id == session.attendance_session_id,
                AttendanceRecord.status == 'present',
                AttendanceRecord.student_id.in_(enrolled_ids)
            ).count()
            
            total_count = AttendanceRecord.query.filter(
                AttendanceRecord.attendance_session_id == session.attendance_session_id,
                AttendanceRecord.student_id.in_(enrolled_ids)
            ).count()
        else:
            # For regular subjects (compulsory/language), count all students in the section
            present_count = AttendanceRecord.query.filter_by(
                attendance_session_id=session.attendance_session_id,
                status='present'
            ).count()
            
            total_count = AttendanceRecord.query.filter_by(
                attendance_session_id=session.attendance_session_id
            ).count()

        diaries.append({
            'sl_no': idx + 1,
            'session_id': session.attendance_session_id,
            'diary_no': session.diary_number or f"DN-{session.attendance_session_id[:8].upper()}",
            'date': session.taken_at.date(),
            'particulars': f"{particulars} {sem_info}",
            'actual_hours': round(actual_hours, 2),
            'claiming_hours': round(claiming_hours, 2),
            'subject_code': subject.subject_code if subject else 'N/A',
            'faculty_name': f"{faculty.first_name} {faculty.last_name}" if faculty else session.taken_by.username,
            'subject_name': subject.subject_name if subject else 'Unknown Subject',
            'section_name': section.section_name if section else 'Unknown Section',
            'status': session.status,
            'is_lab': is_lab,
            'present_count': present_count,
            'total_count': total_count
        })
        
    # Check if Class Teacher
    managed_section = Section.query.filter_by(class_teacher_id=faculty_record.faculty_id).first()
    is_class_teacher = managed_section is not None
        
    return render_template('work_diary.html', 
                         diaries=diaries, 
                         is_class_teacher=is_class_teacher,
                         faculty=faculty_record)


@app.route('/faculty/work-diary/docx')
@faculty_required
def generate_work_diary_docx():
    """
    Generate DOCX report for work diary (?start_date=...&end_date=...&faculty=...).
    Uses date range filtering instead of month-based filtering.
    """
    from work_diary_helpers import (
        build_months_structure_from_sessions,
        filter_and_assign_sl,
        COMBINED_HEADER
    )
    
    current_user = get_current_user()
    start_date_str = request.args.get("start_date")
    end_date_str = request.args.get("end_date")
    faculty_param = request.args.get("faculty", "current")
    
    if not start_date_str or not end_date_str:
        return "Missing start_date or end_date parameter (format: 'YYYY-MM-DD')", 400
    
    try:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
    except ValueError:
        return "Invalid date format. Use YYYY-MM-DD", 400
    
    if start_date > end_date:
        return "Start date must be before or equal to end date", 400
    
    # Get faculty record
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    if not faculty_record:
        return "Faculty record not found", 404
    
    # Determine which faculty's diary to generate
    # Use ClassSchedule.date (class conducted date) as the primary filter
    # This is important because it represents when the class was actually conducted,
    # not when the attendance was submitted (which may differ for past attendance)
    if faculty_param == "current" or not can_view_all_diaries():
        # Current faculty only
        selected_faculty = f"{faculty_record.first_name} {faculty_record.last_name}"
        
        # Query sessions that have a schedule with a date
        query_with_schedule = AttendanceSession.query.join(
            ClassSchedule,
            AttendanceSession.schedule_id == ClassSchedule.schedule_id
        ).filter(
            AttendanceSession.is_deleted == False,
            AttendanceSession.taken_by_user_id == current_user.user_id,
            ClassSchedule.date >= start_date,
            ClassSchedule.date <= end_date
        )
        
        # Also include sessions without schedules (fallback to taken_at date)
        query_without_schedule = AttendanceSession.query.filter(
            AttendanceSession.is_deleted == False,
            AttendanceSession.taken_by_user_id == current_user.user_id,
            AttendanceSession.schedule_id == None,
            db.func.date(AttendanceSession.taken_at) >= start_date,
            db.func.date(AttendanceSession.taken_at) <= end_date
        )
        
        # Combine both queries
        sessions = query_with_schedule.union(query_without_schedule).order_by(
            AttendanceSession.taken_at.desc()
        ).limit(500).all()
    else:
        # Admin/coordinator can view all or specific faculty
        selected_faculty = faculty_param if faculty_param != "All" else "All"
        
        query_with_schedule = AttendanceSession.query.join(
            ClassSchedule,
            AttendanceSession.schedule_id == ClassSchedule.schedule_id
        ).filter(
            AttendanceSession.is_deleted == False,
            ClassSchedule.date >= start_date,
            ClassSchedule.date <= end_date
        )
        
        query_without_schedule = AttendanceSession.query.filter(
            AttendanceSession.is_deleted == False,
            AttendanceSession.schedule_id == None,
            db.func.date(AttendanceSession.taken_at) >= start_date,
            db.func.date(AttendanceSession.taken_at) <= end_date
        )
        
        sessions = query_with_schedule.union(query_without_schedule).order_by(
            AttendanceSession.taken_at.desc()
        ).limit(500).all()
    
    # Note: sessions variable is now populated, no need for separate query line
    
    # Build months structure (still using month grouping internally)
    # If no sessions found, create empty structure for the date range
    if not sessions:
        # Create empty template with the date range
        months = {}
        faculty_list = [selected_faculty]
        all_rendered_weeks = []
        
        # Generate week structure for the date range (Monday to Saturday)
        # Start from the actual start_date, not the first Monday
        current_date = start_date
        week_number = 1
        
        while current_date <= end_date:
            # Find the Monday of the current week
            week_start_monday = current_date - timedelta(days=current_date.weekday())
            
            # For the first week, start from the actual start_date
            if week_number == 1:
                week_start = start_date
            else:
                week_start = week_start_monday
            
            # Week ends on Saturday (6 days from Monday)
            week_end_saturday = week_start_monday + timedelta(days=5)
            
            # Clip week_end to the end_date
            week_end = min(week_end_saturday, end_date)
            
            # Only add week if it's valid
            if week_start <= end_date:
                # Create empty week with blank rows for manual entry
                all_rendered_weeks.append({
                    "week_number": week_number,
                    "week_start": week_start_monday,
                    "week_end": week_end_saturday,
                    "display_start": week_start,
                    "display_end": week_end,
                    "entries": [],  # Empty entries - will create blank rows in DOCX
                    "week_total_actual": 0.0,
                    "week_total_claiming": 0.0
                })
                week_number += 1
            
            # Move to next Monday (7 days forward from the Monday of current week)
            current_date = week_start_monday + timedelta(days=7)
            
            # Safety: limit to 10 weeks max for any single request
            if week_number > 10:
                break
    else:
        months, faculty_list = build_months_structure_from_sessions(sessions)
        
        if not months:
            # Fallback to empty template
            all_rendered_weeks = []
        else:
            # Process all months in the date range
            all_rendered_weeks = []
            for month_label in sorted(months.keys(), key=lambda m: datetime.strptime(m, "%B %Y")):
                rendered_weeks = filter_and_assign_sl(months, month_label, selected_faculty)
                all_rendered_weeks.extend(rendered_weeks)
    
    # Create DOCX with A4 portrait orientation
    doc = Document()
    
    # Set A4 portrait with narrow margins
    sections = doc.sections
    for section in sections:
        section.page_height = Pt(842)       # A4 height (297mm) 
        section.page_width = Pt(595)        # A4 width (210mm)
        section.left_margin = Pt(36)        # 0.5 inch
        section.right_margin = Pt(36)       # 0.5 inch
        section.top_margin = Pt(36)         # 0.5 inch
        section.bottom_margin = Pt(36)      # 0.5 inch
    
    # Configure default style - Times New Roman 11pt
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Header section matching the bcaofficial format
    h = doc.add_paragraph()
    run = h.add_run("BANGALORE UNIVERSITY")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.underline = True
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.space_after = Pt(6)
    
    # Department and Annexure line
    dept_line = doc.add_paragraph()
    dept_run = dept_line.add_run("Department:     BCA")
    dept_run.bold = True
    dept_run.font.size = Pt(11)
    dept_run.font.name = 'Times New Roman'
    dept_run.underline = True
    
    # Add tabs/spaces for annexure
    dept_line.add_run("                    ")
    annex_run = dept_line.add_run("ANNEXURE (time table need to be attached)")
    annex_run.bold = True
    annex_run.font.size = Pt(11)
    annex_run.font.name = 'Times New Roman'
    annex_run.underline = True
    dept_line.space_after = Pt(3)
    
    # Workload line
    workload = doc.add_paragraph()
    workload.add_run("                                                                         ")
    workload_run = workload.add_run("Workload allotted per Week . . . . . .16 hours . . . . . .")
    workload_run.bold = True
    workload_run.font.size = Pt(11)
    workload_run.font.name = 'Times New Roman'
    workload.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    workload.space_after = Pt(8)
    
    # Create tables for each week
    cols = [
        "Sl.\nNo", "Diary\nNo.", "Date", "Particulars / chapter / lectures (as per\nTime Table) I / II / III / IV / V / VI Sem",
        "Actual\nhours", "Claiming hours\n(Lab period\nreduced by 3/4)", "Subject\ncode"
    ]
    
    # Optimized column widths for A4 portrait (595pt width - 72pt margins = ~523pt usable)
    col_widths_pt = [28, 35, 55, 250, 40, 70, 45]  # Total: ~523pt
    
    for w in all_rendered_weeks:
        week_label = f"Week {w['week_number']}: {w['display_start'].strftime('%d %b %Y')} — {w['display_end'].strftime('%d %b %Y')}"
        week_para = doc.add_paragraph(week_label)
        week_para.space_before = Pt(8)
        week_para.space_after = Pt(4)
        for run in week_para.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
        
        rows = []
        for e in w["entries"]:
            rows.append([
                str(e.get("SL No", "")),
                str(e.get("Dairy No.", "")),
                e.get("Date"),
                e.get(COMBINED_HEADER, ""),
                f"{e.get('Actual hours', 0):.1f}",
                f"{e.get('Claiming hours', 0):.1f}",
                e.get("Subject code", "")
            ])
        
        # If no entries, add blank rows for manual entry (10 rows per week)
        if not rows:
            for i in range(10):
                rows.append(["", "", "", "", "", "", ""])
        
        if not rows:
            continue
        
        table = doc.add_table(rows=1 + len(rows) + 1, cols=len(cols))
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths in points
        for i, width in enumerate(col_widths_pt):
            for row in table.rows:
                row.cells[i].width = Pt(width)
        
        # Header row styling - Times New Roman 9pt (smaller for portrait fit)
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(cols):
            p = hdr_cells[i].paragraphs[0]
            p.clear()
            run = p.add_run(c)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.space_before = Pt(2)
            p.space_after = Pt(2)
        
        # Data rows - Times New Roman 9pt with text wrapping (compact for portrait)
        for r_idx, row_data in enumerate(rows, start=1):
            row_cells = table.rows[r_idx].cells
            for c_idx, val in enumerate(row_data):
                p = row_cells[c_idx].paragraphs[0]
                p.clear()
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                p.space_before = Pt(1)
                p.space_after = Pt(1)
                # Center align numeric columns
                if c_idx in [0, 1, 4, 5]:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Enable text wrapping for long content
                row_cells[c_idx].width = Pt(col_widths_pt[c_idx])
        
        # Totals row - Times New Roman 9pt bold (compact for portrait)
        total_row_cells = table.rows[1 + len(rows)].cells
        for i in range(len(cols)):
            p = total_row_cells[i].paragraphs[0]
            p.clear()
            p.space_before = Pt(2)
            p.space_after = Pt(2)
            
            if i == 0:
                run = p.add_run("Weekly Total")
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            elif i == 4:
                run = p.add_run(f"{w['week_total_actual']:.1f}")
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif i == 5:
                run = p.add_run(f"{w['week_total_claiming']:.1f}")
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Calculate total claiming hours for all weeks
    total_claiming = sum(w['week_total_claiming'] for w in all_rendered_weeks)
    
    # Add footer section
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Total monthly working hours and remuneration
    total_line = doc.add_paragraph()
    total_hours_run = total_line.add_run(f"Total monthly working hours: {total_claiming:.1f}")
    total_hours_run.bold = True
    total_hours_run.font.size = Pt(11)
    total_hours_run.font.name = 'Times New Roman'
    
    total_line.add_run("                    ")
    
    total_renum_run = total_line.add_run(f"Total remuneration claiming - {total_claiming * 1000:.0f}/-")
    total_renum_run.bold = True
    total_renum_run.font.size = Pt(11)
    total_renum_run.font.name = 'Times New Roman'
    total_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    total_line.space_after = Pt(24)
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Date and Signature line
    date_sig = doc.add_paragraph()
    date_run = date_sig.add_run(f"Date: {datetime.now().strftime('%d / %m / %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.name = 'Times New Roman'
    
    date_sig.add_run("                                                                              ")
    sig_run = date_sig.add_run("Signature of the Guest Faculty")
    sig_run.font.size = Pt(11)
    sig_run.font.name = 'Times New Roman'
    date_sig.space_after = Pt(36)
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Certification text
    cert = doc.add_paragraph()
    cert.add_run("       ")
    cert_run = cert.add_run(
        "Certified that, the above Guest Faculty has been handled the Classes allotted to him/her as per the Time "
        "Table and as per the Attendance Record maintained in the department. The said dates and hours are is in order."
    )
    cert_run.font.size = Pt(11)
    cert_run.font.name = 'Times New Roman'
    cert.space_after = Pt(48)
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Chairman signature line
    chairman = doc.add_paragraph()
    chairman.add_run("                                                                                                                          ")
    chairman_run = chairman.add_run("Chairman / Chairperson")
    chairman_run.bold = True
    chairman_run.font.size = Pt(11)
    chairman_run.font.name = 'Times New Roman'
    
    # Save to BytesIO buffer
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    # Generate filename with faculty name and month/year
    # Format: FacultyName_MonthYear.docx (e.g., "John_Doe_January_2026.docx")
    month_year = start_date.strftime("%B_%Y")
    faculty_name_clean = selected_faculty.replace(" ", "_")
    filename = f"{faculty_name_clean}_{month_year}.docx"
    
    return send_file(bio, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')




@app.route('/work-diary/create', methods=['GET', 'POST'])
@faculty_required
def create_work_diary():
    """
    Create new work diary entry (for non-class activities)
    """
    if request.method == 'POST':
        return handle_work_diary_creation()
    
    # GET - show form
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    return render_template('work_diary_form.html', 
                         faculty=faculty_record,
                         action='create')


@app.route('/work-diary/<diary_id>/edit', methods=['GET', 'POST'])
@faculty_required
def edit_work_diary(diary_id):
    """
    Edit existing work diary entry
    """
    diary = WorkDiary.query.get_or_404(diary_id)
    
    if not can_edit_work_diary(diary):
        return jsonify({'error': 'Permission denied'}), 403
    
    if request.method == 'POST':
        return handle_work_diary_update(diary)
    
    return render_template('work_diary_form.html', 
                         diary=diary,
                         action='edit')


@app.route('/api/work-diary', methods=['GET', 'POST'])
@faculty_required
def api_work_diary():
    """
    API endpoint for work diary management
    GET: List work diaries (filtered by permissions)
    POST: Create new work diary entry
    """
    if request.method == 'POST':
        data = request.get_json()
        try:
            # Create work diary entry
            diary = create_work_diary_entry(data)
            return jsonify({'success': True, 'diary': diary.to_dict()}), 201
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 400
    
    # GET - retrieve diaries
    current_user = get_current_user()
    
    if can_view_all_diaries():
        diaries = WorkDiary.query.filter_by(is_deleted=False).order_by(
            WorkDiary.date.desc()
        ).limit(100).all()
    else:
        faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
        if faculty_record:
            diaries = WorkDiary.query.filter_by(
                faculty_id=faculty_record.faculty_id,
                is_deleted=False
            ).order_by(WorkDiary.date.desc()).limit(100).all()
        else:
            diaries = []
    
    return jsonify([d.to_dict() for d in diaries])


@app.route('/api/work-diary/<diary_id>', methods=['GET', 'PUT', 'DELETE'])
@faculty_required
def api_work_diary_detail(diary_id):
    """
    API endpoint for individual work diary operations
    """
    diary = WorkDiary.query.get_or_404(diary_id)
    
    if request.method == 'GET':
        return jsonify(diary.to_dict())
    
    elif request.method == 'PUT':
        if not can_edit_work_diary(diary):
            return jsonify({'error': 'Permission denied'}), 403
        
        data = request.get_json()
        try:
            update_work_diary_entry(diary, data)
            return jsonify({'success': True, 'diary': diary.to_dict()})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 400
    
    elif request.method == 'DELETE':
        if not can_edit_work_diary(diary):
            return jsonify({'error': 'Permission denied'}), 403
        
        try:
            diary.is_deleted = True
            db.session.commit()
            return jsonify({'success': True})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/work-diary/<diary_id>/submit', methods=['POST'])
@faculty_required
def submit_work_diary(diary_id):
    """
    Submit work diary entry for approval
    """
    diary = WorkDiary.query.get_or_404(diary_id)
    
    if not can_edit_work_diary(diary):
        return jsonify({'error': 'Permission denied'}), 403
    
    try:
        diary.status = 'submitted'
        diary.submitted_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True, 'message': 'Diary submitted for approval'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/work-diary/<diary_id>/approve', methods=['POST'])
@role_required('admin', 'hod')
def approve_work_diary(diary_id):
    """
    Approve work diary entry (Admin/HOD only)
    """
    diary = WorkDiary.query.get_or_404(diary_id)
    current_user = get_current_user()
    
    data = request.get_json()
    
    try:
        diary.status = 'approved'
        diary.approved_by = current_user.user_id
        diary.approved_at = datetime.utcnow()
        diary.approval_remarks = data.get('remarks', '')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Diary approved successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/work-diary/<diary_id>/reject', methods=['POST'])
@role_required('admin', 'hod')
def reject_work_diary(diary_id):
    """
    Reject work diary entry (Admin/HOD only)
    """
    diary = WorkDiary.query.get_or_404(diary_id)
    current_user = get_current_user()
    
    data = request.get_json()
    
    try:
        diary.status = 'rejected'
        diary.approved_by = current_user.user_id
        diary.approved_at = datetime.utcnow()
        diary.approval_remarks = data.get('remarks', 'Entry rejected')
        db.session.commit()
        return jsonify({'success': True, 'message': 'Diary rejected'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/admin/work-diary')
@admin_required
def admin_work_diary():
    """Admin view of all faculty work diaries"""
    from models import AttendanceSession, ClassSchedule, Faculty, Subject, Section, AttendanceRecord
    
    # Get all faculties for filter
    faculties = Faculty.query.order_by(Faculty.first_name).all()
    
    # Get all attendance sessions with details
    sessions = db.session.query(
        AttendanceSession,
        ClassSchedule,
        Subject,
        Section,
        Faculty
    ).join(
        ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
    ).join(
        Subject, ClassSchedule.subject_id == Subject.subject_id
    ).outerjoin(
        Section, ClassSchedule.section_id == Section.section_id
    ).join(
        Faculty, ClassSchedule.faculty_id == Faculty.faculty_id
    ).filter(
        AttendanceSession.is_deleted == False
    ).order_by(
        AttendanceSession.taken_at.desc()
    ).all()
    
    # Format diary entries
    diary_entries = []
    for session, schedule, subject, section, faculty in sessions:
        # Count present/absent
        records = AttendanceRecord.query.filter_by(attendance_session_id=session.attendance_session_id).all()
        present_count = sum(1 for r in records if r.status == 'present')
        absent_count = sum(1 for r in records if r.status == 'absent')
        
        diary_entries.append({
            'session_id': session.attendance_session_id,
            'diary_number': session.diary_number,
            'faculty_id': faculty.faculty_id,
            'faculty_name': f"{faculty.first_name} {faculty.last_name}",
            'date': schedule.date,
            'subject_name': subject.subject_name,
            'section_name': section.section_name if section else 'N/A',
            'topic': session.topic_taught or 'N/A',
            'present_count': present_count,
            'absent_count': absent_count
        })
    
    return render_template('admin_work_diary.html', 
                         diary_entries=diary_entries,
                         faculties=faculties)



# ============================================
# Admin Management Routes
# ============================================

@app.route('/admin')
@app.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """Admin dashboard with statistics"""
    from models import Unit, Chapter, Concept
    
    faculty_count = Faculty.query.filter_by(is_deleted=False).count()
    student_count = Student.query.filter_by(is_deleted=False).count()
    subject_count = Subject.query.filter_by(is_deleted=False).count()
    section_count = Section.query.filter_by(is_deleted=False).count()
    
    return render_template('admin_dashboard.html',
                         faculty_count=faculty_count,
                         student_count=student_count,
                         subject_count=subject_count,
                         section_count=section_count)


# ---------------------------------------------
# Faculty Management Routes
# ---------------------------------------------

@app.route('/admin/faculty')
@admin_required
def admin_faculty():
    """List all faculty with their subjects"""
    faculty_list = Faculty.query.filter_by(is_deleted=False).all()
    
    # Get subjects for each faculty
    for faculty in faculty_list:
        allocations = SubjectAllocation.query.filter_by(
            faculty_id=faculty.faculty_id,
            is_deleted=False
        ).all()
        faculty.subjects = [alloc.subject for alloc in allocations if alloc.subject and not alloc.subject.is_deleted]
    
    return render_template('admin_faculty.html', faculty_list=faculty_list, programs=Program.query.filter_by(is_deleted=False).order_by(Program.program_name).all())


@app.route('/admin/faculty/add', methods=['GET', 'POST'])
@admin_required
def admin_faculty_add():
    """Add new faculty member"""
    if request.method == 'POST':
        try:
            # Get form data
            employee_id = request.form.get('employee_id')
            first_name = request.form.get('first_name')
            last_name = request.form.get('last_name')
            email = request.form.get('email')
            phone = request.form.get('phone')
            program_id = request.form.get('program_id')
            designation = request.form.get('designation')
            qualification = request.form.get('qualification')
            username = request.form.get('username')
            password = request.form.get('password')
            subject_ids = request.form.getlist('subjects')
            
            # Check if employee_id or username already exists
            existing_faculty = Faculty.query.filter_by(employee_id=employee_id).first()
            if existing_faculty:
                return render_template('admin_faculty_form.html', 
                                     subjects=Subject.query.filter_by(is_deleted=False).all(),
                                     error='Employee ID already exists')
            
            existing_user = User.query.filter_by(username=username).first()
            if existing_user:
                return render_template('admin_faculty_form.html',
                                     subjects=Subject.query.filter_by(is_deleted=False).all(),
                                     error='Username already exists')
            
            # Create user account
            from werkzeug.security import generate_password_hash
            faculty_role = Role.query.filter_by(role_name='faculty').first()
            
            new_user = User(
                username=username,
                email=email,
                password_hash=generate_password_hash(password),
                role_id=faculty_role.role_id
            )
            db.session.add(new_user)
            db.session.flush()
            
            # Create faculty record
            new_faculty = Faculty(
                user_id=new_user.user_id,
                employee_id=employee_id,
                first_name=first_name,
                last_name=last_name,
                name=f"{first_name} {last_name}",
                email=email,
                phone=phone,
                program_id=program_id if program_id else None,
                designation=designation,
                qualification=qualification
            )
            db.session.add(new_faculty)
            db.session.flush()
            
            # Assign subjects with specific sections
            for subject_id in subject_ids:
                section_id = request.form.get(f'section_{subject_id}')
                allocation = SubjectAllocation(
                    subject_id=subject_id,
                    faculty_id=new_faculty.faculty_id,
                    section_id=section_id if section_id else None
                )
                db.session.add(allocation)
            
            db.session.commit()
            return redirect(url_for('admin_faculty'))
            
        except Exception as e:
            db.session.rollback()
            return render_template('admin_faculty_form.html',
                                 subjects=Subject.query.filter_by(is_deleted=False).all(),
                                 error=str(e))
    
    # GET request
    subjects = Subject.query.filter_by(is_deleted=False).all()
    programs = Program.query.filter_by(is_deleted=False).all()
    sections = Section.query.filter_by(is_deleted=False).all()
    
    return render_template('admin_faculty_form.html', 
                         subjects=subjects, 
                         programs=programs,
                         sections=[s.to_dict() for s in sections],
                         faculty_allocations={})


@app.route('/admin/faculty/<faculty_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_faculty_edit(faculty_id):
    """Edit existing faculty member"""
    faculty = Faculty.query.get_or_404(faculty_id)
    
    if request.method == 'POST':
        try:
            faculty.first_name = request.form.get('first_name')
            faculty.last_name = request.form.get('last_name')
            faculty.name = f"{faculty.first_name} {faculty.last_name}"
            faculty.email = request.form.get('email')
            faculty.phone = request.form.get('phone')
            program_id = request.form.get('program_id')
            faculty.program_id = program_id if program_id else None
            faculty.designation = request.form.get('designation')
            faculty.qualification = request.form.get('qualification')
            
            # Update workload hours per week
            workload = request.form.get('workload_hours_per_week', 0)
            faculty.workload_hours_per_week = int(workload) if workload else 0
            
            # Update subject allocations
            subject_ids = request.form.getlist('subjects')
            
            # Remove old allocations
            SubjectAllocation.query.filter_by(faculty_id=faculty_id).delete()
            
            # Add new allocations with sections
            for subject_id in subject_ids:
                section_id = request.form.get(f'section_{subject_id}')
                allocation = SubjectAllocation(
                    subject_id=subject_id,
                    faculty_id=faculty_id,
                    section_id=section_id if section_id else None
                )
                db.session.add(allocation)
            
            db.session.commit()
            return redirect(url_for('admin_faculty'))
            
        except Exception as e:
            db.session.rollback()
            subjects = Subject.query.filter_by(is_deleted=False).all()
            allocations = SubjectAllocation.query.filter_by(faculty_id=faculty_id).all()
            faculty_subject_ids = [a.subject_id for a in allocations]
            return render_template('admin_faculty_form.html', 
                                 faculty=faculty,
                                 subjects=subjects,
                                 faculty_subject_ids=faculty_subject_ids,
                                 error=str(e))
    
    # GET request
    subjects = Subject.query.filter_by(is_deleted=False).all()
    programs = Program.query.filter_by(is_deleted=False).all()
    sections = Section.query.filter_by(is_deleted=False).all()
    
    allocations = SubjectAllocation.query.filter_by(faculty_id=faculty_id, is_deleted=False).all()
    faculty_allocations = {a.subject_id: a.section_id for a in allocations}
    
    return render_template('admin_faculty_form.html', 
                         faculty=faculty,
                         subjects=subjects,
                         programs=programs,
                         sections=[s.to_dict() for s in sections],
                         faculty_subject_ids=list(faculty_allocations.keys()),
                         faculty_allocations=faculty_allocations)


@app.route('/api/admin/faculty/<faculty_id>', methods=['DELETE'])
@admin_required
def api_delete_faculty(faculty_id):
    """Delete faculty (soft delete)"""
    try:
        faculty = Faculty.query.get_or_404(faculty_id)
        faculty.is_deleted = True
        faculty.user.is_deleted = True
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/faculty/<faculty_id>/reset-password', methods=['POST'])
@admin_required
def api_reset_faculty_password(faculty_id):
    """Reset faculty username and password using email prefix (before @)"""
    try:
        faculty = Faculty.query.get_or_404(faculty_id)
        user = faculty.user
        
        if not user:
            return jsonify({'success': False, 'error': 'User account not found'}), 404
        
        if not faculty.email or '@' not in faculty.email:
            return jsonify({'success': False, 'error': 'Valid email required for reset'}), 400
        
        # Extract username from email (part before @)
        email_prefix = faculty.email.split('@')[0]
        
        # Reset username to email prefix
        user.username = email_prefix
        
        # Reset password to email prefix (hashed)
        user.password_hash = hash_password(email_prefix)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Username and password reset to: {email_prefix}',
            'username': email_prefix
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400




# ---------------------------------------------
# Student Management Routes
# ---------------------------------------------

@app.route('/admin/students')
@admin_required
def admin_students():
    """List all students"""
    students = Student.query.filter_by(is_deleted=False).all()
    programs = Program.query.filter_by(is_deleted=False).all()
    sections = Section.query.filter_by(is_deleted=False).all()
    
    total_students = len(students)
    sections_count = len(sections)
    programs_count = len(programs)
    
    return render_template('admin_students.html',
                         students=students,
                         programs=programs,
                         sections=sections,
                         total_students=total_students,
                         sections_count=sections_count,
                         programs_count=programs_count)


@app.route('/admin/students/add', methods=['GET', 'POST'])
@admin_required
def admin_student_add():
    """Add new student"""
    if request.method == 'POST':
        try:
            roll_number = request.form.get('roll_number')
            name = request.form.get('name') # Changed from first_name, last_name
            email = request.form.get('email')
            phone = request.form.get('phone')
            date_of_birth = request.form.get('date_of_birth')
            program_id = request.form.get('program_id')
            section_id = request.form.get('section_id')
            admission_year = request.form.get('admission_year')
            joining_academic_year = request.form.get('joining_academic_year')
            current_semester = request.form.get('current_semester')
            address = request.form.get('address')
            guardian_name = request.form.get('guardian_name')
            guardian_phone = request.form.get('guardian_phone')
            username = request.form.get('username')
            password = request.form.get('password')
            
            # Generate Temp USN if empty
            if not roll_number:
                import random
                import string
                # Generate unique temp ID: TMP-YYYY-XXXXX
                year_suffix = datetime.now().year
                rand_str = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
                roll_number = f"TMP-{year_suffix}-{rand_str}"
            
            # Auto-set username if empty
            if not username:
                username = roll_number

            # Check if roll_number or username exists
            existing_student = Student.query.filter_by(roll_number=roll_number).first()
            if existing_student:
                return render_template('admin_student_form.html',
                                     programs=Program.query.filter_by(is_deleted=False).all(),
                                     sections=Section.query.filter_by(is_deleted=False).all(),
                                     error='Roll number already exists')
            
            existing_user = User.query.filter_by(username=username).first()
            if existing_user:
                return render_template('admin_student_form.html',
                                     programs=Program.query.filter_by(is_deleted=False).all(),
                                     sections=Section.query.filter_by(is_deleted=False).all(),
                                     error='Username already exists')
            
            # Create user account
            from werkzeug.security import generate_password_hash
            student_role = Role.query.filter_by(role_name='student').first()
            
            new_user = User(
                username=username,
                email=email,
                password_hash=generate_password_hash(password),
                role_id=student_role.role_id
            )
            db.session.add(new_user)
            db.session.flush()
            
            # Create student record
            new_student = Student(
                user_id=new_user.user_id,
                roll_number=roll_number,
                usn=roll_number,
                name=name, # Changed from first_name, last_name, and name=f"{first_name} {last_name}"
                email=email,
                phone=phone,
                date_of_birth=datetime.strptime(date_of_birth, '%Y-%m-%d').date() if date_of_birth else None,
                program_id=program_id if program_id else None,
                section_id=section_id if section_id else None,
                admission_year=int(admission_year) if admission_year else None,
                joining_academic_year=joining_academic_year if joining_academic_year else None,
                current_semester=int(current_semester) if current_semester else None,
                address=address,
                guardian_name=guardian_name,
                guardian_phone=guardian_phone,
                category=request.form.get('category'),
                seat_type=request.form.get('seat_type'),
                quota_type=request.form.get('quota_type') if request.form.get('quota_type') else None,
                current_academic_year=request.form.get('current_academic_year')
            )
            db.session.add(new_student)
            db.session.flush()  # Get student ID before creating fee structure
            
            # Auto-create fee structure if minimum fee info provided
            if new_student.seat_type and new_student.joining_academic_year and new_student.current_academic_year:
                from fee_helpers import assign_fee_to_student
                from models import FeeStructure
                import json
                
                # Check if it already exists (unlikely in Add, but good for safety)
                fee_structure = FeeStructure.query.filter_by(
                    student_id=new_student.student_id,
                    academic_year=new_student.current_academic_year,
                    is_deleted=False
                ).first()
                
                if not fee_structure:
                    # Assign base fees
                    assign_result = assign_fee_to_student(new_student.student_id, current_user.user_id)
                    
                    if assign_result['success']:
                        # Get the created fee structure to add additional fees
                        fee_structure = FeeStructure.query.filter_by(
                            student_id=new_student.student_id,
                            academic_year=new_student.current_academic_year,
                            is_deleted=False
                        ).first()
                    else:
                        flash(f"Fee assignment failed: {assign_result.get('error')}", "warning")
                
                if fee_structure:
                        # Parse additional fees from form
                        additional_fees = []
                        fee_names = request.form.getlist('additional_fee_desc[]')
                        fee_amounts = request.form.getlist('additional_fee_amount[]')
                        
                        for name, amount in zip(fee_names, fee_amounts):
                            if name and amount:
                                additional_fees.append({
                                    'name': name,
                                    'amount': float(amount),
                                    'added_by': current_user.user_id,
                                    'added_at': datetime.now().isoformat()
                                })
                        
                        if additional_fees:
                            fee_structure.additional_fees = json.dumps(additional_fees)
                            # Recalculate total
                            additional_total = sum(f['amount'] for f in additional_fees)
                            fee_structure.total_fees = fee_structure.base_fees + additional_total
                            # balance is a calculated property
                            db.session.commit()
            
            db.session.commit()
            
            return redirect(url_for('admin_students'))
            
        except Exception as e:
            db.session.rollback()
            return render_template('admin_student_form.html',
                                 programs=Program.query.filter_by(is_deleted=False).all(),
                                 sections=Section.query.filter_by(is_deleted=False).all(),
                                 error=str(e))
    
    # GET request
    programs = Program.query.filter_by(is_deleted=False).all()
    sections = Section.query.filter_by(is_deleted=False).all()
    return render_template('admin_student_form.html', programs=programs, sections=sections)


@app.route('/admin/students/<student_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_student_edit(student_id):
    """Edit existing student"""
    student = Student.query.get_or_404(student_id)
    
    if request.method == 'POST':
        current_user = get_current_user()
        try:
            # Handle USN/Roll Number Update
            new_roll_number = request.form.get('roll_number')
            if new_roll_number and new_roll_number != student.roll_number:
                # Check for duplicate
                if Student.query.filter_by(roll_number=new_roll_number).first():
                    raise ValueError(f"Roll Number {new_roll_number} already exists")
                
                # Update username if it matches the old roll number
                if student.user.username == student.roll_number:
                    student.user.username = new_roll_number
                
                student.roll_number = new_roll_number
                student.usn = new_roll_number
            
            student.name = request.form.get('name') # Changed from first_name, last_name, and name=f"{student.first_name} {student.last_name}"
            student.email = request.form.get('email')
            student.phone = request.form.get('phone')
            
            date_of_birth = request.form.get('date_of_birth')
            if date_of_birth:
                student.date_of_birth = datetime.strptime(date_of_birth, '%Y-%m-%d').date()
            
            program_id = request.form.get('program_id')
            student.program_id = program_id if program_id else None
            
            section_id = request.form.get('section_id')
            student.section_id = section_id if section_id else None
            
            admission_year = request.form.get('admission_year')
            student.admission_year = int(admission_year) if admission_year else None
            
            joining_academic_year = request.form.get('joining_academic_year')
            student.joining_academic_year = joining_academic_year if joining_academic_year else None
            
            current_semester = request.form.get('current_semester')
            student.current_semester = int(current_semester) if current_semester else None
            
            student.current_academic_year = request.form.get('current_academic_year')
            
            student.address = request.form.get('address')
            student.guardian_name = request.form.get('guardian_name')
            student.guardian_phone = request.form.get('guardian_phone')
            
            # Update fee-related fields
            old_category = student.category
            old_seat_type = student.seat_type
            old_quota_type = student.quota_type
            
            student.category = request.form.get('category')
            student.seat_type = request.form.get('seat_type')
            student.quota_type = request.form.get('quota_type') if request.form.get('quota_type') else None
            
            # Check if fee-related fields changed
            fee_fields_changed = (old_category != student.category or 
                                old_seat_type != student.seat_type or 
                                old_quota_type != student.quota_type)
            
            # Update or create fee structure if minimum fee info provided
            # Skip fee assignment if student has left (status != 'active')
            if student.status == 'active' and student.seat_type and student.joining_academic_year and student.current_academic_year:
                from fee_helpers import assign_fee_to_student
                from models import FeeStructure
                import json
                
                # IMPORTANT: Commit student changes first so assign_fee_to_student reads updated values
                db.session.flush()
                
                # Check if it already exists for this year
                fee_structure = FeeStructure.query.filter_by(
                    student_id=student.student_id,
                    academic_year=student.current_academic_year,
                    is_deleted=False
                ).first()
                
                # Create or update fee structure from template if:
                # 1. Fee structure doesn't exist, OR
                # 2. Fee-related fields (seat_type, category, quota_type) changed
                if not fee_structure or fee_fields_changed:
                    assign_result = assign_fee_to_student(student.student_id, current_user.user_id)
                    if assign_result['success']:
                        fee_structure = FeeStructure.query.filter_by(
                            student_id=student.student_id,
                            academic_year=student.current_academic_year,
                            is_deleted=False
                        ).first()
                        if fee_fields_changed and fee_structure:
                            flash(f"Base fees updated from template: ₹{fee_structure.base_fees}", "info")
                    else:
                        flash(f"Fee assignment failed: {assign_result.get('error')}", "warning")
                
                if fee_structure:
                    # Parse additional fees from form
                    additional_fees = []
                    fee_names = request.form.getlist('additional_fee_desc[]')
                    fee_amounts = request.form.getlist('additional_fee_amount[]')
                    
                    for name, amount in zip(fee_names, fee_amounts):
                        if name and amount:
                            additional_fees.append({
                                'name': name,
                                'amount': float(amount),
                                'added_by': current_user.user_id,
                                'added_at': datetime.now().isoformat()
                            })
                    
                    # Sync additional fees
                    fee_structure.additional_fees = json.dumps(additional_fees) if additional_fees else None
                    
                    # Recalculate total
                    additional_total = sum(f['amount'] for f in additional_fees)
                    fee_structure.total_fees = fee_structure.base_fees + additional_total
                    # balance is a calculated property
                    db.session.flush()
            
            db.session.commit()
            flash('Student updated successfully' + (' with fee structure' if fee_structure else ''), 'success')
            return redirect(url_for('admin_students'))
            
        except Exception as e:
            db.session.rollback()
            return render_template('admin_student_form.html',
                                 student=student,
                                 programs=Program.query.filter_by(is_deleted=False).all(),
                                 sections=Section.query.filter_by(is_deleted=False).all(),
                                 error=str(e))
    
    # GET request
    current_user = get_current_user()
    programs = Program.query.filter_by(is_deleted=False).all()
    sections = Section.query.filter_by(is_deleted=False).all()
    return render_template('admin_student_form.html',
                         student=student,
                         programs=programs,
                         sections=sections,
                         current_user=current_user)


@app.route('/admin/fee-structure/<fee_structure_id>/delete', methods=['POST'])
@admin_required
def admin_delete_fee_structure(fee_structure_id):
    """Delete (hard delete) a fee structure"""
    try:
        from models import FeeStructure
        
        fee_structure = FeeStructure.query.get(fee_structure_id)
        if not fee_structure:
            return jsonify({'success': False, 'message': 'Fee structure not found'}), 404
        
        # Hard delete (not soft delete) to avoid unique constraint issues
        # when recreating fee structure for same student/academic year
        db.session.delete(fee_structure)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Fee structure deleted successfully'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

# Fee Management API Endpoints - Add to app.py after student edit route

# ---------------------------------------------
# Fee Management Routes
# ---------------------------------------------

@app.route('/api/admin/students/<student_id>/assign-fee', methods=['POST'])
@admin_required
def api_assign_fee_to_student(student_id):
    """Auto-assign fee to student based on joining year, academic year, and seat type"""
    from fee_helpers import assign_fee_to_student
    
    current_user = get_current_user()
    result = assign_fee_to_student(student_id, current_user.user_id)
    
    if result['success']:
        return jsonify(result), 200
    else:
        return jsonify(result), 400


@app.route('/api/admin/students/<student_id>/additional-fee', methods=['POST'])
@admin_required
def api_add_additional_fee(student_id):
    """Add additional fee to student"""
    from fee_helpers import add_additional_fee
    
    try:
        data = request.get_json()
        current_user = get_current_user()
        
        result = add_additional_fee(
            student_id=student_id,
            academic_year=data.get('academic_year'),
            fee_name=data.get('fee_name'),
            amount=data.get('amount'),
            remarks=data.get('remarks'),
            user_id=current_user.user_id
        )
        
        if result['success']:
            return jsonify(result), 200
        else:
            return jsonify(result), 400
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/students/<student_id>/additional-fee/<int:fee_index>', methods=['DELETE'])
@admin_required
def api_remove_additional_fee(student_id, fee_index):
    """Remove additional fee from student"""
    from fee_helpers import remove_additional_fee
    
    try:
        data = request.get_json()
        result = remove_additional_fee(
            student_id=student_id,
            academic_year=data.get('academic_year'),
            fee_index=fee_index
        )
        
        if result['success']:
            return jsonify(result), 200
        else:
            return jsonify(result), 400
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/students/<student_id>/fee-breakdown')
@login_required
def api_get_fee_breakdown(student_id):
    """Get fee breakdown for student (accessible by student, faculty, admin)"""
    from fee_helpers import get_student_fee_breakdown
    
    current_user = get_current_user()
    
    # Check permissions
    if current_user.role == 'student':
        # Students can only view their own fees
        from models import Student
        student = Student.query.filter_by(user_id=current_user.user_id).first()
        if not student or student.student_id != student_id:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    elif current_user.role == 'faculty':
        # Faculty can view fees of students in their class
        from models import Faculty, Student, Section
        faculty = Faculty.query.filter_by(user_id=current_user.user_id).first()
        if faculty:
            # Check if student is in faculty's class
            student = Student.query.get(student_id)
            if not student or not student.section_id:
                return jsonify({'success': False, 'error': 'Student not found'}), 404
            
            section = Section.query.get(student.section_id)
            if not section or section.class_teacher_id != faculty.faculty_id:
                return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    academic_year = request.args.get('academic_year')
    result = get_student_fee_breakdown(student_id, academic_year)
    
    if result['success']:
        return jsonify(result), 200
    else:
        return jsonify(result), 400

@app.route('/api/admin/fees/bulk-sync', methods=['POST'])
@admin_required
def api_bulk_sync_fees():
    """Bulk assign fees to all active students"""
    try:
        from fee_helpers import assign_fee_to_student
        from models import Student
        
        students = Student.query.filter_by(is_deleted=False).all()
        success_count = 0
        current_user = get_current_user()
        
        for student in students:
            if student.joining_academic_year and student.current_academic_year and student.seat_type:
                result = assign_fee_to_student(student.student_id, current_user.user_id)
                if result['success']:
                    success_count += 1
                    
        return jsonify({
            'success': True, 
            'message': f'Successfully synced fees for {success_count} students.',
            'total_students': len(students)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400

# ---------------------------------------------
# Semester Promotion Routes
# ---------------------------------------------

@app.route('/admin/students/<student_id>/reset-password', methods=['POST'])
@admin_required
def admin_reset_student_password(student_id):
    """Reset a student's password"""
    from werkzeug.security import generate_password_hash
    
    student = Student.query.get_or_404(student_id)
    
    try:
        new_password = request.form.get('new_password')
        if not new_password or len(new_password) < 6:
            return jsonify({'success': False, 'message': 'Password must be at least 6 characters'}), 400
        
        # Update password
        student.user.password_hash = generate_password_hash(new_password)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Password reset successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/admin/students/promote')
@admin_required
def admin_promote_students():
    """Bulk promote students to next semester"""
    program_id = request.args.get('program_id')
    current_semester = request.args.get('semester')
    
    students = []
    if program_id and current_semester:
        students = Student.query.filter_by(
            program_id=program_id,
            current_semester=int(current_semester),
            status='active',
            is_deleted=False
        ).all()
    
    programs = Program.query.filter_by(is_deleted=False).all()
    return render_template('admin_promote_students.html',
                         students=students,
                         programs=programs,
                         selected_program=program_id,
                         selected_semester=current_semester)


@app.route('/api/admin/students/promote', methods=['POST'])
@admin_required
def api_promote_students():
    """Promote students to next semester"""
    try:
        from models import StudentEnrollment
        
        data = request.get_json()
        student_ids = data.get('student_ids', [])
        
        promoted_count = 0
        for student_id in student_ids:
            student = Student.query.get(student_id)
            if student and student.current_semester < 8:  # Max 8 semesters
                # Mark current semester enrollments as completed
                db.session.query(StudentEnrollment).filter_by(
                    student_id=student_id,
                    semester_enrolled=student.current_semester,
                    enrollment_status='active'
                ).update({'enrollment_status': 'completed'})
                
                # Increment semester
                student.current_semester += 1
                promoted_count += 1
        
        db.session.commit()
        return jsonify({'success': True, 'promoted_count': promoted_count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/students/<student_id>', methods=['DELETE'])
@admin_required
def api_delete_student(student_id):
    """Delete student (soft or permanent)"""
    try:
        student = Student.query.get_or_404(student_id)
        permanent = request.args.get('permanent', 'false').lower() == 'true'
        
        if permanent:
            # Delete associated user if it exists
            if student.user:
                db.session.delete(student.user)
            # Delete student (cascades or manual depending on foreign keys)
            db.session.delete(student)
        else:
            # Soft delete
            student.is_deleted = True
            if student.user:
                student.user.is_deleted = True
                
        db.session.commit()
        return jsonify({'success': True, 'message': 'Student deleted ' + ('permanently' if permanent else 'softly')})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400




# ---------------------------------------------
# Section Management API Routes
# ---------------------------------------------

@app.route('/api/admin/sections', methods=['POST'])
@admin_required
def api_create_section():
    """Create new section"""
    try:
        data = request.get_json()
        
        # Create section
        new_section = Section(
            section_name=data.get('section_name'),
            program_id=data.get('program_id') if data.get('program_id') else None,
            academic_year=data.get('academic_year'),
            current_semester=int(data.get('current_semester')) if data.get('current_semester') else None,
            class_teacher_id=data.get('class_teacher_id') if data.get('class_teacher_id') else None
        )
        
        db.session.add(new_section)
        db.session.commit()
        
        return jsonify({'success': True, 'section': {
            'section_id': new_section.section_id,
            'section_name': new_section.section_name
        }}), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/sections/<section_id>', methods=['PUT'])
@admin_required
def api_update_section(section_id):
    """Update existing section"""
    try:
        section = Section.query.get_or_404(section_id)
        data = request.get_json()
        
        section.section_name = data.get('section_name', section.section_name)
        section.program_id = data.get('program_id') if data.get('program_id') else None
        section.academic_year = data.get('academic_year')
        section.current_semester = int(data.get('current_semester')) if data.get('current_semester') else None
        section.class_teacher_id = data.get('class_teacher_id') if data.get('class_teacher_id') else None
        
        db.session.commit()
        
        return jsonify({'success': True, 'section': {
            'section_id': section.section_id,
            'section_name': section.section_name
        }})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/sections/<section_id>', methods=['GET'])
@admin_required
def api_get_section(section_id):
    """Get section details"""
    try:
        section = Section.query.get_or_404(section_id)
        
        return jsonify({
            'success': True,
            'section': {
                'section_id': section.section_id,
                'section_name': section.section_name,
                'program_id': section.program_id,
                'academic_year': section.academic_year,
                'current_semester': section.current_semester,
                'class_teacher_id': section.class_teacher_id
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/sections/<section_id>', methods=['DELETE'])
@admin_required
def api_delete_section(section_id):
    """Delete section (soft delete)"""
    try:
        section = Section.query.get_or_404(section_id)
        section.is_deleted = True
        db.session.commit()
        
        return jsonify({'success': True})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# ---------------------------------------------
# Subject Management Routes (with Units/Chapters/Concepts)
# ---------------------------------------------

@app.route('/admin/subjects')
@admin_required
def admin_subjects():
    """List all subjects with hierarchy"""
    from models import Unit, Chapter, Concept
    
    subjects = Subject.query.filter_by(is_deleted=False).all()
    
    # Load full hierarchy for each subject
    for subject in subjects:
        subject.units = Unit.query.filter_by(subject_id=subject.subject_id, is_deleted=False).order_by(Unit.unit_number).all()
        for unit in subject.units:
            unit.chapters = Chapter.query.filter_by(unit_id=unit.unit_id, is_deleted=False).order_by(Chapter.chapter_number).all()
            for chapter in unit.chapters:
                chapter.concepts = Concept.query.filter_by(chapter_id=chapter.chapter_id, is_deleted=False).all()
    
    return render_template('admin_subjects.html', subjects=subjects, programs=Program.query.filter_by(is_deleted=False).order_by(Program.program_name).all())


@app.route('/admin/sections')
@admin_required
def admin_sections():
    """Manage sections"""
    sections = Section.query.filter_by(is_deleted=False).all()
    programs = Program.query.filter_by(is_deleted=False).all()
    faculties = Faculty.query.filter_by(is_deleted=False).order_by(Faculty.first_name).all()
    
    # Calculate student counts and get class teacher info
    for section in sections:
        section.student_count = Student.query.filter_by(section_id=section.section_id, is_deleted=False).count()
        
    return render_template('admin_sections.html', sections=sections, programs=programs, faculties=faculties)


@app.route('/admin/subjects/add', methods=['GET', 'POST'])
@admin_required
def admin_subject_add():
    """Add new subject"""
    if request.method == 'POST':
        try:
            subject_code = request.form.get('subject_code')
            subject_name = request.form.get('subject_name')
            description = request.form.get('description')
            program_id = request.form.get('program_id')
            semester_id = request.form.get('semester_id')
            credits = request.form.get('credits')
            subject_type = request.form.get('subject_type')
            total_hours = request.form.get('total_hours')
            
            # Check if subject code exists
            existing_subject = Subject.query.filter_by(subject_code=subject_code).first()
            if existing_subject:
                return render_template('admin_subject_form.html',
                                     programs=Program.query.filter_by(is_deleted=False).all(),
                                     error='Subject code already exists')
            
            new_subject = Subject(
                subject_code=subject_code,
                code=subject_code,
                subject_name=subject_name,
                description=description,
                program_id=program_id if program_id else None,
                semester_id=int(semester_id) if semester_id else None,
                credits=float(credits) if credits else 0,
                subject_type=subject_type,
                total_hours=int(total_hours) if total_hours else None,
                is_specialization=True if request.form.get('subject_category') in ['specialization', 'elective'] else False,
                subject_category=request.form.get('subject_category', 'compulsory'),
                elective_group=request.form.get('elective_group'),
                carries_section=True if request.form.get('carries_section') == 'true' else False
            )
            db.session.add(new_subject)
            db.session.commit()
            
            return redirect(url_for('admin_subjects'))
            
        except Exception as e:
            db.session.rollback()
            return render_template('admin_subject_form.html',
                                 programs=Program.query.filter_by(is_deleted=False).all(),
                                 error=str(e))
    
    # GET request
    programs = Program.query.filter_by(is_deleted=False).all()
    return render_template('admin_subject_form.html', programs=programs)


@app.route('/admin/subjects/<subject_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_subject_edit(subject_id):
    """Edit existing subject"""
    subject = Subject.query.get_or_404(subject_id)
    
    if request.method == 'POST':
        try:
            subject.subject_name = request.form.get('subject_name')
            subject.description = request.form.get('description')
            
            program_id = request.form.get('program_id')
            subject.program_id = program_id if program_id else None
            
            semester_id = request.form.get('semester_id')
            subject.semester_id = int(semester_id) if semester_id else None
            
            credits = request.form.get('credits')
            subject.credits = float(credits) if credits else 0
            
            subject.subject_type = request.form.get('subject_type')
            
            total_hours = request.form.get('total_hours')
            subject.total_hours = int(total_hours) if total_hours else None

            # Update Category and Group
            subject.subject_category = request.form.get('subject_category', 'compulsory')
            subject.elective_group = request.form.get('elective_group')
            
            # Update carries_section checkbox
            subject.carries_section = True if request.form.get('carries_section') == 'true' else False
            
            # Update legacy is_specialization flag based on category
            subject.is_specialization = subject.subject_category in ['specialization', 'elective']
            
            db.session.commit()
            return redirect(url_for('admin_subjects'))
            
        except Exception as e:
            db.session.rollback()
            return render_template('admin_subject_form.html',
                                 subject=subject,
                                 programs=Program.query.filter_by(is_deleted=False).all(),
                                 error=str(e))
    
    # GET request
    programs = Program.query.filter_by(is_deleted=False).all()
    return render_template('admin_subject_form.html', subject=subject, programs=programs)


@app.route('/api/admin/subjects/<subject_id>', methods=['DELETE'])
@admin_required
def api_delete_subject(subject_id):
    """Delete subject (soft delete)"""
    try:
        subject = Subject.query.get_or_404(subject_id)
        subject.is_deleted = True
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# ---------------------------------------------
# Subject Enrollment Management Routes
# ---------------------------------------------

@app.route('/api/admin/subjects/<subject_id>/auto-enroll', methods=['POST'])
@admin_required
def auto_enroll_subject(subject_id):
    """
    Auto-enroll all eligible students in a core subject
    Eligibility: Active students where current_semester matches subject.semester_id
    """
    try:
        from models import StudentEnrollment
        from datetime import datetime
        
        subject = Subject.query.get_or_404(subject_id)
        
        # Find eligible students  
        eligible_students = Student.query.filter_by(
            program_id=subject.program_id,
            current_semester=subject.semester_id,
            status='active',
            is_deleted=False
        ).all()
        
        enrolled_count = 0
        for student in eligible_students:
            # Check if already enrolled
            existing = db.session.query(StudentEnrollment).filter_by(
                student_id=student.student_id,
                subject_id=subject_id,
                is_deleted=False
            ).first()
            
            if not existing:
                # Get current academic year
                now = datetime.now()
                academic_year = f"{now.year}-{str(now.year + 1)[-2:]}" if now.month >= 6 else f"{now.year - 1}-{str(now.year)[-2:]}"
                
                enrollment = StudentEnrollment(
                    student_id=student.student_id,
                    subject_id=subject_id,
                    section_id=student.section_id,
                    enrollment_type='core',
                    enrollment_status='active',
                    semester_enrolled=student.current_semester,
                    academic_year=academic_year
                )
                db.session.add(enrollment)
                enrolled_count += 1
        
        db.session.commit()
        return jsonify({'success': True, 'enrolled_count': enrolled_count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/admin/subjects/<subject_id>/enroll-students')
@admin_required
def admin_enroll_students(subject_id):
    """Page for manually enrolling students in elective subjects"""
    from models import StudentEnrollment
    
    subject = Subject.query.get_or_404(subject_id)
    
    # Debug output
    print(f"DEBUG ENROLLMENT:")
    print(f"  Subject: {subject.subject_name}")
    print(f"  Program ID: {subject.program_id}")
    print(f"  Semester ID: {subject.semester_id} (type: {type(subject.semester_id)})")
    
    # Get active students matching program and semester
    eligible_students = Student.query.filter_by(
        program_id=subject.program_id,
        current_semester=subject.semester_id,
        status='active',
        is_deleted=False
    ).all()
    
    print(f"  Found {len(eligible_students)} eligible students")
    
    # Also check all students in program
    all_program_students = Student.query.filter_by(
        program_id=subject.program_id,
        status='active',
        is_deleted=False
    ).all()
    print(f"  Total students in program: {len(all_program_students)}")
    for s in all_program_students[:5]:  # Show first 5
        print(f"    - {s.roll_number}: semester={s.current_semester} (type: {type(s.current_semester)})")
    
    # Get already enrolled student IDs
    enrolled_ids_query = db.session.query(StudentEnrollment.student_id).filter_by(
        subject_id=subject_id,
        is_deleted=False
    )
    enrolled_student_ids = [row[0] for row in enrolled_ids_query.all()]
    
    # Filter available students
    available_students = [s for s in eligible_students if s.student_id not in enrolled_student_ids]
    
    # Get enrolled students with their enrollment details
    enrollments = db.session.query(StudentEnrollment).filter_by(
        subject_id=subject_id,
        is_deleted=False
    ).all()
    enrolled_students = [{'student': e.student, 'enrollment': e} for e in enrollments]
    
    return render_template('admin_enroll_students.html',
                         subject=subject,
                         available_students=available_students,
                         enrolled_students=enrolled_students)


@app.route('/api/admin/enrollments/add', methods=['POST'])
@admin_required
def api_add_enrollments():
    """Enroll multiple students in a subject"""
    try:
        from models import StudentEnrollment
        from datetime import datetime
        
        data = request.get_json()
        subject_id = data.get('subject_id')
        student_ids = data.get('student_ids', [])
        
        # Get current academic year
        now = datetime.now()
        academic_year = f"{now.year}-{str(now.year + 1)[-2:]}" if now.month >= 6 else f"{now.year - 1}-{str(now.year)[-2:]}"
        
        enrolled_count = 0
        for student_id in student_ids:
            student = Student.query.get(student_id)
            if student:
                enrollment = StudentEnrollment(
                    student_id=student_id,
                    subject_id=subject_id,
                    section_id=student.section_id,
                    enrollment_type='elective',
                    enrollment_status='active',
                    semester_enrolled=student.current_semester,
                    academic_year=academic_year
                )
                db.session.add(enrollment)
                enrolled_count += 1
        
        db.session.commit()
        return jsonify({'success': True, 'enrolled_count': enrolled_count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/enrollments/<enrollment_id>', methods=['DELETE'])
@admin_required
def api_delete_enrollment(enrollment_id):
    """Remove student enrollment"""
    try:
        from models import StudentEnrollment
        
        enrollment = db.session.query(StudentEnrollment).get_or_404(enrollment_id)
        enrollment.is_deleted = True
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# Unit/Chapter/Concept CRUD APIs
@app.route('/api/admin/units', methods=['POST'])
@admin_required
def api_create_unit():
    """Create a unit within a subject"""
    try:
        from models import Unit
        data = request.get_json()
        
        unit = Unit(
            subject_id=data['subject_id'],
            unit_number=data['unit_number'],
            unit_name=data['unit_name'],
            description=data.get('description')
        )
        db.session.add(unit)
        db.session.commit()
        return jsonify({'success': True, 'unit_id': unit.unit_id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/units/<unit_id>', methods=['DELETE'])
@admin_required
def api_delete_unit(unit_id):
    """Delete unit (soft delete)"""
    try:
        from models import Unit
        unit = Unit.query.get_or_404(unit_id)
        unit.is_deleted = True
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/chapters', methods=['POST'])
@admin_required
def api_create_chapter():
    """Create a chapter within a unit"""
    try:
        from models import Chapter
        data = request.get_json()
        
        chapter = Chapter(
            unit_id=data['unit_id'],
            chapter_number=data['chapter_number'],
            chapter_name=data['chapter_name'],
            description=data.get('description')
        )
        db.session.add(chapter)
        db.session.commit()
        return jsonify({'success': True, 'chapter_id': chapter.chapter_id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/chapters/<chapter_id>', methods=['DELETE'])
@admin_required
def api_delete_chapter(chapter_id):
    """Delete chapter (soft delete)"""
    try:
        from models import Chapter
        chapter = Chapter.query.get_or_404(chapter_id)
        chapter.is_deleted = True
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/concepts', methods=['POST'])
@admin_required
def api_create_concept():
    """Create a concept within a chapter"""
    try:
        from models import Concept
        data = request.get_json()
        
        concept = Concept(
            chapter_id=data['chapter_id'],
            concept_name=data['concept_name'],
            description=data.get('description')
        )
        db.session.add(concept)
        db.session.commit()
        return jsonify({'success': True, 'concept_id': concept.concept_id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/concepts/<concept_id>', methods=['DELETE'])
@admin_required
def api_delete_concept(concept_id):
    """Delete concept (soft delete)"""
    try:
        from models import Concept
        concept = Concept.query.get_or_404(concept_id)
        concept.is_deleted = True
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# ---------------------------------------------
# Program Management Routes
# ---------------------------------------------

@app.route('/admin/programs')
@admin_required
def admin_programs():
    """List all programs (departments)"""
    programs = Program.query.filter_by(is_deleted=False).all()
    
    # Calculate section counts for each program
    for program in programs:
        program.section_count = Section.query.filter_by(
            program_id=program.program_id, 
            is_deleted=False
        ).count()
    
    return render_template('admin_programs.html', programs=programs)


@app.route('/api/admin/programs', methods=['POST'])
@admin_required
def api_create_program():
    """Create a new program"""
    try:
        data = request.get_json()
        
        # Validate required fields
        program_code = data.get('program_code')
        program_name = data.get('program_name')
        
        if not program_code or not program_name:
            return jsonify({
                'success': False, 
                'error': 'Program code and name are required'
            }), 400
        
        # Check if program code already exists
        existing = Program.query.filter_by(
            program_code=program_code, 
            is_deleted=False
        ).first()                                                                                                     
        
        if existing:
            return jsonify({
                'success': False, 
                'error': f'Program code "{program_code}" already exists'
            }), 400
        
        program = Program(
            program_code=program_code,
            program_name=program_name,
            name=program_name,
           duration_years=int(data.get('duration_years', 3)),
            duration=int(data.get('duration_years', 3))
        )
        db.session.add(program)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'program': {
                'program_id': program.program_id,
                'program_code': program.program_code,
                'program_name': program.program_name,
                'duration_years': program.duration_years
            }
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/programs/<program_id>', methods=['GET'])
@admin_required
def api_get_program(program_id):
    """Get program details"""
    try:
        program = Program.query.get_or_404(program_id)
        
        if program.is_deleted:
            return jsonify({'success': False, 'error': 'Program not found'}), 404
        
        return jsonify({
            'success': True,
            'program': {
                'program_id': program.program_id,
                'program_code': program.program_code,
                'program_name': program.program_name,
                'duration_years': program.duration_years
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/programs/<program_id>', methods=['PUT'])
@admin_required
def api_update_program(program_id):
    """Update existing program"""
    try:
        program = Program.query.get_or_404(program_id)
        
        if program.is_deleted:
            return jsonify({'success': False, 'error': 'Program not found'}), 404
        
        data = request.get_json()
        
        # Update fields
        if 'program_name' in data:
            program.program_name = data['program_name']
            program.name = data['program_name']
        
        if 'duration_years' in data:
            duration = int(data['duration_years'])
            program.duration_years = duration
            program.duration = duration
        
        if 'current_academic_year' in data:
            new_academic_year = data['current_academic_year']
            program.current_academic_year = new_academic_year
            
            # Sync to all students in this program
            students = Student.query.filter_by(program_id=program_id, is_deleted=False).all()
            for student in students:
                student.current_academic_year = new_academic_year
                
                # Trigger fee re-assignment for the new academic year
                if student.seat_type and student.joining_academic_year:
                    try:
                        from fee_helpers import assign_fee_to_student
                        assign_fee_to_student(student.student_id)
                    except Exception as e:
                        print(f"Fee assignment failed for {student.student_id}: {str(e)}")
            
            # Sync to all faculty in this program
            faculties = Faculty.query.filter_by(program_id=program_id, is_deleted=False).all()
            for faculty in faculties:
                faculty.current_academic_year = new_academic_year
        
        # Note: program_code is typically not editable after creation
        # to maintain referential integrity
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'program': {
                'program_id': program.program_id,
                'program_code': program.program_code,
                'program_name': program.program_name,
                'duration_years': program.duration_years,
                'current_academic_year': program.current_academic_year
            },
            'sync_stats': {
                'students': len(students) if 'current_academic_year' in data else 0,
                'faculty': len(faculties) if 'current_academic_year' in data else 0
            }
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/programs/<program_id>', methods=['DELETE'])
@admin_required
def api_delete_program(program_id):
    """Delete program (soft delete)"""
    try:
        program = Program.query.get_or_404(program_id)
        
        # Check if there are active sections using this program
        active_sections = Section.query.filter_by(
            program_id=program_id, 
            is_deleted=False
        ).count()
        
        if active_sections > 0:
            return jsonify({
                'success': False, 
                'error': f'Cannot delete program. {active_sections} active section(s) are using this program.'
            }), 400
        
        program.is_deleted = True
        db.session.commit()
        
        return jsonify({'success': True})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400



# ============================================
# Student Assignment Routes (Admin Only)
# ============================================

@app.route('/admin/assign-students')
@admin_required
def admin_assign_students():
    """Student-section assignment and specialization enrollment page"""
    # Get optional program filter from query params
    program_id = request.args.get('program_id')
    
    # Get all programs for the selector
    programs = Program.query.filter_by(is_deleted=False).order_by(Program.program_name).all()
    
    # Filter students and sections by program if specified
    if program_id:
        students = Student.query.filter_by(
            program_id=program_id,
            is_deleted=False
        ).all()
        sections = Section.query.filter_by(
            program_id=program_id,
            is_deleted=False
        ).all()
        selected_program = Program.query.get(program_id)
        
        # Get unique semesters from students
        available_semesters = sorted(set(
            s.current_semester for s in students 
            if s.current_semester is not None
        ))
    else:
        # If no program selected, show empty state
        students = []
        sections = []
        selected_program = None
        available_semesters = []
    
    subjects = Subject.query.filter_by(is_deleted=False).all()
    
    return render_template('admin_assign_students.html', 
                         students=students, 
                         sections=sections,
                         subjects=subjects,
                         programs=programs,
                         selected_program=selected_program,
                         available_semesters=available_semesters)


@app.route('/api/admin/students/assign-section', methods=['POST'])
@admin_required
def api_assign_students_to_section():
    """Bulk assign students to a section"""
    try:
        data = request.get_json()
        student_ids = data.get('student_ids', [])
        section_id = data.get('section_id')
        
        if not student_ids or not section_id:
            return jsonify({'success': False, 'error': 'Missing student_ids or section_id'}), 400
        
        # Update all students
        updated_count = 0
        for student_id in student_ids:
            student = Student.query.get(student_id)
            if student:
                student.section_id = section_id
                updated_count += 1
        
        db.session.commit()
        return jsonify({'success': True, 'updated': updated_count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/students/update-sections', methods=['POST'])
@admin_required
def api_update_student_sections():
    """Bulk update student sections with attendance transfer logic"""
    try:
        data = request.get_json()
        assignments = data.get('assignments', [])
        transfer_attendance = data.get('transfer_attendance', True)
        
        updated_count = 0
        
        for item in assignments:
            student_id = item.get('student_id')
            new_section_id = item.get('section_id')
            
            student = Student.query.get(student_id)
            if not student:
                continue

            # Check if status or section is changing
            current_section = student.section_id
            current_active = (student.status == 'active')
            
            # Helper to determine new state
            target_section = new_section_id
            target_active = True
            target_status = 'active'
            
            if new_section_id == 'left_college':
                target_section = None
                target_active = False
                target_status = 'inactive'
            elif new_section_id == 'none' or new_section_id is None:
                target_section = None
                target_active = True  # Unassigned but still active
                target_status = 'active'
            
            # Detect change
            is_section_changed = (current_section != target_section)
            is_status_changed = (current_active != target_active)
            
            if is_section_changed or is_status_changed:
                # If section changed (and active) AND NOT transferring, reset attendance
                # (We don't reset if just marking inactive, records should stay for history)
                if is_section_changed and target_active and not transfer_attendance:
                     AttendanceRecord.query.filter_by(student_id=student_id, is_deleted=False).update({'is_deleted': True})
                
                student.section_id = target_section
                student.status = target_status
                updated_count += 1
        
        db.session.commit()
        return jsonify({'success': True, 'updated': updated_count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/students/enroll-subject', methods=['POST'])
@admin_required
def api_enroll_students_in_subject():
    """Enroll students in a specialization subject"""
    try:
        data = request.get_json()
        student_ids = data.get('student_ids', [])
        subject_id = data.get('subject_id')
        academic_year = data.get('academic_year')
        semester = data.get('semester')
        
        if not student_ids or not subject_id:
            return jsonify({'success': False, 'error': 'Missing student_ids or subject_id'}), 400
        
        enrolled_count = 0
        errors = []
        
        for student_id in student_ids:
            # Check if already enrolled
            existing = StudentSubjectEnrollment.query.filter_by(
                student_id=student_id,
                subject_id=subject_id,
                academic_year=academic_year
            ).first()
            
            if not existing:
                enrollment = StudentSubjectEnrollment(
                    student_id=student_id,
                    subject_id=subject_id,
                    academic_year=academic_year,
                    semester=semester
                )
                db.session.add(enrollment)
                enrolled_count += 1
            else:
                errors.append(f"Student {student_id} already enrolled")
        
        db.session.commit()
        return jsonify({
            'success': True, 
            'enrolled': enrolled_count,
            'errors': errors
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/students/unenroll-subject', methods=['POST'])
@admin_required
def api_unenroll_student_from_subject():
    """Remove student enrollment from specialization subject"""
    try:
        data = request.get_json()
        enrollment_id = data.get('enrollment_id')
        
        enrollment = StudentSubjectEnrollment.query.get_or_404(enrollment_id)
        enrollment.is_deleted = True
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# ============================================
# Bulk Import Routes (Admin Only)
# ============================================

@app.route('/admin/import')
@admin_required
def admin_import():
    """
    Display bulk import interface for admin
    """
    try:
        recent_imports = ImportLog.query.order_by(ImportLog.created_at.desc()).limit(10).all()
    except Exception:
        # Handle case where old enum values exist in database
        recent_imports = []
    return render_template('admin_import.html', imports=recent_imports)


@app.route('/api/admin/import/template/<import_type>')
@admin_required
def download_import_template(import_type):
    """
    Download sample Excel template for import
    """
    import os
    from flask import send_file
    
    template_files = {
        'student': 'students_template.csv',
        'faculty': 'faculty_template.csv',
        'subject': 'subjects_template.csv',
        'schedule': 'schedules_template.csv'
    }
    
    if import_type not in template_files:
        return jsonify({'success': False, 'error': 'Invalid import type'}), 400
    
    template_path = os.path.join(app.root_path, 'sample_imports', template_files[import_type])
    
    if not os.path.exists(template_path):
        return jsonify({'success': False, 'error': 'Template not found'}), 404
    
    # Determine mimetype based on file extension
    if template_files[import_type].endswith('.xlsx'):
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    else:
        mimetype = 'text/csv'
    
    return send_file(
        template_path,
        as_attachment=True,
        download_name=template_files[import_type],
        mimetype=mimetype
    )


@app.route('/api/admin/import', methods=['POST'])
@admin_required
def api_bulk_import():
    """
    Handle bulk data import from CSV/Excel files
    Supports: students, faculty, subjects, schedules
    """
    print("[IMPORT] Starting import request...")  # Debug
    current_user = get_current_user()
    
    if 'file' not in request.files:
        print("[IMPORT] No file in request")  # Debug
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    import_type = request.form.get('import_type')
    
    print(f"[IMPORT] File: {file.filename}, Type: {import_type}")  # Debug
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not import_type:
        return jsonify({'success': False, 'error': 'Import type not specified'}), 400
    
    # Allowed file extensions
    allowed_extensions = {'.csv', '.xlsx', '.xls'}
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        return jsonify({'success': False, 'error': 'Invalid file type. Use CSV or Excel'}), 400
    
    try:
        # Read file data
        import pandas as pd
        print("[IMPORT] Reading file...")  # Debug
        
        if file_ext == '.csv':
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)
        
        print(f"[IMPORT] File read. Rows: {len(df)}, Columns: {list(df.columns)}")  # Debug
        
        # Process import based on type
        print("[IMPORT] Starting process_bulk_import...")  # Debug
        result = process_bulk_import(df, import_type, current_user, file.filename)
        print(f"[IMPORT] Completed: {result}")  # Debug
        
        return jsonify(result)
    
    except Exception as e:
        print(f"[IMPORT] ERROR: {str(e)}")  # Debug
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Import failed: {str(e)}'}), 400


# ============================================
# Database Initialization
# ============================================

def init_db():
    """
    Initialize the database - create all tables
    """
    with app.app_context():
        db.create_all()
        print("âœ“ Database tables created successfully!")
        
        # Create default roles if they don't exist
        create_default_roles()
        print("âœ“ Default roles initialized!")


def create_default_roles():
    """
    Create default system roles
    Students can expand this to add more roles
    """
    default_roles = [
        {'role_name': 'admin', 'description': 'System Administrator'},
        {'role_name': 'hod', 'description': 'Head of Department'},
        {'role_name': 'faculty', 'description': 'Faculty Member'},
        {'role_name': 'student', 'description': 'Student'},
        {'role_name': 'parent', 'description': 'Parent/Guardian'},
    ]
    
    for role_data in default_roles:
        existing_role = Role.query.filter_by(role_name=role_data['role_name']).first()
        if not existing_role:
            role = Role(**role_data)
            db.session.add(role)
    
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Error creating roles: {e}")


# ============================================
# Work Diary Helper Functions
# ============================================

def auto_create_work_diary_from_attendance(attendance_session):
    """
    Automatically create work diary entry when attendance is taken
    """
    schedule = attendance_session.schedule
    if not schedule:
        return None
    
    # Generate diary number
    year = datetime.utcnow().year
    last_diary = WorkDiary.query.filter(
        WorkDiary.diary_number.like(f'WD-{year}-%')
    ).order_by(WorkDiary.diary_number.desc()).first()
    
    if last_diary:
        last_num = int(last_diary.diary_number.split('-')[-1])
        diary_num = f'WD-{year}-{last_num + 1:04d}'
    else:
        diary_num = f'WD-{year}-0001'
    
    # Determine activity type based on schedule
    activity_type = 'practical_class' if 'lab' in schedule.room_number.lower() else 'theory_class'
    
    # Count students present (assuming records exist)
    students_present = AttendanceRecord.query.filter_by(
        attendance_session_id=attendance_session.attendance_session_id,
        status='present'
    ).count()
    
    students_total = AttendanceRecord.query.filter_by(
        attendance_session_id=attendance_session.attendance_session_id
    ).count()
    
    # Create diary entry
    diary = WorkDiary(
        diary_number=diary_num,
        faculty_id=schedule.faculty_id,
        subject_id=schedule.subject_id,
        section_id=schedule.section.section_id if schedule.section else None,
        date=attendance_session.session_date or datetime.utcnow().date(),
        start_time=schedule.start_time,
        end_time=schedule.end_time,
        activity_type=activity_type,
        attendance_session_id=attendance_session.attendance_session_id,
        students_present=students_present,
        students_total=students_total,
        status='draft'
    )
    
    db.session.add(diary)
    return diary


def create_work_diary_entry(data):
    """
    Create work diary entry from form/API data
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        raise ValueError('Faculty record not found')
    
    # Generate diary number
    year = datetime.utcnow().year
    last_diary = WorkDiary.query.filter(
        WorkDiary.diary_number.like(f'WD-{year}-%')
    ).order_by(WorkDiary.diary_number.desc()).first()
    
    if last_diary:
        last_num = int(last_diary.diary_number.split('-')[-1])
        diary_num = f'WD-{year}-{last_num + 1:04d}'
    else:
        diary_num = f'WD-{year}-0001'
    
    # Parse time strings
    start_time = datetime.strptime(data['start_time'], '%H:%M').time() if isinstance(data.get('start_time'), str) else data.get('start_time')
    end_time = datetime.strptime(data['end_time'], '%H:%M').time() if isinstance(data.get('end_time'), str) else data.get('end_time')
    date = datetime.strptime(data['date'], '%Y-%m-%d').date() if isinstance(data.get('date'), str) else data.get('date')
    
    diary = WorkDiary(
        diary_number=diary_num,
        faculty_id=faculty_record.faculty_id,
        subject_id=data.get('subject_id'),
        section_id=data.get('section_id'),
        date=date,
        start_time=start_time,
        end_time=end_time,
        activity_type=data['activity_type'],
        students_present=data.get('students_present', 0),
        students_total=data.get('students_total', 0),
        activity_title=data.get('activity_title'),
        activity_description=data.get('activity_description'),
        topics_covered=data.get('topics_covered'),
        status='draft'
    )
    
    db.session.add(diary)
    db.session.commit()
    return diary


def update_work_diary_entry(diary, data):
    """
    Update existing work diary entry
    """
    # Only allow updates to draft/rejected entries
    if diary.status not in ['draft', 'rejected']:
        raise ValueError('Cannot edit approved or submitted entries')
    
    # Update fields
    if 'date' in data:
        diary.date = datetime.strptime(data['date'], '%Y-%m-%d').date() if isinstance(data['date'], str) else data['date']
    if 'start_time' in data:
        diary.start_time = datetime.strptime(data['start_time'], '%H:%M').time() if isinstance(data['start_time'], str) else data['start_time']
    if 'end_time' in data:
        diary.end_time = datetime.strptime(data['end_time'], '%H:%M').time() if isinstance(data['end_time'], str) else data['end_time']
    if 'activity_type' in data:
        diary.activity_type = data['activity_type']
    if 'subject_id' in data:
        diary.subject_id = data['subject_id']
    if 'section_id' in data:
        diary.section_id = data['section_id']
    if 'students_present' in data:
        diary.students_present = data['students_present']
    if 'students_total' in data:
        diary.students_total = data['students_total']
    if 'activity_title' in data:
        diary.activity_title = data['activity_title']
    if 'activity_description' in data:
        diary.activity_description = data['activity_description']
    if 'topics_covered' in data:
        diary.topics_covered = data['topics_covered']
    
    db.session.commit()
    return diary


def handle_work_diary_creation():
    """
    Handle POST request for creating work diary
    """
    try:
        data = {
            'date': request.form.get('date'),
            'start_time': request.form.get('start_time'),
            'end_time': request.form.get('end_time'),
            'activity_type': request.form.get('activity_type'),
            'subject_id': request.form.get('subject_id'),
            'section_id': request.form.get('section_id'),
            'students_present': int(request.form.get('students_present', 0)),
            'students_total': int(request.form.get('students_total', 0)),
            'activity_title': request.form.get('activity_title'),
            'activity_description': request.form.get('activity_description'),
            'topics_covered': request.form.get('topics_covered')
        }
        
        diary = create_work_diary_entry(data)
        return redirect(url_for('work_diary_list'))
    except Exception as e:
        return f"Error: {str(e)}", 400


def handle_work_diary_update(diary):
    """
    Handle POST request for updating work diary
    """
    try:
        data = {
            'date': request.form.get('date'),
            'start_time': request.form.get('start_time'),
            'end_time': request.form.get('end_time'),
            'activity_type': request.form.get('activity_type'),
            'subject_id': request.form.get('subject_id'),
            'section_id': request.form.get('section_id'),
            'students_present': int(request.form.get('students_present', 0)),
            'students_total': int(request.form.get('students_total', 0)),
            'activity_title': request.form.get('activity_title'),
            'activity_description': request.form.get('activity_description'),
            'topics_covered': request.form.get('topics_covered')
        }
        
        update_work_diary_entry(diary, data)
        return redirect(url_for('work_diary_list'))
    except Exception as e:
        return f"Error: {str(e)}", 400


def process_bulk_import(df, import_type, current_user, filename):
    """
    Process bulk import data from DataFrame
    Returns result dict with success/failure counts
    """
    success_count = 0
    error_count = 0
    errors = []
    
    try:
        # Process import based on type
        if import_type == 'student':
            success_count, error_count, errors = import_students(df)
        elif import_type == 'faculty':
            success_count, error_count, errors = import_faculty(df)
        elif import_type == 'subject':
            success_count, error_count, errors = import_subjects(df)
        elif import_type == 'schedule':
            success_count, error_count, errors = import_schedules(df)
        else:
            raise ValueError(f'Unknown import type: {import_type}')
        
        # Create import log AFTER successful import
        import_log = ImportLog(
            import_type=import_type,
            imported_by=current_user.user_id,
            file_name=filename,
            total_rows=len(df),
            successful_rows=success_count,
            failed_rows=error_count,
            status='completed' if error_count == 0 else 'completed_with_errors',
            error_log='\n'.join(errors[:50]) if errors else None
        )
        db.session.add(import_log)
        db.session.commit()
        
        return {
            'success': True,
            'message': f'Import completed: {success_count} succeeded, {error_count} failed',
            'successful': success_count,
            'failed': error_count,
            'errors': errors[:10]  # Return first 10 errors
        }
    
    except Exception as e:
        db.session.rollback()
        # Create failed import log
        try:
            import_log = ImportLog(
                import_type=import_type,
                imported_by=current_user.user_id,
                file_name=filename,
                total_rows=len(df),
                status='failed',
                error_log=str(e)
            )
            db.session.add(import_log)
            db.session.commit()
        except:
            pass
        raise


def import_students(df):
    """
    Import students from DataFrame
    Creates User account for each student:
    - Username = USN (roll_number)
    - Password = date_of_birth (DDMMYYYY format)
    
    Supports both old format (first_name, last_name, roll_number) and 
    new format (name, usn)
    """
    import pandas as pd
    from auth import hash_password
    
    success_count = 0
    error_count = 0
    errors = []
    
    # Normalize column names - support both USN and roll_number
    if 'usn' in df.columns and 'roll_number' not in df.columns:
        df['roll_number'] = df['usn']
    
    # Support single 'name' field - split into first_name and last_name
    if 'name' in df.columns:
        # If 'name' column exists, use it directly for the 'name' field in Student model
        # No need to split into first_name/last_name if the model only uses 'name'
        pass
    elif 'first_name' in df.columns or 'last_name' in df.columns:
        # If separate first_name/last_name, combine them into a 'name' column for processing
        df['name'] = df.apply(lambda row: f"{row.get('first_name', '')} {row.get('last_name', '')}".strip(), axis=1)
    else:
        # If neither 'name' nor 'first_name'/'last_name' are present, raise error or handle default
        pass # Will be caught by required_columns check below if 'name' is required
    
    # Required columns (minimum) - now requires 'name'
    required_columns = ['roll_number', 'name', 'email', 'date_of_birth']
    
    # Validate columns
    missing_cols = set(required_columns) - set(df.columns)
    if missing_cols:
        raise ValueError(f'Missing required columns: {", ".join(missing_cols)}. Ensure "name" column is present or "first_name" and "last_name" to be combined.')
    
    # Get student role
    student_role = Role.query.filter_by(role_name='student').first()
    if not student_role:
        raise ValueError('Student role not found in database')
    
    
    total_rows = len(df)
    print(f"[IMPORT] Starting to process {total_rows} students...")
    
    for idx, row in df.iterrows():
        # Progress logging every 50 rows
        if (idx + 1) % 50 == 0:
            print(f"[IMPORT] Progress: {idx + 1}/{total_rows} rows processed...")
        
        try:
            roll_number = str(row['roll_number']).strip()
            
            # Check if student already exists
            existing = Student.query.filter_by(roll_number=roll_number).first()
            if existing:
                errors.append(f"Row {idx+2}: Student {roll_number} already exists")
                error_count += 1
                continue
            
            # Check if username already exists
            existing_user = User.query.filter_by(username=roll_number).first()
            if existing_user:
                errors.append(f"Row {idx+2}: Username {roll_number} already exists")
                error_count += 1
                continue
            
            # Parse date of birth for password
            dob = row['date_of_birth']
            if pd.isna(dob):
                errors.append(f"Row {idx+2}: Date of birth is required for password generation")
                error_count += 1
                continue
            
            # Convert to date if string
            if isinstance(dob, str):
                try:
                    dob = pd.to_datetime(dob).date()
                except:
                    errors.append(f"Row {idx+2}: Invalid date format for date_of_birth")
                    error_count += 1
                    continue
            elif hasattr(dob, 'date'):
                dob = dob.date()
            
            # Password = DDMMYYYY format
            password = dob.strftime('%d%m%Y')
            
            # Create User account
            user = User(
                username=roll_number,
                password_hash=hash_password(password),
                email=str(row['email']).strip(),
                role_id=student_role.role_id,
                is_active=True
            )
            db.session.add(user)
            db.session.flush()  # Get user_id
            
            # Parse optional fields
            phone = row.get('phone', '')
            if pd.isna(phone):
                phone = None
            else:
                phone = str(phone).strip() if phone else None
            
            admission_year = row.get('admission_year')
            if pd.isna(admission_year) if hasattr(pd, 'isna') else admission_year is None:
                admission_year = None
            else:
                admission_year = int(admission_year)
            
            guardian_name = row.get('guardian_name', '')
            if pd.isna(guardian_name):
                guardian_name = None
            
            guardian_phone = row.get('guardian_phone', '')
            if pd.isna(guardian_phone):
                guardian_phone = None
            
            address = row.get('address', '')
            if pd.isna(address):
                address = None
            
            # Get program_id if provided
            program_id = None
            if 'program_code' in row and not pd.isna(row.get('program_code')):
                program = Program.query.filter_by(program_code=str(row['program_code']).strip()).first()
                if program:
                    program_id = program.program_id
            
            # Get the full name from the 'name' column
            full_name = str(row.get('name', '')).strip() if pd.notna(row.get('name')) else ''
            
            # Fallback: if no name, use roll_number
            if not full_name:
                full_name = roll_number
            
            # Create Student record
            student = Student(
                user_id=user.user_id,
                roll_number=roll_number,
                usn=roll_number,
                name=full_name, # Use the combined/provided name
                email=str(row['email']).strip(),
                phone=phone,
                date_of_birth=dob,
                guardian_name=guardian_name,
                guardian_phone=str(guardian_phone).strip() if guardian_phone else None,
                address=address,
                admission_year=admission_year,
                program_id=program_id,
                status='active',
                gender=str(row.get('gender', '')).strip().upper()[:1] if pd.notna(row.get('gender')) else None  # M, F, or O
            )
            db.session.add(student)
            
            # NOTE: Parent accounts are no longer created separately
            # Parents now login using the same student credentials with "Parent" radio option
            
            success_count += 1
            
        except Exception as e:
            print(f"[IMPORT ERROR] Row {idx+2}: {str(e)}")  # Debug logging
            import traceback
            traceback.print_exc()  # Print full stack trace
            errors.append(f"Row {idx+2}: {str(e)}")
            error_count += 1
    
    # Log summary
    print(f"[IMPORT COMPLETE] Success: {success_count}, Failed: {error_count}")
    if errors:
        print(f"[IMPORT ERRORS] {errors}")
    
    db.session.commit()
    return success_count, error_count, errors


def import_faculty(df):
    """Import faculty from DataFrame"""
    import pandas as pd
    from auth import hash_password
    
    success_count = 0
    error_count = 0
    errors = []
    
    required_columns = ['employee_id', 'first_name', 'last_name', 'email']
    
    missing_cols = set(required_columns) - set(df.columns)
    if missing_cols:
        raise ValueError(f'Missing required columns: {", ".join(missing_cols)}')
    
    # Get faculty role
    faculty_role = Role.query.filter_by(role_name='faculty').first()
    if not faculty_role:
        raise ValueError('Faculty role not found in database')
    
    for idx, row in df.iterrows():
        try:
            employee_id = str(row['employee_id']).strip()
            
            # Check if faculty already exists
            existing = Faculty.query.filter_by(employee_id=employee_id).first()
            if existing:
                errors.append(f"Row {idx+2}: Faculty {employee_id} already exists")
                error_count += 1
                continue
            
            # Create User account first
            # Username = employee_id, Password = employee_id (default)
            user = User.query.filter_by(username=employee_id).first()
            
            # Check if email is already in use by a DIFFERENT user
            email = str(row['email']).strip()
            existing_email_user = User.query.filter_by(email=email).first()
            
            if existing_email_user:
                if user and existing_email_user.user_id != user.user_id:
                     pass 
                elif not user:
                     errors.append(f"Row {idx+2}: Email {email} is already in use by user '{existing_email_user.username}'")
                     error_count += 1
                     continue
            
            if not user:
                user = User(
                    username=employee_id,
                    email=email,
                    password_hash=hash_password(employee_id), # Default password is employee_id
                    role_id=faculty_role.role_id,
                    is_active=True
                )
                db.session.add(user)
                db.session.flush() # Flush to get user_id
            
            faculty = Faculty(
                employee_id=employee_id,
                user_id=user.user_id,
                first_name=row['first_name'],
                last_name=row['last_name'],
                email=row['email'],
                phone=row.get('phone') if pd.notna(row.get('phone')) else None,
                designation=row.get('designation') if pd.notna(row.get('designation')) else None,
                department=row.get('department') if pd.notna(row.get('department')) else None
            )
            db.session.add(faculty)
            db.session.commit() # Commit each successful row immediately
            success_count += 1
            
        except Exception as e:
            db.session.rollback() # Rollback transaction on any error to clean session
            errors.append(f"Row {idx+2}: {str(e)}")
            error_count += 1
    
    # db.session.commit() # Removed final commit as we commit per row
    return success_count, error_count, errors


def import_subjects(df):
    """Import subjects from DataFrame"""
    import pandas as pd
    
    success_count = 0
    error_count = 0
    errors = []
    
    required_columns = ['subject_code', 'subject_name', 'semester']
    
    missing_cols = set(required_columns) - set(df.columns)
    if missing_cols:
        raise ValueError(f'Missing required columns: {", ".join(missing_cols)}')
    
    for idx, row in df.iterrows():
        try:
            existing = Subject.query.filter_by(subject_code=str(row['subject_code']).strip()).first()
            if existing:
                errors.append(f"Row {idx+2}: Subject {row['subject_code']} already exists")
                error_count += 1
                continue
            
            # Get credits, handle NaN
            credits_val = row.get('credits', 4) if not pd.isna(row.get('credits', 4)) else 4
            subject_type_val = row.get('subject_type', 'theory') if not pd.isna(row.get('subject_type', 'theory')) else 'theory'
            
            subject = Subject(
                subject_code=str(row['subject_code']).strip(),
                subject_name=str(row['subject_name']).strip(),
                semester_id=int(row['semester']),  # Use semester_id field
                credits=float(credits_val),
                subject_type=str(subject_type_val).strip()
            )
            db.session.add(subject)
            success_count += 1
            
        except Exception as e:
            print(f"[IMPORT ERROR] Row {idx+2}: {str(e)}")
            import traceback
            traceback.print_exc()
            errors.append(f"Row {idx+2}: {str(e)}")
            error_count += 1
    
    print(f"[IMPORT COMPLETE] Subjects - Success: {success_count}, Failed: {error_count}")
    if errors:
        print(f"[IMPORT ERRORS] {errors}")
    
    db.session.commit()
    return success_count, error_count, errors


def import_schedules(df):
    """Import class schedules from DataFrame"""
    success_count = 0
    error_count = 0
    errors = []
    
    required_columns = ['subject_code', 'faculty_employee_id', 'day_of_week', 'start_time', 'end_time', 'room_number']
    
    missing_cols = set(required_columns) - set(df.columns)
    if missing_cols:
        raise ValueError(f'Missing required columns: {", ".join(missing_cols)}')
    
    for idx, row in df.iterrows():
        try:
            # Lookup subject and faculty
            subject = Subject.query.filter_by(subject_code=row['subject_code']).first()
            faculty = Faculty.query.filter_by(employee_id=row['faculty_employee_id']).first()
            
            if not subject:
                errors.append(f"Row {idx+2}: Subject {row['subject_code']} not found")
                error_count += 1
                continue
            
            if not faculty:
                errors.append(f"Row {idx+2}: Faculty {row['faculty_employee_id']} not found")
                error_count += 1
                continue
            
            schedule = ClassSchedule(
                subject_id=subject.subject_id,
                faculty_id=faculty.faculty_id,
                day_of_week=row['day_of_week'],
                start_time=row['start_time'],
                end_time=row['end_time'],
                room_number=row['room_number']
            )
            db.session.add(schedule)
            success_count += 1
            
        except Exception as e:
            errors.append(f"Row {idx+2}: {str(e)}")
            error_count += 1
    
    db.session.commit()
    return success_count, error_count, errors


# ============================================
# Campus Check-In System - Geo-fenced Attendance
# ============================================

def haversine_distance(lat1, lon1, lat2, lon2):
    """
    Calculate the great-circle distance between two points on Earth using Haversine formula.
    Returns distance in meters.
    """
    R = 6371000  # Earth's radius in meters
    
    # Convert to radians
    phi1 = math.radians(lat1)
    phi2 = math.radians(lat2)
    delta_phi = math.radians(lat2 - lat1)
    delta_lambda = math.radians(lon2 - lon1)
    
    # Haversine formula
    a = math.sin(delta_phi / 2) ** 2 + \
        math.cos(phi1) * math.cos(phi2) * math.sin(delta_lambda / 2) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    
    return R * c


def get_campus_config():
    """
    Get campus configuration including coordinates and check-in settings.
    Creates default config if not exists.
    """
    config = CollegeConfig.query.filter_by(config_key='campus_location').first()
    if not config:
        # Default config - Bangalore University approximate coordinates
        config = CollegeConfig(
            config_key='campus_location',
            campus_latitude=12.9393,  # Default: Bangalore University
            campus_longitude=77.5829,
            campus_radius_meters=200,  # 200 meter radius
            college_name='Bangalore University - BCA Department',
            checkin_start_time=time(7, 0),   # 7:00 AM
            checkin_end_time=time(18, 0)     # 6:00 PM
        )
        db.session.add(config)
        db.session.commit()
    return config


@app.route('/api/campus-checkin', methods=['POST'])
@login_required
def api_campus_checkin():
    """
    Student campus check-in endpoint.
    Validates location is within campus radius and records check-in.
    """
    current_user = get_current_user()
    
    # Verify user is a student
    if not current_user.role or current_user.role.role_name.lower() != 'student':
        return jsonify({'success': False, 'error': 'Only students can check in'}), 403
    
    # Get student record
    student = Student.query.filter_by(user_id=current_user.user_id, is_deleted=False).first()
    if not student:
        return jsonify({'success': False, 'error': 'Student record not found'}), 404
    
    data = request.get_json()
    latitude = data.get('latitude')
    longitude = data.get('longitude')
    device_info = data.get('device_info', '')
    
    if latitude is None or longitude is None:
        return jsonify({'success': False, 'error': 'Location coordinates required'}), 400
    
    try:
        latitude = float(latitude)
        longitude = float(longitude)
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'Invalid coordinates format'}), 400
    
    # Get campus config
    config = get_campus_config()
    
    # Check if already checked in today
    today = date.today()
    existing_checkin = CampusCheckIn.query.filter_by(
        student_id=student.student_id,
        checkin_date=today,
        is_deleted=False
    ).first()
    
    if existing_checkin:
        return jsonify({
            'success': False, 
            'error': 'Already checked in today',
            'checkin_time': existing_checkin.checkin_time.isoformat()
        }), 400
    
    # Check time restrictions (optional)
    current_time = datetime.now().time()
    if config.checkin_start_time and config.checkin_end_time:
        if current_time < config.checkin_start_time or current_time > config.checkin_end_time:
            return jsonify({
                'success': False, 
                'error': f'Check-in only allowed between {config.checkin_start_time.strftime("%H:%M")} and {config.checkin_end_time.strftime("%H:%M")}'
            }), 400
    
    # Calculate distance from campus
    distance = haversine_distance(
        latitude, longitude,
        config.campus_latitude, config.campus_longitude
    )
    
    is_valid_location = distance <= config.campus_radius_meters
    
    if not is_valid_location:
        return jsonify({
            'success': False, 
            'error': f'You are {int(distance)} meters from campus. Please be within {config.campus_radius_meters} meters to check in.',
            'distance': int(distance),
            'required_radius': config.campus_radius_meters
        }), 400
    
    # Create check-in record
    try:
        checkin = CampusCheckIn(
            student_id=student.student_id,
            checkin_date=today,
            checkin_time=current_time,
            latitude=latitude,
            longitude=longitude,
            is_valid_location=is_valid_location,
            device_info=device_info[:255] if device_info else None
        )
        db.session.add(checkin)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Checked in successfully!',
            'checkin': checkin.to_dict()
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/campus-checkin/status')
@login_required
def api_checkin_status():
    """
    Get today's check-in status for the current student.
    """
    current_user = get_current_user()
    
    if not current_user.role or current_user.role.role_name.lower() != 'student':
        return jsonify({'success': False, 'error': 'Only students can view check-in status'}), 403
    
    student = Student.query.filter_by(user_id=current_user.user_id, is_deleted=False).first()
    if not student:
        return jsonify({'success': False, 'error': 'Student record not found'}), 404
    
    today = date.today()
    checkin = CampusCheckIn.query.filter_by(
        student_id=student.student_id,
        checkin_date=today,
        is_deleted=False
    ).first()
    
    config = get_campus_config()
    
    return jsonify({
        'success': True,
        'checked_in': checkin is not None,
        'checkin': checkin.to_dict() if checkin else None,
        'campus_config': {
            'radius_meters': config.campus_radius_meters,
            'start_time': config.checkin_start_time.isoformat() if config.checkin_start_time else None,
            'end_time': config.checkin_end_time.isoformat() if config.checkin_end_time else None
        }
    })


# ============================================
# Class Teacher Dashboard Routes
# ============================================

@app.route('/faculty')
@faculty_required
def faculty_dashboard_view():
    """
    Main Faculty Dashboard
    Displays:
    - Weekly/Monthly teaching hours
    - Assigned Subjects (Academic & Admin)
    - Quick Actions
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        return redirect(url_for('login'))
        
    # 1. Calculate Stats (Hours)
    today = date.today()
    start_of_week = today - timedelta(days=today.weekday()) # Monday
    start_of_month = today.replace(day=1)
    
    # Query sessions for this faculty
    week_sessions = AttendanceSession.query.filter(
        AttendanceSession.taken_by_user_id == current_user.user_id,
        AttendanceSession.taken_at >= datetime.combine(start_of_week, time.min),
        AttendanceSession.is_deleted == False
    ).all()
    
    month_sessions = AttendanceSession.query.filter(
        AttendanceSession.taken_by_user_id == current_user.user_id,
        AttendanceSession.taken_at >= datetime.combine(start_of_month, time.min),
        AttendanceSession.is_deleted == False
    ).all()
    
    def calculate_session_duration(session):
        try:
            if session.schedule and session.schedule.start_time and session.schedule.end_time:
                # Calculate duration from time objects
                dummy_date = date.today()
                start = datetime.combine(dummy_date, session.schedule.start_time)
                end = datetime.combine(dummy_date, session.schedule.end_time)
                duration = (end - start).total_seconds() / 3600
                return round(duration, 2)
        except:
            pass
        return 1.0  # Default to 1 hour if calculation fails
    
    hours_week = sum([calculate_session_duration(s) for s in week_sessions])
    hours_month = sum([calculate_session_duration(s) for s in month_sessions])
    
    # 2. Fetch Assigned Work
    allocations = SubjectAllocation.query.filter_by(faculty_id=faculty_record.faculty_id).all()
    
    assigned_work = []
    for alloc in allocations:
        subject = alloc.subject
        section = alloc.section
        
        # Check if Admin Work (custom type 'admin')
        is_admin_work = subject.subject_type and subject.subject_type.lower() == 'admin'
        
        # Get progress (Classes conducted for this subject/section)
        classes_taken = AttendanceSession.query.join(ClassSchedule).filter(
            ClassSchedule.subject_id == subject.subject_id,
            ClassSchedule.section_id == section.section_id if section else True,
            AttendanceSession.taken_by_user_id == current_user.user_id,
            AttendanceSession.is_deleted == False
        ).count() if not is_admin_work else 0
        
        assigned_work.append({
            'subject_name': subject.subject_name,
            'section': section.section_name if section else 'N/A',
            'type': subject.subject_type,
            'is_admin': is_admin_work,
            'description': subject.description,
            'classes_taken': classes_taken
        })
    
    # 3. Recent Attendance (Last 5 Sessions)
    recent_sessions = AttendanceSession.query.filter(
        AttendanceSession.taken_by_user_id == current_user.user_id,
        AttendanceSession.is_deleted == False
    ).order_by(AttendanceSession.taken_at.desc()).limit(5).all()
    
    # 4. Check-in/Out Status
    from models import FacultyAttendance
    attendance = FacultyAttendance.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        date=today
    ).first()
    
    current_checkin = attendance.to_dict() if attendance else None
    has_checked_in = bool(attendance and attendance.check_in_time)
    has_checked_out = bool(attendance and attendance.check_out_time)

    # 5. Calculate student count for each recent session
    recent_data = []
    for sess in recent_sessions:
        present_count = AttendanceRecord.query.filter_by(attendance_session_id=sess.attendance_session_id, status='present').count()
        total_count = AttendanceRecord.query.filter_by(attendance_session_id=sess.attendance_session_id).count()
        
        recent_data.append({
            'session_id': sess.attendance_session_id,
            'subject_name': sess.schedule.subject.subject_name if sess.schedule and sess.schedule.subject else "N/A",
            'section_name': sess.schedule.section.section_name if sess.schedule and sess.schedule.section else "N/A",
            'date': sess.taken_at.strftime('%d %b, %H:%M'),
            'present_count': present_count,
            'total_count': total_count
        })

    # 6. Check if Class Teacher
    managed_section = Section.query.filter_by(class_teacher_id=faculty_record.faculty_id).first()
    is_class_teacher = managed_section is not None
    
    # Check if today is a holiday
    holiday = Holiday.query.filter_by(holiday_date=today).first()
    is_holiday = holiday is not None
    
    return render_template('faculty/dashboard.html',
                         faculty=faculty_record,
                         hours_week=round(hours_week, 1),
                         hours_month=round(hours_month, 1),
                         assigned_work=assigned_work,
                         recent_sessions=recent_data,
                         is_class_teacher=is_class_teacher,
                         current_checkin=current_checkin,
                         has_checked_in=has_checked_in,
                         has_checked_out=has_checked_out,
                         is_holiday=is_holiday,
                         holiday=holiday)


@app.route('/faculty/attendance-report')
@faculty_required
def faculty_attendance_report():
    """
    Generate professional monthly/periodical attendance report for faculty.
    Supports start_date and end_date query params.
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    # Get range from query params
    start_str = request.args.get('start_date')
    end_str = request.args.get('end_date')
    
    today = date.today()
    if not start_str or not end_str:
        # Default to current month
        start_date = today.replace(day=1)
        # Last day of current month
        if today.month == 12:
            end_date = date(today.year, 12, 31)
        else:
            end_date = (today.replace(day=1, month=today.month+1) - timedelta(days=1))
    else:
        try:
            start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
        except ValueError:
            start_date = today.replace(day=1)
            end_date = today
            
    # Fetch attendance records in range
    attendance_records = {
        r.date: r for r in FacultyAttendance.query.filter(
            FacultyAttendance.faculty_id == faculty_record.faculty_id,
            FacultyAttendance.date >= start_date,
            FacultyAttendance.date <= end_date
        ).all()
    }
    
    # Fetch holidays in range
    holidays = {
        h.holiday_date: h for h in Holiday.query.filter(
            Holiday.holiday_date >= start_date,
            Holiday.holiday_date <= end_date
        ).all()
    }
    
    # Build daily report data
    report_data = []
    current_day = start_date
    delta = timedelta(days=1)
    
    while current_day <= end_date:
        status = "Absent"
        remarks = ""
        check_in = None
        check_out = None
        
        # Check if holiday
        if current_day in holidays:
            status = "Holiday"
            remarks = holidays[current_day].holiday_name
        elif current_day.weekday() == 6: # Sunday fallback
            status = "Holiday"
            remarks = "Sunday"
        
        # Check if record exists
        if current_day in attendance_records:
            record = attendance_records[current_day]
            if record.check_in_time:
                status = "Present"
                from timezone_utils import format_ist_time
                check_in = format_ist_time(record.check_in_time, '%H:%M')
            if record.check_out_time:
                check_out = format_ist_time(record.check_out_time, '%H:%M')
                
        # Future dates should be blank/upcoming
        if current_day > today and status == "Absent":
            status = "-"
            
        report_data.append({
            'date': current_day.strftime('%d-%m-%Y'),
            'day': current_day.strftime('%A'),
            'status': status,
            'check_in': check_in or '-',
            'check_out': check_out or '-',
            'remarks': remarks
        })
        current_day += delta
        
    return render_template('faculty/attendance_report.html',
                         faculty=faculty_record,
                         start_date=start_date.strftime('%d %b %Y'),
                         end_date=end_date.strftime('%d %b %Y'),
                         report_data=report_data,
                         generated_at=datetime.now().strftime('%d-%m-%Y %H:%M'))


@app.route('/faculty/session/<session_id>')
@faculty_required
def faculty_attendance_session_detail(session_id):
    """
    View details of a specific attendance session
    """
    session = AttendanceSession.query.filter_by(attendance_session_id=session_id).first_or_404()
    
    records = AttendanceRecord.query.filter_by(attendance_session_id=session.attendance_session_id).join(Student).order_by(Student.roll_number).all()
    
    return render_template('faculty/session_detail.html', session=session, records=records)


@app.route('/faculty/session/<session_id>/delete', methods=['POST'])
@faculty_required
def faculty_delete_attendance_session(session_id):
    """
    Secure delete of attendance session
    Requires confirmation_date in JSON body
    """
    session = AttendanceSession.query.filter_by(attendance_session_id=session_id).first_or_404()
    
    current_user = get_current_user()
    if session.taken_by_user_id != current_user.user_id:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403

    data = request.get_json()
    entered_date = data.get('confirmation_date')
    expected_date = session.taken_at.strftime('%d-%m-%Y')
    
    if not entered_date or entered_date.strip() != expected_date:
        return jsonify({'success': False, 'error': 'Incorrect date confirmation'}), 400
        
    try:
        # Soft Delete
        session.is_deleted = True
        
        # Soft delete records
        AttendanceRecord.query.filter_by(attendance_session_id=session.attendance_session_id).update({'is_deleted': True})
        
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/faculty/my-class')
@faculty_required
def faculty_my_class_view():
    """
    Class Teacher Dashboard - My Class
    Displays:
    - Overview of the class (Attendance, Absentees)
    - List of Students (with edit capability)
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        return redirect(url_for('login'))
        
    managed_section = Section.query.filter_by(class_teacher_id=faculty_record.faculty_id).first()
    
    # If not a class teacher, redirect or show error (or empty state)
    if not managed_section:
        # Check if they are authorized to view this page effectively
        # For now, let's render the template but it will show "No Class Assigned"
        return render_template('faculty/my_class.html', faculty=faculty_record, class_data=None)

    # Class Data Logic (Moved from Dashboard)
    total_students = managed_section.students.filter_by(is_deleted=False).count()
    
    # Campus Check-ins for Today
    checkins_today = CampusCheckIn.query.filter(
        CampusCheckIn.checkin_date == date.today(),
        CampusCheckIn.student_id.in_([s.student_id for s in managed_section.students])
    ).all()
    present_count = len(checkins_today)
    present_student_ids = {c.student_id for c in checkins_today}
    
    # Identify Absentees and Student List
    students_list = []
    absentees = []
    
    all_students = managed_section.students.filter_by(is_deleted=False).order_by(Student.roll_number).all()
    
    for s in all_students:
        student_data = {
            'student_id': s.student_id,
            'name': s.name,
            'roll_number': s.roll_number,
            'phone': s.phone,
            'guardian_phone': s.guardian_phone,
            'address': s.address,
            'is_present': s.student_id in present_student_ids
        }
        students_list.append(student_data)
        
        if s.student_id not in present_student_ids:
            absentees.append(student_data)
    
    # Sort students: Low attendance first
    # Calculate percentage for each student first
    for s in students_list:
        # Get overall attendance % (Total sessions vs Attended)
        # Note: Ideally this should be a DB aggregation for performance, 
        # but iterating works for class sizes ~60.
        
        # Count all UNIQUE sessions applicable to this student's section
        # Use distinct to avoid counting same session multiple times due to duplicate ClassSchedule entries
        total_sessions = db.session.query(AttendanceSession.attendance_session_id).join(ClassSchedule).filter(
            ClassSchedule.section_id == managed_section.section_id,
            AttendanceSession.is_deleted == False
        ).distinct().count()
        
        # Count UNIQUE attended sessions
        attended = db.session.query(AttendanceSession.attendance_session_id).join(
            AttendanceRecord, AttendanceRecord.attendance_session_id == AttendanceSession.attendance_session_id
        ).join(ClassSchedule).filter(
            AttendanceRecord.student_id == s['student_id'],
            AttendanceRecord.status == 'present',
            ClassSchedule.section_id == managed_section.section_id
        ).distinct().count()
        
        s['attendance_percentage'] = int((attended / total_sessions * 100)) if total_sessions > 0 else 0
        s['total_attended'] = attended
        s['total_sessions'] = total_sessions

    # Sort: Ascending order (Low attendance on top)
    students_list.sort(key=lambda x: x['attendance_percentage'])

    # Syllabus / Days Left Logic (Simplified)
    last_day = date.today() + timedelta(days=60) 
    days_left = (last_day - date.today()).days if last_day else 0

    class_data = {
        'section_name': managed_section.section_name,
        'semester': managed_section.current_semester,
        'present_count': present_count,
        'total_students': total_students,
        'absentees': absentees,
        'students': students_list,
        'days_left': days_left,
        'attendance_percentage': int((present_count / total_students * 100)) if total_students > 0 else 0
    }

    return render_template('faculty/my_class.html', faculty=faculty_record, class_data=class_data, is_class_teacher=True)



@app.route('/faculty/take-attendance')
@faculty_required
def faculty_take_attendance_app():
    """
    Render the Swipe Attendance App Container
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    # Get subjects assigned to this faculty
    allocations = SubjectAllocation.query.filter_by(faculty_id=faculty_record.faculty_id).all()
    
    # Get unique assigned subjects for this faculty
    unique_allocations = []
    seen_subjects = set()
    for a in allocations:
        if a.subject_id not in seen_subjects:
            unique_allocations.append({
                'subject_id': a.subject_id,
                'subject_name': a.subject.subject_name if a.subject else 'Unknown',
                'subject_type': a.subject.subject_type if a.subject else 'theory',
                'semester_id': a.subject.semester_id if a.subject else None
            })
            seen_subjects.add(a.subject_id)
    
    # Get only sections relevant to this faculty
    # 1. Sections from their class schedules
    # 2. Virtual sections for their assigned subjects
    faculty_section_ids = set()
    
    # Get sections from faculty's schedules
    schedule_sections = db.session.query(ClassSchedule.section_id).filter(
        ClassSchedule.faculty_id == faculty_record.faculty_id,
        ClassSchedule.is_deleted == False
    ).distinct().all()
    faculty_section_ids.update([s.section_id for s in schedule_sections])
    
    # Get virtual sections for faculty's assigned subjects
    faculty_subject_ids = [a['subject_id'] for a in unique_allocations]
    if faculty_subject_ids:
        virtual_sections = Section.query.filter(
            Section.linked_subject_id.in_(faculty_subject_ids),
            Section.is_elective == True,
            Section.is_deleted == False
        ).all()
        faculty_section_ids.update([s.section_id for s in virtual_sections])
    
    # Fetch only relevant sections
    all_sections = Section.query.filter(
        Section.section_id.in_(faculty_section_ids),
        Section.is_deleted == False
    ).order_by(Section.section_name).all() if faculty_section_ids else []
    
    return render_template('faculty/take_attendance_swipe.html', 
                         allocations=unique_allocations,
                         sections=all_sections,
                         allocations_json=json.dumps(unique_allocations),
                         date=date)


@app.route('/faculty/profile', methods=['GET', 'POST'])
@faculty_required
def faculty_profile():
    """
    Faculty Profile - View and Edit
    """
    from werkzeug.security import generate_password_hash
    
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        flash('Faculty record not found', 'error')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        try:
            # Update faculty information
            faculty_record.first_name = request.form.get('first_name')
            faculty_record.last_name = request.form.get('last_name')
            faculty_record.phone = request.form.get('phone')
            faculty_record.designation = request.form.get('designation')
            
            # Update user information
            current_user.email = request.form.get('email')
            new_username = request.form.get('username')
            
            # Check if username is being changed and if it's already taken
            if new_username != current_user.username:
                existing_user = User.query.filter_by(username=new_username).first()
                if existing_user:
                    flash('Username already taken. Please choose another.', 'error')
                    return render_template('faculty/profile.html', faculty=faculty_record, user=current_user)
                current_user.username = new_username
            
            # Update password if provided
            new_password = request.form.get('new_password')
            confirm_password = request.form.get('confirm_password')
            
            if new_password:
                if new_password != confirm_password:
                    flash('Passwords do not match', 'error')
                    return render_template('faculty/profile.html', faculty=faculty_record, user=current_user)
                
                if len(new_password) < 6:
                    flash('Password must be at least 6 characters long', 'error')
                    return render_template('faculty/profile.html', faculty=faculty_record, user=current_user)
                
                current_user.password_hash = generate_password_hash(new_password)
            
            db.session.commit()
            flash('Profile updated successfully!', 'success')
            return redirect(url_for('faculty_profile'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating profile: {str(e)}', 'error')
            return render_template('faculty/profile.html', faculty=faculty_record, user=current_user)
    
    # Check if Class Teacher
    managed_section = Section.query.filter_by(class_teacher_id=faculty_record.faculty_id).first()
    is_class_teacher = managed_section is not None
    
    return render_template('faculty/profile.html', faculty=faculty_record, user=current_user, is_class_teacher=is_class_teacher)


@app.route('/api/faculty/sections-for-subject/<subject_id>')
@faculty_required
def api_faculty_sections_for_subject(subject_id):
    """
    Get sections assigned to the current faculty for a specific subject
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    allocations = SubjectAllocation.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        subject_id=subject_id
    ).all()
    
    sections = []
    seen = set()
    for a in allocations:
        if a.section and a.section_id not in seen:
            sections.append({
                'section_id': a.section_id,
                'section_name': a.section.section_name
            })
            seen.add(a.section_id)
            
    return jsonify(sections)


@app.route('/api/faculty/subjects-for-section/<section_id>')
@faculty_required
def api_faculty_subjects_for_section(section_id):
    """
    Get subjects assigned to the current faculty for a specific section
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    allocations = SubjectAllocation.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        section_id=section_id
    ).all()
    
    subjects = []
    seen = set()
    for a in allocations:
        if a.subject and a.subject_id not in seen and a.subject.subject_type != 'admin':
            subjects.append({
                'subject_id': a.subject_id,
                'subject_name': a.subject.subject_name,
                'subject_type': a.subject.subject_type
            })
            seen.add(a.subject_id)
            
    return jsonify(subjects)


@app.route('/api/attendance/submit', methods=['POST'])
@faculty_required
def api_submit_attendance():
    """
    Submit attendance for a class session with validation
    Ensures students can only be marked in sessions for their section or enrolled subjects
    """
    try:
        data = request.get_json()
        subject_id = data.get('subject_id')
        section_id = data.get('section_id')
        topic = data.get('topic', '')
        class_start_time = data.get('class_start_time')
        class_end_time = data.get('class_end_time')
        records = data.get('records', [])
        
        if not subject_id or not records:
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400
        
        # Get current faculty
        current_user = get_current_user()
        faculty = Faculty.query.filter_by(user_id=current_user.user_id).first()
        
        if not faculty:
            return jsonify({'success': False, 'error': 'Faculty not found'}), 404
        
        # Create or get schedule for this session
        schedule = None
        if section_id:
            # Try to find existing schedule
            schedule = ClassSchedule.query.filter_by(
                subject_id=subject_id,
                section_id=section_id,
                faculty_id=faculty.faculty_id,
                is_deleted=False
            ).first()
            
            # If no schedule exists, create a temporary one
            if not schedule:
                # Use the attendance date if provided, otherwise use today
                attendance_date_str = data.get('attendance_date')
                if attendance_date_str:
                    attendance_date = datetime.strptime(attendance_date_str, '%Y-%m-%d').date()
                else:
                    attendance_date = datetime.now().date()
                
                schedule = ClassSchedule(
                    subject_id=subject_id,
                    section_id=section_id,
                    faculty_id=faculty.faculty_id,
                    date=attendance_date,
                    start_time=datetime.strptime(class_start_time, '%H:%M').time() if class_start_time else None,
                    end_time=datetime.strptime(class_end_time, '%H:%M').time() if class_end_time else None
                )
                db.session.add(schedule)
                db.session.flush()  # Get the schedule_id
        
        # Create attendance session
        session_obj = AttendanceSession(
            schedule_id=schedule.schedule_id if schedule else None,
            taken_by_user_id=current_user.user_id,
            topic_taught=topic,
            status='finalized'
        )
        db.session.add(session_obj)
        db.session.flush()  # Get the session_id
        
        # Validate and create attendance records
        subject = Subject.query.get(subject_id)
        section = Section.query.get(section_id) if section_id else None
        
        # Generate program-specific diary number
        program_code = None
        if section and section.program:
            program_code = section.program.program_code
        session_obj.diary_number = AttendanceSession.generate_diary_number(program_code)
        
        valid_student_ids = set()
        
        # Get valid students for this session
        if section:
            # For section-based subjects: students in the section
            valid_student_ids.update([s.student_id for s in section.students.filter_by(is_deleted=False).all()])
        
        # Also include students enrolled in this subject (for electives/languages)
        from models import StudentSubjectEnrollment
        enrolled_students = StudentSubjectEnrollment.query.filter_by(
            subject_id=subject_id,
            is_deleted=False
        ).all()
        valid_student_ids.update([e.student_id for e in enrolled_students])
        
        # Create attendance records only for valid students
        created_count = 0
        skipped_count = 0
        
        for record in records:
            student_id = record.get('student_id')
            status = record.get('status', 'absent')
            
            # VALIDATION: Only create record if student is valid for this session
            if student_id not in valid_student_ids:
                skipped_count += 1
                continue
            
            attendance_record = AttendanceRecord(
                attendance_session_id=session_obj.attendance_session_id,
                student_id=student_id,
                status=status
            )
            db.session.add(attendance_record)
            created_count += 1
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Attendance submitted successfully! {created_count} records created.',
            'session_id': session_obj.attendance_session_id,
            'diary_number': session_obj.diary_number,
            'created_count': created_count,
            'skipped_count': skipped_count
        })
        
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/section/<section_id>/students')
@login_required # Allow faculty/admin
def api_get_section_students(section_id):
    """
    Get all students in a specific section
    For Swipe Attendance App - Updated to handle virtual sections
    """
    section = Section.query.get(section_id)
    if not section:
        return jsonify([]), 404
    
    # Use get_students method which handles both regular and virtual sections
    students = section.get_students()
    
    return jsonify([{
        'student_id': s.student_id,
        'name': s.name,
        'usn': s.roll_number or s.usn,
        'roll_number': s.roll_number
    } for s in students])


@app.route('/api/subject/<subject_id>/students')
@login_required
def api_get_subject_students(subject_id):
    """
    Get all students enrolled in a subject (for elective subjects)
    """
    try:
        from models import StudentSubjectEnrollment
        
        # Get students enrolled in this subject
        enrollments = StudentSubjectEnrollment.query.filter_by(
            subject_id=subject_id,
            is_deleted=False
        ).all()
        
        students_data = [{
            'student_id': e.student.student_id,
            'name': e.student.name,
            'usn': e.student.roll_number or e.student.usn,
            'roll_number': e.student.roll_number
        } for e in enrollments if e.student and not e.student.is_deleted]
        
        return jsonify(students_data)
    except Exception as e:
        return jsonify([]), 400

@app.route('/class-teacher/dashboard')
@faculty_required
def class_teacher_dashboard():
    """
    Class teacher dashboard showing today's check-ins and attendance comparison.
    """
    current_user = get_current_user()
    faculty = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty:
        flash('Faculty record not found', 'error')
        return redirect(url_for('index'))
    
    # Get sections where this faculty is class teacher
    sections = Section.query.filter_by(
        class_teacher_id=faculty.faculty_id,
        is_deleted=False
    ).all()
    
    if not sections:
        flash('You are not assigned as class teacher for any section', 'warning')
        return redirect(url_for('index'))
    
    # Get today's check-in data for each section
    today = date.today()
    section_data = []
    
    for section in sections:
        # Get all active students in section
        students = Student.query.filter_by(
            section_id=section.section_id,
            is_deleted=False,
            is_active=True
        ).all()
        
        total_students = len(students)
        student_ids = [s.student_id for s in students]
        
        # Get today's check-ins for these students
        checkins = CampusCheckIn.query.filter(
            CampusCheckIn.student_id.in_(student_ids),
            CampusCheckIn.checkin_date == today,
            CampusCheckIn.is_deleted == False
        ).all()
        
        checked_in_ids = {c.student_id for c in checkins}
        
        # Separate checked-in and not checked-in students
        checked_in_students = []
        not_checked_in_students = []
        
        for student in students:
            if student.student_id in checked_in_ids:
                checkin = next(c for c in checkins if c.student_id == student.student_id)
                checked_in_students.append({
                    'student': student,
                    'checkin_time': checkin.checkin_time
                })
            else:
                not_checked_in_students.append(student)
        
        section_data.append({
            'section': section,
            'total_students': total_students,
            'checked_in_count': len(checked_in_students),
            'checked_in_students': checked_in_students,
            'not_checked_in_students': not_checked_in_students
        })
    
    return render_template('class_teacher_dashboard.html',
                         faculty=faculty,
                         section_data=section_data,
                         today=today)


@app.route('/api/class-teacher/today-checkins')
@faculty_required
def api_class_teacher_checkins():
    """
    API endpoint for class teacher to get today's check-in data.
    """
    current_user = get_current_user()
    faculty = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty:
        return jsonify({'success': False, 'error': 'Faculty record not found'}), 404
    
    section_id = request.args.get('section_id')
    today = date.today()
    
    # Get section(s)
    if section_id:
        sections = Section.query.filter_by(
            section_id=section_id,
            class_teacher_id=faculty.faculty_id,
            is_deleted=False
        ).all()
    else:
        sections = Section.query.filter_by(
            class_teacher_id=faculty.faculty_id,
            is_deleted=False
        ).all()
    
    result = []
    for section in sections:
        students = Student.query.filter_by(
            section_id=section.section_id,
            is_deleted=False
        ).all()
        
        student_ids = [s.student_id for s in students]
        
        checkins = CampusCheckIn.query.filter(
            CampusCheckIn.student_id.in_(student_ids) if student_ids else False,
            CampusCheckIn.checkin_date == today,
            CampusCheckIn.is_deleted == False
        ).all()
        
        result.append({
            'section': section.to_dict(),
            'total_students': len(students),
            'checked_in_count': len(checkins),
            'checkins': [c.to_dict() for c in checkins]
        })
    
    return jsonify({'success': True, 'data': result})


# ============================================
# Monthly Attendance Report Routes
# ============================================

@app.route('/attendance/monthly-report')
@faculty_required
def monthly_attendance_report():
    """
    Monthly attendance report page for class teachers.
    """
    current_user = get_current_user()
    faculty = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    # Get sections for this faculty (class teacher sections)
    sections = Section.query.filter_by(
        class_teacher_id=faculty.faculty_id,
        is_deleted=False
    ).all() if faculty else []
    
    # If admin/HOD, show all sections
    if current_user.role and current_user.role.role_name.lower() in ['admin', 'hod']:
        sections = Section.query.filter_by(is_deleted=False).all()
    
    return render_template('monthly_report.html',
                         sections=sections,
                         current_month=datetime.now().month,
                         current_year=datetime.now().year)


@app.route('/api/attendance/monthly-report')
@faculty_required
def api_monthly_attendance_report():
    """
    API endpoint for generating monthly attendance report.
    Parameters: section_id, month, year
    """
    section_id = request.args.get('section_id')
    month = request.args.get('month', type=int, default=datetime.now().month)
    year = request.args.get('year', type=int, default=datetime.now().year)
    
    if not section_id:
        return jsonify({'success': False, 'error': 'Section ID required'}), 400
    
    # Verify permission
    current_user = get_current_user()
    faculty = Faculty.query.filter_by(user_id=current_user.user_id).first()
    section = Section.query.get(section_id)
    
    if not section:
        return jsonify({'success': False, 'error': 'Section not found'}), 404
    
    # Check if user has permission (class teacher, admin, or HOD)
    is_admin = current_user.role and current_user.role.role_name.lower() in ['admin', 'hod']
    is_class_teacher = faculty and section.class_teacher_id == faculty.faculty_id
    
    if not is_admin and not is_class_teacher:
        return jsonify({'success': False, 'error': 'Permission denied'}), 403
    
    # Get date range for the month
    first_day = date(year, month, 1)
    if month == 12:
        last_day = date(year + 1, 1, 1) - timedelta(days=1)
    else:
        last_day = date(year, month + 1, 1) - timedelta(days=1)
    
    # Get students in section
    students = Student.query.filter_by(
        section_id=section_id,
        is_deleted=False
    ).order_by(Student.roll_number).all()
    
    # Get all attendance sessions for the section in this month
    sessions = AttendanceSession.query.join(ClassSchedule).filter(
        ClassSchedule.section_id == section_id,
        AttendanceSession.taken_at >= datetime.combine(first_day, time.min),
        AttendanceSession.taken_at <= datetime.combine(last_day, time.max),
        AttendanceSession.is_deleted == False
    ).all()
    
    total_classes = len(sessions)
    session_ids = [s.attendance_session_id for s in sessions]
    
    # Calculate attendance for each student
    report_data = []
    for student in students:
        # Get ALL attendance records for this student in this month (across any section if transferred)
        # We filter by date range on the joined AttendanceSession
        records = AttendanceRecord.query.join(AttendanceSession).filter(
            AttendanceRecord.student_id == student.student_id,
            AttendanceSession.taken_at >= datetime.combine(first_day, time.min),
            AttendanceSession.taken_at <= datetime.combine(last_day, time.max),
            AttendanceRecord.is_deleted == False
        ).all()
        
        present_count = sum(1 for r in records if r.status == 'present')
        absent_count = sum(1 for r in records if r.status == 'absent')
        late_count = sum(1 for r in records if r.status == 'late')
        
        # Use student's total records as denominator if they have any (handles transfer cases)
        # If no records, use section's total classes (new student with 0 attendance)
        student_total_classes = len(records)
        effective_total = student_total_classes if student_total_classes > 0 else total_classes
        
        percentage = round((present_count / effective_total) * 100, 1) if effective_total > 0 else 0
        
        report_data.append({
            'student_id': student.student_id,
            'roll_number': student.roll_number,
            'name': f"{student.name}",
            'total_classes': total_classes,
            'present': present_count,
            'absent': absent_count,
            'late': late_count,
            'percentage': percentage
        })
    
    # Also include daily check-in data
    student_ids = [s.student_id for s in students]
    checkins = CampusCheckIn.query.filter(
        CampusCheckIn.student_id.in_(student_ids) if student_ids else False,
        CampusCheckIn.checkin_date >= first_day,
        CampusCheckIn.checkin_date <= last_day,
        CampusCheckIn.is_deleted == False
    ).all()
    
    # Count unique check-in days per student
    checkin_days = {}
    for checkin in checkins:
        if checkin.student_id not in checkin_days:
            checkin_days[checkin.student_id] = set()
        checkin_days[checkin.student_id].add(checkin.checkin_date)
    
    # Add check-in days to report
    for data in report_data:
        data['checkin_days'] = len(checkin_days.get(data['student_id'], set()))
    
    return jsonify({
        'success': True,
        'section': section.to_dict(),
        'month': month,
        'year': year,
        'total_classes': total_classes,
        'total_working_days': (last_day - first_day).days + 1,
        'report': report_data
    })


@app.route('/api/admin/campus-config', methods=['GET', 'PUT'])
@admin_required
def api_campus_config():
    """
    Admin endpoint to view/update campus configuration.
    """
    config = get_campus_config()
    
    if request.method == 'PUT':
        data = request.get_json()
        try:
            if 'campus_latitude' in data:
                config.campus_latitude = float(data['campus_latitude'])
            if 'campus_longitude' in data:
                config.campus_longitude = float(data['campus_longitude'])
            if 'campus_radius_meters' in data:
                config.campus_radius_meters = int(data['campus_radius_meters'])
            if 'college_name' in data:
                config.college_name = data['college_name']
            if 'checkin_start_time' in data:
                config.checkin_start_time = datetime.strptime(data['checkin_start_time'], '%H:%M').time()
            if 'checkin_end_time' in data:
                config.checkin_end_time = datetime.strptime(data['checkin_end_time'], '%H:%M').time()
            
            db.session.commit()
            return jsonify({'success': True, 'config': config.to_dict()})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 400
    
    return jsonify({'success': True, 'config': config.to_dict()})




from math import radians, sin, cos, sqrt, atan2

def calculate_distance(lat1, lon1, lat2, lon2):
    """
    Calculate distance between two GPS coordinates using Haversine formula
    Returns distance in meters
    """
    R = 6371000  # Earth's radius in meters
    
    lat1, lon1, lat2, lon2 = map(radians, [lat1, lon1, lat2, lon2])
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    distance = R * c
    
    return distance


@app.route('/api/faculty/check-in', methods=['POST'])
@faculty_required
def api_faculty_checkin():
    """Faculty check-in (can be done from anywhere)"""
    from models import FacultyAttendance
    
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    data = request.get_json()
    latitude = data.get('latitude')
    longitude = data.get('longitude')
    
    if not latitude or not longitude:
        return jsonify({'success': False, 'error': 'GPS location required'}), 400
    
    today = date.today()
    
    # Check if already checked in today
    existing = FacultyAttendance.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        date=today
    ).first()
    
    if existing and existing.check_in_time:
        return jsonify({'success': False, 'error': 'Already checked in today'}), 400
    
    # Create or update attendance record
    if not existing:
        attendance = FacultyAttendance(
            faculty_id=faculty_record.faculty_id,
            date=today
        )
    else:
        attendance = existing
    
    attendance.check_in_time = datetime.now()
    attendance.check_in_latitude = latitude
    attendance.check_in_longitude = longitude
    
    if not existing:
        db.session.add(attendance)
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'check_in_time': attendance.check_in_time.isoformat(),
        'message': 'Checked in successfully!'
    })


@app.route('/api/faculty/check-out', methods=['POST'])
@faculty_required
def api_faculty_checkout():
    """Faculty check-out (must be at department location)"""
    from models import FacultyAttendance, CollegeConfig
    
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    data = request.get_json()
    latitude = data.get('latitude')
    longitude = data.get('longitude')
    
    if not latitude or not longitude:
        return jsonify({'success': False, 'error': 'GPS location required'}), 400
    
    today = date.today()
    
    # Get today's attendance record
    attendance = FacultyAttendance.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        date=today
    ).first()
    
    if not attendance or not attendance.check_in_time:
        return jsonify({'success': False, 'error': 'Please check in first'}), 400
    
    if attendance.check_out_time:
        return jsonify({'success': False, 'error': 'Already checked out today'}), 400
    
    # Get campus location from config
    config = CollegeConfig.query.first()
    if not config or not config.campus_latitude or not config.campus_longitude:
        # If no config, allow checkout (fallback)
        attendance.check_out_time = datetime.now()
        attendance.check_out_latitude = latitude
        attendance.check_out_longitude = longitude
        attendance.check_out_valid = True
        db.session.commit()
        return jsonify({
            'success': True,
            'message': 'Checked out successfully!',
            'hours_worked': attendance.get_hours_worked()
        })
    
    # Calculate distance from campus
    distance = calculate_distance(
        latitude, longitude,
        config.campus_latitude, config.campus_longitude
    )
    
    radius = config.campus_radius_meters or 100  # Default 100m
    
    if distance > radius:
        return jsonify({
            'success': False,
            'error': f'You are not at the department! You are {int(distance)}m away. Please come to campus to check out.',
            'distance': int(distance),
            'required_radius': radius
        }), 400
    
    # Valid checkout
    attendance.check_out_time = datetime.now()
    attendance.check_out_latitude = latitude
    attendance.check_out_longitude = longitude
    attendance.check_out_valid = True
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'check_out_time': attendance.check_out_time.isoformat(),
        'hours_worked': attendance.get_hours_worked(),
        'message': 'Checked out successfully!'
    })


@app.route('/api/faculty/attendance-status')
@faculty_required
def api_faculty_attendance_status():
    """Get today's check-in/out status"""
    from models import FacultyAttendance
    
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    today = date.today()
    attendance = FacultyAttendance.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        date=today
    ).first()
    
    if not attendance:
        return jsonify({
            'checked_in': False,
            'checked_out': False
        })
    
    return jsonify({
        'checked_in': bool(attendance.check_in_time),
        'checked_out': bool(attendance.check_out_time),
        'check_in_time': attendance.check_in_time.isoformat() if attendance.check_in_time else None,
        'check_out_time': attendance.check_out_time.isoformat() if attendance.check_out_time else None,
        'hours_worked': attendance.get_hours_worked()
    })


@app.route('/api/faculty/students/<student_id>/analytics')
@faculty_required
def api_faculty_student_analytics(student_id):
    """
    Get detailed attendance analytics for a student
    """
    print(f"=" * 80)
    print(f"ANALYTICS ENDPOINT CALLED - Student ID: {student_id}")
    print(f"=" * 80)
    
    try:
        from models import Student, AttendanceRecord, AttendanceSession, ClassSchedule, Subject
        
        print(f"Step 1: Fetching student with ID: {student_id}")
        student = Student.query.get(student_id)
        if not student:
            print(f"ERROR: Student not found with ID: {student_id}")
            return jsonify({'success': False, 'error': 'Student not found'}), 404
        
        print(f"Step 2: Found student: {student.name}")
        print(f"Step 3: Student section_id: {student.section_id}")
        
        # Get the student's section to find the semester
        from models import Section
        section = Section.query.get(student.section_id)
        if not section:
            return jsonify({'success': False, 'error': 'Section not found'}), 404
        
        print(f"Step 4: Section current_semester: {section.current_semester}")
        
        # Get ALL subjects for this semester
        subjects = Subject.query.filter_by(semester_id=section.current_semester).all()
        
        print(f"Step 5: Found {len(subjects)} subjects for semester {section.current_semester}")
        
        # Initialize stats for ALL subjects in this semester
        subject_stats = {}
        for subject in subjects:
            subject_stats[subject.subject_id] = {
                'subject_name': subject.subject_name,
                'subject_code': subject.subject_code,
                'total_classes': 0,
                'attended': 0
            }
        
        # Get all sessions for the student's section
        # This counts ALL classes conducted, even if multiple classes of same subject on same day
        all_sessions = db.session.query(
            AttendanceSession, 
            ClassSchedule,
            Subject
        ).join(
            ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
        ).join(
            Subject, ClassSchedule.subject_id == Subject.subject_id
        ).filter(
            ClassSchedule.section_id == student.section_id,
            AttendanceSession.is_deleted == False
        ).all()
        
        print(f"Step 5: Found {len(all_sessions)} total sessions conducted")
        
        # Get all attendance records for this student
        student_records = AttendanceRecord.query.filter_by(
            student_id=student_id,
            is_deleted=False
        ).all()
        
        print(f"Step 6: Found {len(student_records)} attendance records for student")
        
        # Create a map of session_id -> status
        attendance_map = {r.attendance_session_id: r.status for r in student_records}
        
        # Process Stats - Count EVERY session (allows multiple classes per day)
        day_analysis = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0} # 0=Mon, 6=Sun
        days_map = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
        
        for session, schedule, subject in all_sessions:
            sub_id = subject.subject_id
            
            # Only count if this subject is allocated to the section
            if sub_id in subject_stats:
                # Count every session (even if multiple per day)
                subject_stats[sub_id]['total_classes'] += 1
                
                # Check attendance
                status = attendance_map.get(session.attendance_session_id, 'absent')
                
                if status == 'present':
                    subject_stats[sub_id]['attended'] += 1
                else:
                    # Analyze absence
                    day_idx = session.taken_at.weekday()
                    day_analysis[day_idx] += 1
                
        print(f"Step 6: Processed {len(subject_stats)} subjects")
        
        # Format Subject Stats
        final_subject_stats = []
        for sub_id, stats in subject_stats.items():
            total = stats['total_classes']
            attended = stats['attended']
            percentage = round((attended / total * 100), 1) if total > 0 else 0
            
            final_subject_stats.append({
                'subject_name': stats['subject_name'],
                'subject_code': stats['subject_code'],
                'total_classes': total,
                'attended': attended,
                'percentage': percentage
            })
            
        # Format Day Analysis (only if count > 0)
        final_day_analysis = []
        total_absences = sum(day_analysis.values())
        
        for day_idx, count in day_analysis.items():
            if count > 0:
                percentage = round((count / total_absences * 100), 1) if total_absences > 0 else 0
                final_day_analysis.append({
                    'day': days_map[day_idx],
                    'count': count,
                    'percentage': percentage
                })
                
        # Sort days by count desc
        final_day_analysis.sort(key=lambda x: x['count'], reverse=True)
        
        print(f"Step 7: Returning success response")
        print(f"=" * 80)
        
        return jsonify({
            'success': True,
            'student_name': student.name,
            'roll_number': student.roll_number,
            'subject_stats': final_subject_stats,
            'day_analysis': final_day_analysis
        })
    except Exception as e:
        import traceback
        print(f"=" * 80)
        print(f"ANALYTICS ERROR:")
        print(f"Error Type: {type(e).__name__}")
        print(f"Error Message: {str(e)}")
        print(f"Full Traceback:")
        print(traceback.format_exc())
        print(f"=" * 80)
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/faculty/students/<student_id>', methods=['PUT'])
@faculty_required
def api_faculty_update_student(student_id):
    """
    Update student details (Phone, Address)
    """
    from models import Student
    
    student = Student.query.get_or_404(student_id)
    data = request.get_json()
    
    if 'phone' in data:
        student.phone = data['phone']
    if 'guardian_phone' in data:
        student.guardian_phone = data['guardian_phone']
    if 'address' in data:
        student.address = data['address']
        
    try:
        db.session.commit()
        return jsonify({'success': True, 'message': 'Updated successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


# ============================================
# Main Application Entry
# ============================================


# ============================================
# Student Dashboard Routes
# ============================================

# Motivational quotes list
MOTIVATIONAL_QUOTES = [
    "Success is not final, failure is not fatal: it is the courage to continue that counts.",
    "Believe you can and you're halfway there.",
    "The future belongs to those who believe in the beauty of their dreams.",
    "Education is the most powerful weapon which you can use to change the world.",
    "Your attitude, not your aptitude, will determine your altitude.",
    "The only way to do great work is to love what you do.",
    "Don't watch the clock; do what it does. Keep going.",
    "The expert in anything was once a beginner.",
    "Success is the sum of small efforts repeated day in and day out.",
    "Dream big, work hard, stay focused, and surround yourself with good people.",
    "The beautiful thing about learning is that no one can take it away from you.",
    "Strive for progress, not perfection.",
    "Your only limit is you.",
    "Great things never come from comfort zones.",
    "The harder you work for something, the greater you'll feel when you achieve it."
]

@app.route('/student/dashboard')
@student_required
def student_dashboard():
    """
    Student Dashboard - Main page for students
    Shows attendance analytics, holidays, check-in status, and motivational quote
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        flash('Student record not found.', 'danger')
        return redirect(url_for('login'))
    
    # Get student's section
    section = Section.query.get(student_record.section_id)
    
    # Get Enrolled Elective Sections
    enrollments = StudentSubjectEnrollment.query.filter_by(
        student_id=student_record.student_id,
        semester=section.current_semester
    ).all()
    enrolled_section_ids = [e.section_id for e in enrollments if e.section_id]
    all_section_ids = [student_record.section_id] + enrolled_section_ids

    # Calculate overall attendance percentage (Core + Electives)
    total_sessions = db.session.query(AttendanceSession.attendance_session_id).join(ClassSchedule).filter(
        ClassSchedule.section_id.in_(all_section_ids),
        AttendanceSession.is_deleted == False
    ).distinct().count()
    
    attended = db.session.query(AttendanceSession.attendance_session_id).join(
        AttendanceRecord, AttendanceRecord.attendance_session_id == AttendanceSession.attendance_session_id
    ).join(ClassSchedule).filter(
        AttendanceRecord.student_id == student_record.student_id,
        AttendanceRecord.status == 'present',
        ClassSchedule.section_id.in_(all_section_ids)
    ).distinct().count()
    
    overall_percentage = int((attended / total_sessions * 100)) if total_sessions > 0 else 0
    
    # Check if in red zone (<75%)
    is_red_zone = overall_percentage < 75
    
    # Get random motivational quote
    import random
    motivational_quote = random.choice(MOTIVATIONAL_QUOTES)
    

    # --- Fetch Subjects and Calculate Per-Subject Attendance ---
    # 1. Compulsory Subjects
    compulsory_subjects = Subject.query.filter(
        Subject.program_id == section.program_id,
        Subject.semester_id == section.current_semester,
        Subject.subject_category == 'compulsory'
    ).all()
    
    # 2. Enrolled Elective Subjects
    # Filter by Semester instead of Academic Year to avoid string mismatch issues
    enrolled_electives = db.session.query(Subject).join(StudentSubjectEnrollment).filter(
        StudentSubjectEnrollment.student_id == student_record.student_id,
        StudentSubjectEnrollment.semester == section.current_semester
    ).all()
    
    all_subjects_objs = compulsory_subjects + enrolled_electives
    
    subjects_data = []
    for subj in all_subjects_objs:
        # Verify if subject has a virtual section (for electives) or uses main section
        target_section_id = student_record.section_id # Default to core section
        
        # If elective, find the virtual section id
        if subj.subject_category != 'compulsory':
            # Find the enrollment for this subject to get section_id
            enrollment = next((e for e in enrollments if e.subject_id == subj.subject_id), None)
            if enrollment and enrollment.section_id:
                target_section_id = enrollment.section_id
        
        # Get total classes for this subject/section
        # Note: This might be slightly inefficient in a loop, but safe for N < 10 subjects
        subj_total = db.session.query(AttendanceSession).join(ClassSchedule).filter(
            ClassSchedule.section_id == target_section_id,
            ClassSchedule.subject_id == subj.subject_id,
            AttendanceSession.is_deleted == False
        ).count()
        
        # Get attended classes
        subj_attended = db.session.query(AttendanceRecord).join(AttendanceSession).join(ClassSchedule).filter(
            AttendanceRecord.student_id == student_record.student_id,
            AttendanceRecord.status == 'present',
            AttendanceRecord.is_deleted == False,
            ClassSchedule.subject_id == subj.subject_id,
            ClassSchedule.section_id == target_section_id
        ).count()
        
        pct = int((subj_attended / subj_total * 100)) if subj_total > 0 else 0
        
        subjects_data.append({
            'subject_name': subj.subject_name,
            'subject_code': subj.subject_code,
            'attendance_pct': pct if subj_total > 0 else None
        })
    
    # Check if today is a holiday
    today = date.today()
    holiday = Holiday.query.filter_by(holiday_date=today).first()
    is_holiday = holiday is not None
    
    dashboard_data = {
        'student': student_record,
        'section': section,
        'overall_percentage': overall_percentage,
        'total_sessions': total_sessions,
        'attended': attended,
        'is_red_zone': is_red_zone,
        'motivational_quote': motivational_quote,
        'subjects': subjects_data,
        'is_holiday': is_holiday,
        'holiday': holiday
    }
    
    return render_template('student/dashboard.html', **dashboard_data)


@app.route('/api/student/analytics')
@student_required
def api_student_analytics():
    """
    Get detailed attendance analytics for the logged-in student
    Returns subject-wise breakdown with percentages
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
    
    try:
        # Get student's section
        section = Section.query.get(student_record.section_id)
        if not section:
            return jsonify({'success': False, 'error': 'Section not found'}), 404
        
        # Get ALL subjects for this semester (Compulsory + Enrolled Electives)
        # 1. Compulsory Subjects
        compulsory_subjects = Subject.query.filter(
            Subject.program_id == section.program_id,
            Subject.semester_id == section.current_semester,
            Subject.subject_category == 'compulsory'
        ).all()
        
        # 2. Enrolled Elective Subjects
        enrolled_subjects = db.session.query(Subject).join(StudentSubjectEnrollment).filter(
            StudentSubjectEnrollment.student_id == student_record.student_id,
            StudentSubjectEnrollment.semester == section.current_semester
        ).all()
        
        subjects = compulsory_subjects + enrolled_subjects
        
        # Initialize stats for ALL subjects in this semester
        subject_stats = {}
        for subject in subjects:
            subject_stats[subject.subject_id] = {
                'subject_name': subject.subject_name,
                'subject_code': subject.subject_code,
                'total_classes': 0,
                'attended': 0
            }
        
        # Get all sessions for the student's section
        all_sessions = db.session.query(
            AttendanceSession, 
            ClassSchedule,
            Subject
        ).join(
            ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
        ).join(
            Subject, ClassSchedule.subject_id == Subject.subject_id
        ).filter(
            ClassSchedule.section_id == student_record.section_id,
            AttendanceSession.is_deleted == False,
            ClassSchedule.is_deleted == False,
            Subject.is_deleted == False
        ).all()
        
        # ALSO get sessions for enrolled subjects (for virtual sections/electives)
        # These sessions might not have a schedule_id or might be for virtual sections
        enrolled_subject_ids = [s.subject_id for s in enrolled_subjects]
        if enrolled_subject_ids:
            enrolled_sessions = db.session.query(
                AttendanceSession,
                Subject
            ).join(
                ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
            ).join(
                Subject, ClassSchedule.subject_id == Subject.subject_id
            ).filter(
                Subject.subject_id.in_(enrolled_subject_ids),
                AttendanceSession.is_deleted == False,
                ClassSchedule.is_deleted == False,
                Subject.is_deleted == False
            ).all()
            
            # Add these to all_sessions (with None for schedule since we already have it)
            for session, subject in enrolled_sessions:
                # Check if this session is already in all_sessions
                session_ids = [s[0].attendance_session_id for s in all_sessions]
                if session.attendance_session_id not in session_ids:
                    all_sessions.append((session, None, subject))
        
        # Get all attendance records for this student
        student_records = AttendanceRecord.query.filter_by(
            student_id=student_record.student_id,
            is_deleted=False
        ).all()
        
        # Create a map of session_id -> status
        attendance_map = {r.attendance_session_id: r.status for r in student_records}
        
        # Process Stats - Count EVERY session
        for session_data in all_sessions:
            session = session_data[0]
            schedule = session_data[1] if len(session_data) > 1 else None
            subject = session_data[2] if len(session_data) > 2 else None
            
            sub_id = subject.subject_id if subject else None
            
            # Only count if this subject is in the semester
            if sub_id and sub_id in subject_stats:
                subject_stats[sub_id]['total_classes'] += 1
                
                # Check attendance
                status = attendance_map.get(session.attendance_session_id, 'absent')
                
                if status == 'present':
                    subject_stats[sub_id]['attended'] += 1
        
        # Format Subject Stats
        final_subject_stats = []
        for sub_id, stats in subject_stats.items():
            total = stats['total_classes']
            attended = stats['attended']
            percentage = round((attended / total * 100), 1) if total > 0 else 0
            
            final_subject_stats.append({
                'subject_id': sub_id,
                'subject_name': stats['subject_name'],
                'subject_code': stats['subject_code'],
                'total_classes': total,
                'attended': attended,
                'percentage': percentage
            })
        
        # Calculate overall percentage
        total_all = sum(s['total_classes'] for s in final_subject_stats)
        attended_all = sum(s['attended'] for s in final_subject_stats)
        overall_percentage = round((attended_all / total_all * 100), 1) if total_all > 0 else 0
        
        return jsonify({
            'success': True,
            'student_name': f"{student_record.name}",
            'roll_number': student_record.roll_number,
            'overall_percentage': overall_percentage,
            'is_red_zone': overall_percentage < 75,
            'subject_stats': final_subject_stats
        })
    except Exception as e:
        import traceback
        print(f"Analytics Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/student/subject/<subject_id>/history')
@student_required
def api_student_subject_history(subject_id):
    """
    Get detailed class history for a specific subject for the logged-in student
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
        
    try:
        subject = Subject.query.get_or_404(subject_id)
        
        # Determine the target section (consider electives)
        enrollment = StudentSubjectEnrollment.query.filter_by(
            student_id=student_record.student_id,
            subject_id=subject_id,
            is_deleted=False
        ).first()
        
        target_section_id = enrollment.section_id if enrollment and enrollment.section_id else student_record.section_id
        
        # Get Teacher Name - Check both section-specific and global allocations
        # Try to find allocation for this specific section OR a global allocation (section_id is NULL)
        allocation = SubjectAllocation.query.filter(
            SubjectAllocation.subject_id == subject_id,
            db.or_(
                SubjectAllocation.section_id == target_section_id,
                SubjectAllocation.section_id == None
            )
        ).first()
        
        teacher_name = None
        if allocation and allocation.faculty:
            faculty = allocation.faculty
            teacher_name = faculty.name or f"{faculty.first_name} {faculty.last_name}"
        
        # If no allocation found, get teacher from the most recent attendance session
        if not teacher_name:
            recent_session = db.session.query(AttendanceSession).join(
                ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
            ).filter(
                ClassSchedule.subject_id == subject_id,
                ClassSchedule.section_id == target_section_id,
                AttendanceSession.is_deleted == False
            ).order_by(AttendanceSession.taken_at.desc()).first()
            
            if recent_session and recent_session.taken_by:
                faculty = Faculty.query.filter_by(user_id=recent_session.taken_by.user_id).first()
                if faculty:
                    teacher_name = faculty.name or f"{faculty.first_name} {faculty.last_name}"
        
        # Final fallback
        if not teacher_name:
            teacher_name = "Not Assigned"
        
        # Get all attendance sessions for this subject/section
        # Join with ClassSchedule and AttendanceRecord
        sessions = db.session.query(
            AttendanceSession,
            AttendanceRecord.status
        ).join(
            ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
        ).outerjoin(
            AttendanceRecord, db.and_(
                AttendanceRecord.attendance_session_id == AttendanceSession.attendance_session_id,
                AttendanceRecord.student_id == student_record.student_id
            )
        ).filter(
            ClassSchedule.subject_id == subject_id,
            ClassSchedule.section_id == target_section_id,
            AttendanceSession.is_deleted == False
        ).order_by(AttendanceSession.taken_at.desc()).all()
        
        history = []
        for session, status in sessions:
            history.append({
                'date': session.taken_at.strftime('%d %b, %Y'),
                'time': session.taken_at.strftime('%I:%M %p'),
                'topic': session.topic_taught or "Regular Class",
                'status': status or "absent"  # Default to absent if no record found
            })
            
        return jsonify({
            'success': True,
            'subject_name': subject.subject_name,
            'subject_code': subject.subject_code,
            'teacher_name': teacher_name,
            'history': history
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 400




@app.route('/api/student/check-in', methods=['POST'])
@student_required
def api_student_checkin():
    """
    Student campus check-in with GPS verification
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
    
    data = request.get_json()
    latitude = data.get('latitude')
    longitude = data.get('longitude')
    
    if not latitude or not longitude:
        return jsonify({'success': False, 'error': 'GPS location required'}), 400
    
    today = date.today()
    
    try:
        # Check if already checked in today
        existing = CampusCheckIn.query.filter_by(
            student_id=student_record.student_id,
            checkin_date=today
        ).first()
        
        if existing:
            return jsonify({'success': False, 'error': 'Already checked in today'}), 400
        
        # Get campus location from config
        config = CollegeConfig.query.first()
        
        # Ensure lat/long are floats
        try:
            latitude = float(latitude)
            longitude = float(longitude)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'error': 'Invalid GPS coordinates'}), 400
        
        # Verify location if config exists and has valid coordinates
        if config and config.campus_latitude and config.campus_longitude:
            from math import radians, sin, cos, sqrt, atan2
            
            # Calculate distance
            R = 6371000  # Earth's radius in meters
            lat1, lon1, lat2, lon2 = map(radians, [latitude, longitude, config.campus_latitude, config.campus_longitude])
            dlat = lat2 - lat1
            dlon = lon2 - lon1
            a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
            c = 2 * atan2(sqrt(a), sqrt(1-a))
            distance = R * c
            
            radius = config.campus_radius_meters or 500  # Default 500m
            
            if distance > radius:
                return jsonify({
                    'success': False,
                    'error': f'You are not at campus! You are {int(distance)}m away.',
                    'distance': int(distance)
                }), 400
        
        # Create check-in record
        now_dt = datetime.now()
        checkin = CampusCheckIn(
            student_id=student_record.student_id,
            checkin_date=today,
            checkin_time=now_dt.time(), # Use .time() to match db.Time column
            latitude=latitude,
            longitude=longitude
        )
        
        db.session.add(checkin)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Checked in successfully!',
            'checkin_time': now_dt.strftime('%I:%M %p')
        })
    except Exception as e:
        print(f"CHECKIN ERROR: {str(e)}") # Debug to console
        db.session.rollback()
        return jsonify({'success': False, 'error': f'Server Error: {str(e)}'}), 500


@app.route('/api/student/check-out', methods=['POST'])
@student_required
def api_student_checkout():
    """
    Student campus check-out
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
        
    data = request.get_json()
    latitude = data.get('latitude')
    longitude = data.get('longitude')
    
    today = date.today()
    
    # Find today's check-in
    checkin = CampusCheckIn.query.filter_by(
        student_id=student_record.student_id,
        checkin_date=today
    ).first()
    
    if not checkin:
        return jsonify({'success': False, 'error': 'You have not checked in today!'}), 400
        
    if checkin.checkout_time:
        return jsonify({'success': False, 'error': 'Already checked out today'}), 400

    if latitude and longitude:
        print(f"CHECKOUT ATTEMPT for {student_record.roll_number}: Lat={latitude}, Lon={longitude}")
        
    try:
        now_dt = datetime.now()
        checkin.checkout_time = now_dt.time()
        
        # Save checkout location if provided
        if latitude and longitude:
            try:
                checkin.checkout_latitude = float(latitude)
                checkin.checkout_longitude = float(longitude)
            except:
                pass # Ignore invalid coords for checkout
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Checked out successfully!',
            'checkout_time': now_dt.strftime('%I:%M %p')
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': f'Server Error: {str(e)}'}), 500


@app.route('/admin/config', methods=['GET', 'POST'])
@admin_required
def admin_config():
    """
    Admin page to configure campus GPS location
    """
    config = CollegeConfig.query.first()
    
    if request.method == 'POST':
        try:
            latitude = float(request.form.get('latitude'))
            longitude = float(request.form.get('longitude'))
            radius = int(request.form.get('radius'))
            
            if not config:
                config = CollegeConfig(
                    campus_latitude=latitude,
                    campus_longitude=longitude,
                    campus_radius_meters=radius
                )
                db.session.add(config)
            else:
                config.campus_latitude = latitude
                config.campus_longitude = longitude
                config.campus_radius_meters = radius
                
            db.session.commit()
            flash('Campus configuration updated successfully!', 'success')
        except ValueError:
            flash('Invalid input values. Please enter valid numbers.', 'danger')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating config: {str(e)}', 'danger')
            
    return render_template('admin_config.html', config=config)


@app.route('/api/student/recent-absences')
@student_required
def api_student_recent_absences():
    """
    Get student's absences from the past 3 days
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
    
    # Get absences from past 3 days
    three_days_ago = date.today() - timedelta(days=3)
    
    absences = db.session.query(
        AttendanceRecord,
        AttendanceSession,
        ClassSchedule,
        Subject
    ).join(
        AttendanceSession, AttendanceRecord.attendance_session_id == AttendanceSession.attendance_session_id
    ).join(
        ClassSchedule, AttendanceSession.schedule_id == ClassSchedule.schedule_id
    ).join(
        Subject, ClassSchedule.subject_id == Subject.subject_id
    ).filter(
        AttendanceRecord.student_id == student_record.student_id,
        AttendanceRecord.status == 'absent',
        AttendanceSession.taken_at >= three_days_ago,
        AttendanceRecord.is_deleted == False
    ).order_by(AttendanceSession.taken_at.desc()).all()
    
    absence_list = []
    for record, session, schedule, subject in absences:
        absence_list.append({
            'date': session.taken_at.strftime('%d %b %Y'),
            'day': session.taken_at.strftime('%A'),
            'subject_name': subject.subject_name,
            'subject_code': subject.subject_code
        })
    
    return jsonify({
        'success': True,
        'absences': absence_list
    })


@app.route('/api/student/holidays')
@student_required
def api_student_holidays():
    """
    Get upcoming and recent holidays
    """
    from models import Holiday
    
    today = date.today()
    tomorrow = today + timedelta(days=1)
    
    # Get upcoming holidays (next 30 days)
    upcoming = Holiday.query.filter(
        Holiday.holiday_date >= today,
        Holiday.holiday_date <= today + timedelta(days=30)
    ).order_by(Holiday.holiday_date).all()
    
    # Get recent holidays (past 7 days)
    recent = Holiday.query.filter(
        Holiday.holiday_date >= today - timedelta(days=7),
        Holiday.holiday_date < today
    ).order_by(Holiday.holiday_date.desc()).all()
    
    # Check if tomorrow is a holiday
    tomorrow_holiday = Holiday.query.filter_by(holiday_date=tomorrow).first()
    
    upcoming_list = [{
        'date': h.holiday_date.strftime('%d %b %Y'),
        'day': h.holiday_date.strftime('%A'),
        'name': h.holiday_name,
        'is_tomorrow': h.holiday_date == tomorrow
    } for h in upcoming]
    
    recent_list = [{
        'date': h.holiday_date.strftime('%d %b %Y'),
        'day': h.holiday_date.strftime('%A'),
        'name': h.holiday_name
    } for h in recent]
    
    return jsonify({
        'success': True,
        'tomorrow_is_holiday': bool(tomorrow_holiday),
        'tomorrow_holiday_name': tomorrow_holiday.holiday_name if tomorrow_holiday else None,
        'upcoming': upcoming_list,
        'recent': recent_list
    })




MOTIVATIONAL_QUOTES = [
    "The only way to do great work is to love what you do. â€“ Steve Jobs",
    "Believe you can and you're halfway there. â€“ Theodore Roosevelt",
    "Success is not final, failure is not fatal: It is the courage to continue that counts. â€“ Winston Churchill",
    "Don't watch the clock; do what it does. Keep going. â€“ Sam Levenson",
    "The future belongs to those who believe in the beauty of their dreams. â€“ Eleanor Roosevelt",
    "Education is the most powerful weapon which you can use to change the world. â€“ Nelson Mandela",
    "Expert in anything was once a beginner. â€“ Helen Hayes",
    "There are no shortcuts to any place worth going. â€“ Beverly Sills",
    "Motivation is what gets you started. Habit is what keeps you going. â€“ Jim Ryun",
    "Your time is limited, so don't waste it living someone else's life. â€“ Steve Jobs",
    "The beautiful thing about learning is that no one can take it away from you. â€“ B.B. King",
    "Start where you are. Use what you have. Do what you can. â€“ Arthur Ashe",
    "Success doesn't come to you, you've got to go to it. â€“ Marva Collins",
    "Dream big and dare to fail. â€“ Norman Vaughan",
    "It always seems impossible until it's done. â€“ Nelson Mandela"
]

@app.route('/api/student/motivational-quote')
@student_required
def api_student_motivational_quote():
    """
    Get a random motivational quote
    """
    import random
    quote = random.choice(MOTIVATIONAL_QUOTES)
    
    return jsonify({
        'success': True,
        'quote': quote
    })


@app.route('/student/profile')
@student_required
def student_profile():
    """
    Student profile view and edit page
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        flash('Student record not found.', 'danger')
        return redirect(url_for('login'))
    
    section = Section.query.get(student_record.section_id)
    
    return render_template('student/profile.html', student=student_record, section=section)


@app.route('/api/student/profile', methods=['PUT'])
@student_required
def api_student_update_profile():
    """
    Update student profile (phone, address only)
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
    
    data = request.get_json()
    
    # Students can only update specific fields
    if 'phone' in data:
        student_record.phone = data['phone']
    if 'address' in data:
        student_record.address = data['address']
    
    try:
        db.session.commit()
        return jsonify({'success': True, 'message': 'Profile updated successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400



# ============================================
# Parent Dashboard Routes
# ============================================



# ============================================
# Fee Management Page Routes
# ============================================

@app.route('/student/fees')
@student_required
def student_fees_page():
    """Student fee management page"""
    from models import Student, FeeStructure
    import json

    current_user = get_current_user()
    student = Student.query.filter_by(user_id=current_user.user_id).first()

    if not student:
        flash('Student record not found', 'error')
        return redirect(url_for('student_dashboard'))
    
    # Get all fee structures for this student
    fee_structures = FeeStructure.query.filter_by(
        student_id=student.student_id,
        is_deleted=False
    ).order_by(FeeStructure.academic_year.desc()).all()
    
    # Debug: Print what we found
    print(f"DEBUG: Student ID: {student.student_id}")
    print(f"DEBUG: Found {len(fee_structures)} fee structures")
    for fs in fee_structures:
        print(f"  - {fs.academic_year}: Base={fs.base_fees}, Total={fs.total_fees}")
    
    # Parse additional fees from JSON and fetch receipts
    for fee_structure in fee_structures:
        # Additional fees list
        if fee_structure.additional_fees:
            try:
                fee_structure.additional_fees_list = json.loads(fee_structure.additional_fees)
            except:
                fee_structure.additional_fees_list = []
        else:
            fee_structure.additional_fees_list = []
            
        # Fetch associated receipts
        from models import FeeReceipt
        fee_structure.receipts = FeeReceipt.query.filter_by(
            fee_structure_id=fee_structure.fee_structure_id,
            is_deleted=False
        ).order_by(FeeReceipt.payment_date.desc()).all()
    
    return render_template('student/fees.html', 
                         student=student,
                         fee_structures=fee_structures)


@app.route('/api/student/fees/submit-receipt', methods=['POST'])
@student_required
def api_submit_fee_receipt():
    """Submit a fee receipt for verification"""
    from models import Student, FeeStructure, FeeReceipt
    import datetime
    
    current_user = get_current_user()
    student = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student:
        return jsonify({'success': False, 'error': 'Student record not found'}), 404
        
    try:
        data = request.get_json()
        fee_structure_id = data.get('fee_structure_id')
        receipt_number = data.get('receipt_number')
        receipt_phone = data.get('receipt_phone')
        payment_date_str = data.get('payment_date')
        amount_paid = data.get('amount_paid')
        
        if not all([fee_structure_id, receipt_number, receipt_phone, payment_date_str, amount_paid]):
            return jsonify({'success': False, 'error': 'All fields are required'}), 400
            
        fee_structure = FeeStructure.query.get(fee_structure_id)
        if not fee_structure or fee_structure.student_id != student.student_id:
            return jsonify({'success': False, 'error': 'Invalid fee structure'}), 400
            
        # Check for duplicate receipt number
        existing_receipt = FeeReceipt.query.filter_by(receipt_number=receipt_number).first()
        if existing_receipt:
            return jsonify({'success': False, 'error': 'Receipt number already submitted'}), 400
            
        payment_date = datetime.datetime.strptime(payment_date_str, '%Y-%m-%d').date()
        
        new_receipt = FeeReceipt(
            fee_structure_id=fee_structure_id,
            student_id=student.student_id,
            receipt_number=receipt_number,
            receipt_phone=receipt_phone,
            amount_paid=float(amount_paid),
            payment_date=payment_date,
            payment_mode='Online/Self-Reported',
            entered_by_user_id=current_user.user_id,
            entered_by_role='student',
            approved=False # Requires faculty approval
        )
        
        db.session.add(new_receipt)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Receipt submitted successfully for verification.'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/admin/holidays')
@admin_required
def admin_holidays():
    """Admin holiday management page"""
    academic_year = SystemSettings.get_current_academic_year()
    return render_template('admin_holidays.html', academic_year=academic_year)


@app.route('/api/admin/holidays', methods=['GET'])
@admin_required
def api_get_holidays():
    from models import Holiday
    holidays = Holiday.query.order_by(Holiday.holiday_date).all()
    return jsonify({
        'success': True,
        'holidays': [h.to_dict() for h in holidays]
    })


@app.route('/api/admin/holidays/settings', methods=['POST'])
@admin_required
def api_save_holiday_settings():
    try:
        data = request.get_json()
        start_date = data.get('academic_year_start')
        end_date = data.get('academic_year_end')
        
        if not start_date or not end_date:
            return jsonify({'success': False, 'error': 'Start and end dates required'}), 400
            
        SystemSettings.query.filter_by(setting_key='academic_year_start').delete()
        SystemSettings.query.filter_by(setting_key='academic_year_end').delete()
        
        db.session.add(SystemSettings(setting_key='academic_year_start', setting_value=start_date))
        db.session.add(SystemSettings(setting_key='academic_year_end', setting_value=end_date))
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Holidays settings updated'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/holidays/auto-generate', methods=['POST'])
@admin_required
def api_auto_generate_holidays():
    from models import Holiday, SystemSettings
    from datetime import datetime, timedelta
    
    try:
        start_setting = SystemSettings.query.get('academic_year_start')
        end_setting = SystemSettings.query.get('academic_year_end')
        academic_year = SystemSettings.get_current_academic_year()
        
        if not start_setting or not end_setting:
            return jsonify({'success': False, 'error': 'Please set academic year start and end dates first'}), 400
            
        start_date = datetime.strptime(start_setting.setting_value, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_setting.setting_value, '%Y-%m-%d').date()
        
        curr = start_date
        created_count = 0
        
        while curr <= end_date:
            # Sunday
            if curr.weekday() == 6:
                name = "Sunday"
                h_type = "sunday"
                add = True
            # Saturday
            elif curr.weekday() == 5:
                # 2nd Saturday or 4th Saturday
                day_num = curr.day
                if 8 <= day_num <= 14:
                    name = "2nd Saturday"
                    h_type = "saturday"
                    add = True
                elif 22 <= day_num <= 28:
                    name = "4th Saturday"
                    h_type = "saturday"
                    add = True
                else:
                    add = False
            else:
                add = False
                
            if add:
                existing = Holiday.query.filter_by(holiday_date=curr).first()
                if not existing:
                    new_h = Holiday(
                        holiday_date=curr,
                        holiday_name=name,
                        holiday_type=h_type,
                        academic_year=academic_year
                    )
                    db.session.add(new_h)
                    created_count += 1
            
            curr += timedelta(days=1)
            
        db.session.commit()
        return jsonify({'success': True, 'message': f'Generated {created_count} holidays automatically.'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/holidays', methods=['POST'])
@admin_required
def api_add_holiday():
    from models import Holiday, SystemSettings
    from datetime import datetime
    
    try:
        data = request.get_json()
        h_date_str = data.get('holiday_date')
        h_name = data.get('holiday_name')
        h_type = data.get('holiday_type', 'public')
        
        if not h_date_str or not h_name:
            return jsonify({'success': False, 'error': 'Date and Name required'}), 400
            
        h_date = datetime.strptime(h_date_str, '%Y-%m-%d').date()
        academic_year = SystemSettings.get_current_academic_year()
        
        existing = Holiday.query.filter_by(holiday_date=h_date).first()
        if existing:
            existing.holiday_name = h_name
            existing.holiday_type = h_type
        else:
            new_h = Holiday(
                holiday_date=h_date,
                holiday_name=h_name,
                holiday_type=h_type,
                academic_year=academic_year
            )
            db.session.add(new_h)
            
        db.session.commit()
        return jsonify({'success': True, 'message': 'Holiday saved'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/api/admin/holidays/<holiday_id>', methods=['DELETE'])
@admin_required
def api_delete_holiday(holiday_id):
    from models import Holiday
    try:
        h = Holiday.query.get(holiday_id)
        if not h:
            return jsonify({'success': False, 'error': 'Holiday not found'}), 404
            
        db.session.delete(h)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Holiday removed (Marked as working day)'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 400


@app.route('/faculty/fees')
@faculty_required
def faculty_fees_page():
    """Faculty fee management page"""
    return render_template('faculty_fee_management.html')


@app.route('/admin/fees')
@admin_required
def admin_fees_page():
    """Admin fee management and defaulter reports"""
    return render_template('admin_fees.html')


# ---------------------------------------------
# Fee Template Management Routes (Admin)
# ---------------------------------------------
@app.route('/admin/fee-templates')
@admin_required
def admin_fee_templates():
    """List all fee templates grouped by batch year"""
    from models import FeeTemplate
    from sqlalchemy import desc
    
    # Get all templates, ordered by batch year and academic year
    templates = FeeTemplate.query.filter_by(is_deleted=False).order_by(
        desc(FeeTemplate.batch_year),
        desc(FeeTemplate.academic_year),
        FeeTemplate.seat_type,
        FeeTemplate.quota_type
    ).all()
    
    # Group by batch year, then by academic year
    templates_by_batch = {}
    for template in templates:
        batch = template.batch_year
        if batch not in templates_by_batch:
            templates_by_batch[batch] = {}
        
        year = template.academic_year
        if year not in templates_by_batch[batch]:
            templates_by_batch[batch][year] = []
        templates_by_batch[batch][year].append(template)
    
    return render_template('admin_fee_templates.html', templates_by_batch=templates_by_batch)


@app.route('/admin/fee-templates/add', methods=['GET', 'POST'])
@admin_required
def admin_add_fee_template():
    """Add a new fee template"""
    from models import FeeTemplate
    
    if request.method == 'POST':
        try:
            academic_year = request.form.get('academic_year')
            batch_year = request.form.get('batch_year')
            seat_type = request.form.get('seat_type')
            quota_type = request.form.get('quota_type') if request.form.get('quota_type') else None
            base_fees = float(request.form.get('base_fees'))
            description = request.form.get('description')
            
            # Validate
            if not academic_year or not batch_year or not seat_type or base_fees < 0:
                flash('Please fill all required fields', 'error')
                return redirect(url_for('admin_add_fee_template'))
            
            # Check for duplicates
            existing = FeeTemplate.query.filter_by(
                academic_year=academic_year,
                batch_year=batch_year,
                seat_type=seat_type,
                quota_type=quota_type,
                is_deleted=False
            ).first()
            
            if existing:
                flash(f'Fee template already exists for Batch {batch_year}, Year {academic_year}, {seat_type} - {quota_type or "N/A"}', 'error')
                return redirect(url_for('admin_add_fee_template'))
            
            # Create template
            current_user = get_current_user()
            template = FeeTemplate(
                academic_year=academic_year,
                batch_year=batch_year,
                seat_type=seat_type,
                quota_type=quota_type,
                base_fees=base_fees,
                description=description,
                created_by_user_id=current_user.user_id
            )
            
            db.session.add(template)
            db.session.commit()
            
            flash(f'Fee template created successfully for Batch {batch_year}, Year {academic_year}', 'success')
            return redirect(url_for('admin_fee_templates'))
            
        except Exception as e:
            db.session.rollback()

            # debug code
            print(f'Error creating fee template: {str(e)}')
            flash(f'Error creating fee template: {str(e)}', 'error')
            return redirect(url_for('admin_add_fee_template'))
    
    return render_template('admin_fee_template_form.html', template=None)


@app.route('/admin/fee-templates/<template_id>/edit', methods=['GET', 'POST'])
@admin_required
def admin_edit_fee_template(template_id):
    """Edit an existing fee template"""
    from models import FeeTemplate
    
    template = FeeTemplate.query.filter_by(fee_template_id=template_id, is_deleted=False).first()
    if not template:
        flash('Fee template not found', 'error')
        return redirect(url_for('admin_fee_templates'))
    
    if request.method == 'POST':
        try:
            template.base_fees = float(request.form.get('base_fees'))
            template.description = request.form.get('description')
            
            db.session.commit()
            
            flash(f'Fee template updated successfully', 'success')
            return redirect(url_for('admin_fee_templates'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating fee template: {str(e)}', 'error')
    
    return render_template('admin_fee_template_form.html', template=template)


@app.route('/admin/fee-templates/<template_id>/delete', methods=['POST'])
@admin_required
def admin_delete_fee_template(template_id):
    """Soft delete a fee template"""
    from models import FeeTemplate
    
    template = FeeTemplate.query.filter_by(fee_template_id=template_id, is_deleted=False).first()
    if not template:
        return jsonify({'success': False, 'message': 'Template not found'}), 404
    
    try:
        template.is_deleted = True
        db.session.commit()
        return jsonify({'success': True, 'message': 'Template deleted successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/admin/assign-fees', methods=['GET', 'POST'])
@admin_required
def admin_assign_fees():
    """Bulk assign fees to students without fee structures"""
    from models import Student, FeeTemplate, FeeStructure
    import uuid
    
    if request.method == 'POST':
        try:
            # Get students without fee structures
            students = Student.query.filter_by(is_deleted=False).all()
            
            assigned_count = 0
            errors = []
            
            for student in students:
                # Check if student already has fee structure
                existing = FeeStructure.query.filter_by(
                    student_id=student.student_id,
                    is_deleted=False
                ).first()
                
                if existing:
                    continue
                
                # Find matching template
                template = FeeTemplate.query.filter_by(
                    academic_year=student.joining_academic_year,
                    batch_year=student.joining_academic_year,
                    seat_type=student.seat_type,
                    quota_type=student.quota_type,
                    is_deleted=False
                ).first()
                
                if not template:
                    errors.append(f"{student.roll_number}: No matching template")
                    continue
                
                # Create fee structure
                fee_structure = FeeStructure(
                    fee_structure_id=str(uuid.uuid4()),
                    student_id=student.student_id,
                    academic_year=student.joining_academic_year,
                    base_fees=template.base_fees,
                    total_fees=template.base_fees,
                    additional_fees=None,
                    fee_template_id=template.fee_template_id
                )
                
                db.session.add(fee_structure)
                assigned_count += 1
            
            db.session.commit()
            
            if assigned_count > 0:
                flash(f'Successfully assigned fees to {assigned_count} students!', 'success')
            else:
                flash('No new fee assignments needed', 'info')
            
            if errors:
                flash(f'Errors: {", ".join(errors[:5])}', 'warning')
            
            return redirect(url_for('admin_assign_fees'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error: {str(e)}', 'error')
            return redirect(url_for('admin_assign_fees'))
    
    # GET request - show students without fees
    from models import Student, FeeStructure
    
    students = Student.query.filter_by(is_deleted=False).all()
    students_without_fees = []
    
    for student in students:
        has_fees = FeeStructure.query.filter_by(
            student_id=student.student_id,
            is_deleted=False
        ).first() is not None
        
        if not has_fees:
            students_without_fees.append(student)
    
    return render_template('admin_assign_fees.html', 
                         students=students_without_fees,
                         total_students=len(students),
                         students_with_fees=len(students) - len(students_without_fees))


@app.route('/parent/dashboard')
@student_required
def parent_dashboard():
    """
    Parent Dashboard - Read Only View
    """
    if not session.get('login_as_parent'):
        return redirect(url_for('student_dashboard'))

    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    # Calculate attendance logic (Same as student dashboard)
    enrollments = StudentSubjectEnrollment.query.filter_by(student_id=student_record.student_id).all()
    subject_ids = [e.subject_id for e in enrollments]
    
    total_sessions = 0
    attended = 0
    
    for sub_id in subject_ids:
        # Get schedules for this subject
        schedules = ClassSchedule.query.filter_by(subject_id=sub_id, is_deleted=False).all()
        schedule_ids = [s.schedule_id for s in schedules]
        
        # Get sessions
        sessions = AttendanceSession.query.filter(
            AttendanceSession.schedule_id.in_(schedule_ids),
            AttendanceSession.status == 'finalized',
            AttendanceSession.is_deleted == False
        ).all()
        
        for sess in sessions:
            total_sessions += 1
            record = AttendanceRecord.query.filter_by(
                attendance_session_id=sess.attendance_session_id,
                student_id=student_record.student_id
            ).first()
            if record and record.status == 'present':
                attended += 1

    overall_percentage = round((attended / total_sessions * 100), 1) if total_sessions > 0 else 0
    
    # Red zone calculation
    is_red_zone = overall_percentage < 75
    
    # Get current checkin status
    today = date.today()
    checkin_today = CampusCheckIn.query.filter_by(
        student_id=student_record.student_id,
        checkin_date=today,
        is_deleted=False
    ).first()
    
    has_checked_in = bool(checkin_today)
    
    # Motivational Quote (for parent context, maybe just a welcome)
    import random
    quote = random.choice(MOTIVATIONAL_QUOTES)

    section = Section.query.get(student_record.section_id) if student_record.section_id else None

    dashboard_data = {
        'student': student_record,
        'section': section,
        'overall_percentage': overall_percentage,
        'total_sessions': total_sessions,
        'attended': attended,
        'is_red_zone': is_red_zone,
        'motivational_quote': quote,
        'has_checked_in': has_checked_in,
        'current_checkin': checkin_today,
        'is_parent_view': True # Flag for template
    }
    
    return render_template('parent/dashboard.html', **dashboard_data)

@app.route('/parent/profile')
@student_required
def parent_profile():
    """
    Parent View: Read-only student profile
    """
    if not session.get('login_as_parent'):
        return redirect(url_for('student_dashboard'))

    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    section = Section.query.get(student_record.section_id) if student_record.section_id else None
    
    return render_template('parent/profile.html', student=student_record, section=section)

@app.route('/parent/contact-teacher')
@student_required
def parent_contact_teacher():
    """
    Parent View: Contact Class Teacher
    """
    if not session.get('login_as_parent'):
        return redirect(url_for('student_dashboard'))

    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    section = Section.query.get(student_record.section_id) if student_record.section_id else None
    
    class_teacher = None
    if section and section.class_teacher_id:
        class_teacher = Faculty.query.get(section.class_teacher_id)
        
    return render_template('parent/contact_teacher.html', 
                         student=student_record, 
                         section=section,
                         teacher=class_teacher)

@app.route('/api/parent/student-status')
@student_required
def api_parent_student_status():
    """
    API for polling student status
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    today = date.today()
    checkin = CampusCheckIn.query.filter_by(
        student_id=student_record.student_id,
        checkin_date=today,
        is_deleted=False
    ).first()
    
    status = 'not_reported'
    if checkin:
        if checkin.checkout_time:
            status = 'checked_out'
        else:
            status = 'checked_in'
            
    return jsonify({
        'success': True,
        'status': status, # checked_in, checked_out, not_reported
        'last_updated': datetime.now().isoformat(),
        'student_name': student_record.first_name
    })



# ============================================================================
# INTERNAL MARKS SYSTEM - FACULTY ROUTES
# ============================================================================

@app.route('/faculty/internal-marks')
@login_required
def faculty_internal_marks():
    """
    Faculty Internal Marks - List all assessments created by faculty
    """
    current_user = get_current_user()
    
    # Get faculty record
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    if not faculty_record:
        flash('Faculty record not found.', 'danger')
        return redirect(url_for('login'))
    
    # Get all tests created by this faculty
    tests = Test.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        is_deleted=False
    ).order_by(Test.test_date.desc()).all()
    
    # Get faculty's subject allocations for dropdown
    allocations = SubjectAllocation.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        is_deleted=False
    ).all()
    
    return render_template('faculty/internal_marks.html', 
                         tests=tests, 
                         allocations=allocations,
                         faculty=faculty_record)


@app.route('/faculty/internal-marks/create', methods=['GET', 'POST'])
@login_required
def faculty_create_assessment():
    """
    Create new internal assessment
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        flash('Faculty record not found.', 'danger')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        try:
            subject_id = request.form.get('subject_id')
            section_id = request.form.get('section_id')
            if not section_id: # Handle empty string from form as None for DB matching
                section_id = None
            
            test_name = request.form.get('test_name')
            component_type = request.form.get('component_type', 'test')
            max_marks = float(request.form.get('max_marks', 0))
            weightage = request.form.get('weightage')
            test_date = request.form.get('test_date')
            description = request.form.get('description')
            
            # Validate that faculty teaches this subject in this section
            allocation = SubjectAllocation.query.filter_by(
                faculty_id=faculty_record.faculty_id,
                subject_id=subject_id,
                section_id=section_id,
                is_deleted=False
            ).first()
            
            if not allocation:
                flash('You are not authorized to create assessments for this subject/section.', 'danger')
                return redirect(url_for('faculty_internal_marks'))
            
            # Create new test
            new_test = Test(
                test_name=test_name,
                subject_id=subject_id,
                faculty_id=faculty_record.faculty_id,
                section_id=section_id,
                component_type=component_type,
                max_marks=max_marks,
                weightage=float(weightage) if weightage else None,
                test_date=datetime.strptime(test_date, '%Y-%m-%d').date() if test_date else None,
                description=description,
                is_published=False
            )
            
            db.session.add(new_test)
            db.session.commit()
            
            flash(f'Assessment "{test_name}" created successfully!', 'success')
            return redirect(url_for('faculty_enter_marks', test_id=new_test.test_id))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating assessment: {str(e)}', 'danger')
            return redirect(url_for('faculty_internal_marks'))
    
    # GET request - show form
    allocations = SubjectAllocation.query.filter_by(
        faculty_id=faculty_record.faculty_id,
        is_deleted=False
    ).all()
    
    # Get all potential sections for each subject
    current_year = SystemSettings.get_current_academic_year()
    
    subject_sections = {}
    for allocation in allocations:
        subject = allocation.subject
        if not subject or subject.subject_id in subject_sections:
            continue
            
        # 1. Get regular sections for this subject's semester and program
        # Regular sections match the subject's program and semester
        matching_sections = Section.query.filter_by(
            program_id=subject.program_id,
            current_semester=subject.semester_id,
            academic_year=current_year,
            is_deleted=False
        ).all()
        
        # 2. Get elective/virtual sections linked to this subject
        elective_sections = Section.query.filter_by(
            linked_subject_id=subject.subject_id,
            is_elective=True,
            academic_year=current_year,
            is_deleted=False
        ).all()
        
        # Combine and ensure unique
        all_sections = list(set(matching_sections + elective_sections))
        
        subject_sections[subject.subject_id] = [
            {
                'section_id': s.section_id,
                'section_name': s.section_name
            } for s in all_sections
        ]
    
    return render_template('faculty/create_assessment.html', 
                         allocations=allocations,
                         subject_sections=subject_sections,
                         faculty=faculty_record)


@app.route('/faculty/internal-marks/<test_id>/enter-marks')
@login_required
def faculty_enter_marks(test_id):
    """
    Enter/Edit marks for an assessment
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        flash('Faculty record not found.', 'danger')
        return redirect(url_for('login'))
    
    # Get test
    test = Test.query.get_or_404(test_id)
    
    # Verify faculty owns this test
    if test.faculty_id != faculty_record.faculty_id:
        flash('You are not authorized to edit this assessment.', 'danger')
        return redirect(url_for('faculty_internal_marks'))
    
    # Get students in the section
    section = Section.query.get(test.section_id)
    students = section.get_students()
    
    # Get existing results
    existing_results = {r.student_id: r for r in test.results.all()}
    
    # Prepare student data with marks
    student_data = []
    for student in students:
        result = existing_results.get(student.student_id)
        student_data.append({
            'student': student,
            'marks': result.marks_obtained if result else None,
            'remarks': result.remarks if result else ''
        })
    
    return render_template('faculty/enter_marks.html',
                         test=test,
                         student_data=student_data,
                         faculty=faculty_record)


@app.route('/api/faculty/internal-marks/<test_id>/save-marks', methods=['POST'])
@login_required
def api_save_marks(test_id):
    """
    Save marks for multiple students
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        return jsonify({'success': False, 'error': 'Faculty not found'}), 404
    
    test = Test.query.get_or_404(test_id)
    
    # Verify ownership
    if test.faculty_id != faculty_record.faculty_id:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    try:
        marks_data = request.json.get('marks', [])
        
        for item in marks_data:
            student_id = item.get('student_id')
            marks = item.get('marks')
            remarks = item.get('remarks', '')
            
            # Validate marks
            if marks is not None:
                marks = float(marks)
                if marks < 0 or marks > test.max_marks:
                    return jsonify({
                        'success': False, 
                        'error': f'Marks must be between 0 and {test.max_marks}'
                    }), 400
            
            # Check if result exists
            result = TestResult.query.filter_by(
                test_id=test_id,
                student_id=student_id
            ).first()
            
            if result:
                # Update existing
                result.marks_obtained = marks
                result.remarks = remarks
            else:
                # Create new
                result = TestResult(
                    test_id=test_id,
                    student_id=student_id,
                    marks_obtained=marks,
                    remarks=remarks
                )
                db.session.add(result)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Marks saved successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/faculty/internal-marks/<test_id>/publish', methods=['POST'])
@login_required
def api_publish_marks(test_id):
    """
    Toggle publish status of marks
    """
    current_user = get_current_user()
    faculty_record = Faculty.query.filter_by(user_id=current_user.user_id).first()
    
    if not faculty_record:
        return jsonify({'success': False, 'error': 'Faculty not found'}), 404
    
    test = Test.query.get_or_404(test_id)
    
    # Verify ownership
    if test.faculty_id != faculty_record.faculty_id:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    try:
        test.is_published = not test.is_published
        db.session.commit()
        
        return jsonify({
            'success': True,
            'is_published': test.is_published,
            'message': 'Marks published' if test.is_published else 'Marks unpublished'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# INTERNAL MARKS SYSTEM - STUDENT ROUTES
# ============================================================================

@app.route('/student/internal-marks')
@student_required
def student_internal_marks():
    """
    Student view of internal marks
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        flash('Student record not found.', 'danger')
        return redirect(url_for('login'))
    
    # Get student's section
    section = Section.query.get(student_record.section_id)
    
    # Get all published tests for student's subjects
    # Get compulsory subjects
    compulsory_subjects = Subject.query.filter(
        Subject.program_id == section.program_id,
        Subject.semester_id == section.current_semester,
        Subject.subject_category == 'compulsory'
    ).all()
    
    # Get enrolled electives
    enrollments = StudentSubjectEnrollment.query.filter_by(
        student_id=student_record.student_id,
        semester=section.current_semester,
        is_deleted=False
    ).all()
    enrolled_subjects = [Subject.query.get(e.subject_id) for e in enrollments]
    
    all_subjects = compulsory_subjects + enrolled_subjects
    
    # Get marks for each subject
    subject_marks = {}
    for subject in all_subjects:
        # Find target section (main or elective)
        target_section_id = student_record.section_id
        enrollment = next((e for e in enrollments if e.subject_id == subject.subject_id), None)
        if enrollment and enrollment.section_id:
            target_section_id = enrollment.section_id
        
        # Get published tests for this subject/section
        tests = Test.query.filter_by(
            subject_id=subject.subject_id,
            section_id=target_section_id,
            is_published=True,
            is_deleted=False
        ).order_by(Test.test_date.desc()).all()
        
        # Get student's results
        marks_list = []
        for test in tests:
            result = TestResult.query.filter_by(
                test_id=test.test_id,
                student_id=student_record.student_id
            ).first()
            
            marks_list.append({
                'test': test,
                'marks': result.marks_obtained if result else None,
                'percentage': round((result.marks_obtained / test.max_marks * 100), 2) if result and result.marks_obtained is not None else None
            })
        
        if marks_list:
            subject_marks[subject.subject_id] = {
                'subject': subject,
                'marks': marks_list
            }
    
    return render_template('student/internal_marks.html',
                         subject_marks=subject_marks,
                         student=student_record,
                         section=section)


@app.route('/api/student/internal-marks')
@student_required
def api_student_internal_marks():
    """
    API endpoint for student internal marks
    """
    current_user = get_current_user()
    student_record = Student.query.filter_by(user_id=current_user.user_id).first()
    
    if not student_record:
        return jsonify({'success': False, 'error': 'Student not found'}), 404
    
    # Similar logic as above but return JSON
    # Implementation similar to student_internal_marks route
    return jsonify({'success': True, 'message': 'API endpoint for future use'})



if __name__ == '__main__':
    # Create database tables if they don't exist
    init_db()
    
    # Run the Flask application
    # debug=True for development, set to False in production
    print("\n" + "="*50)
    print("  BCA BUB Attendance System")
    print("  Mobile-First PWA Application")
    print("="*50)
    print("\nðŸ“± Access the app at:")
    print("   â€¢ Local: http://localhost:5000")
    print("   â€¢ Mobile: http://YOUR_IP:5000")
    print("\nðŸ”§ Running in DEBUG mode")
    print("="*50 + "\n")
    
    app.run(debug=True, host='0.0.0.0', port=5000)
# Faculty Check-In/Checkout API Endpoints
# Add these to app.py# ============================================================================
